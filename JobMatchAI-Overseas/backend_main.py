"""
JobMatchAI - 智能求职助手
支持：中国、北欧、全球市场
核心功能：简历精修 + 邮件职位聚合 + 智能求职信 + 职位匹配

Copyright © 2026 JobMatchAI. All rights reserved.
"""
from fastapi import FastAPI, UploadFile, File, Form, Request, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict
import uvicorn
import os
import re
import io
import json
import jwt
import unicodedata
import asyncio
import threading

# 学习推荐引擎
from learning_recommendations import router as learning_router

# ===== AI驱动的Skill Gap分析（完全由AI完成，不使用硬编码）=====

def analyze_skill_gap_ai(resume_text: str, job_description: str, job_title: str = '', detected_skills: List[str] = None, lang: str = 'en') -> Dict:
    """
    使用AI分析简历与职位的技能差距
    完全由AI驱动，不使用硬编码关键词匹配
    """
    if not AI_AVAILABLE:
        return {
            'missing_skills': [],
            'matched_skills': detected_skills or [],
            'critical_gaps': [],
            'score': 50,
            'reasoning': 'AI unavailable'
        }
    
    try:
        # 构建简历技能列表（用于上下文）
        resume_skills_str = ', '.join(detected_skills) if detected_skills else '从简历内容中提取'
        
        prompts = {
            'en': f"""You are an expert technical recruiter. Analyze the skill gap between this resume and job description.

JOB TITLE: {job_title if job_title else 'Not specified'}

JOB DESCRIPTION:
{job_description[:2000]}

RESUME SKILLS (extracted): {resume_skills_str}

RESUME CONTENT:
{resume_text[:2000]}

Analyze thoroughly and return ONLY valid JSON:
{{
    "score": 数字(1-100, match score based on skill alignment),
    "matched_skills": ["skills the resume has that match the job"],
    "missing_skills": ["critical skills required by job but missing in resume - be specific and accurate, only list skills actually mentioned in job description"],
    "critical_gaps": ["⚠️ MUST-HAVE requirements that will likely cause rejection if missing - focus on explicit requirements like certifications, degrees, years of experience, specific tools/technologies"],
    "reasoning": "brief explanation of why these gaps matter for this specific role"
}}

CRITICAL RULES:
1. Only list skills in missing_skills if they are EXPLICITLY mentioned in the job description
2. Do NOT guess or hallucinate skills (e.g., do NOT list Power BI if job doesn't mention Power BI)
3. Focus on what the job actually REQUIRES vs what would be nice to have
4. Look for: specific tools named, certifications required, years of experience, language requirements""",
            
            'zh': f"""你是专业的技术招聘顾问。分析简历与职位的技能差距。

职位名称：{job_title if job_title else '未指定'}

职位描述：
{job_description[:2000]}

简历已检测技能：{resume_skills_str}

简历内容：
{resume_text[:2000]}

仔细分析后返回JSON：
{{
    "score": 数字(1-100, 技能匹配度),
    "matched_skills": ["简历中与职位匹配的技能"],
    "missing_skills": ["职位明确要求但简历中缺少的技能 - 要准确，只列出职位描述中提到的技能"],
    "critical_gaps": ["⚠️ 关键要求 - 缺失会导致被拒的硬性要求，如认证、学位、工作年限、特定工具],
    "reasoning": "简要说明为什么这些差距对这个职位重要"
}}

关键规则：
1. 只在 missing_skills 中列出职位描述中明确提到的技能
2. 不要猜测或编造技能（如：职位没提Power BI就不要列）
3. 关注职位实际要求的是什么 vs 最好有的
4. 查找：明确提到的工具、需要的认证、工作年限要求、语言要求"""
        }
        
        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a job analyzer. Return ONLY valid JSON, no markdown, no explanation."},
                {"role": "user", "content": prompts.get(lang, prompts['en'])}
            ],
            preferred_provider="groq",
            temperature=0.1,
            max_tokens=600
        )
        
        content = response.choices[0].message.content.strip()
        # 清理可能的markdown格式
        if content.startswith('```'):
            content = content.split('```')[1]
            if content.startswith('json'):
                content = content[4:]
        
        result = json.loads(content)
        
        return {
            'missing_skills': result.get('missing_skills', [])[:5],
            'matched_skills': result.get('matched_skills', []),
            'critical_gaps': result.get('critical_gaps', []),
            'score': result.get('score', 50),
            'reasoning': result.get('reasoning', '')
        }
        
    except Exception as e:
        print(f"AI Skill Gap analysis failed: {e}")
        return {
            'missing_skills': [],
            'matched_skills': detected_skills or [],
            'critical_gaps': [],
            'score': 50,
            'reasoning': f'Analysis failed: {str(e)}'
        }


# ===== AI驱动的职位分析函数 =====

def identify_critical_requirements(job_description: str, resume_text: str, lang: str = 'en') -> List[str]:
    """
    使用AI分析职位描述，识别可能导致申请失败的关键要求
    与简历对比，标记缺失的硬技能
    """
    if not AI_AVAILABLE:
        # 后备：使用规则匹配
        return identify_critical_requirements_fallback(job_description, resume_text, lang)
    
    try:
        prompts = {
            'zh': f"""你是一位专业的招聘顾问。请分析以下职位描述，识别所有关键要求，并与简历对比。

职位描述：
{job_description[:1500]}

简历摘要：
{resume_text[:500]}

请返回JSON数组格式：
{{"critical": ["关键要求1（必须满足否则极可能被拒）", "关键要求2"], "soft_skills": ["建议添加到简历的软技能1"]}}

关键要求包括：
1. 硬性技能要求（如：日语 fluency、特定软件认证、学位要求）
2. 语言要求（如：日语 fluent、普通话 native）
3. 工作经验年限要求
4. 特定工具/软件要求

请只返回JSON，不要其他文字。""",
            'en': f"""You are a professional recruiter. Analyze this job description and identify ALL critical requirements, then compare with the resume.

Job Description:
{job_description[:1500]}

Resume Summary:
{resume_text[:500]}

Return JSON array format:
{{"critical": ["critical requirement 1 (may cause rejection if missing)", "critical requirement 2"], "soft_skills": ["soft skill to add to resume 1"]}}

Critical requirements include:
1. Hard skills (e.g., Japanese fluency, specific software certifications, degree)
2. Language requirements (e.g., Japanese fluent, Mandarin native)
3. Years of experience requirements
4. Specific tools/software requirements

Return ONLY JSON, no other text."""
        }
        
        prompt = prompts.get(lang, prompts['en'])
        
        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a professional job analyzer. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            preferred_provider="groq",
            temperature=0.2,
            max_tokens=500
        )
        
        content = response.choices[0].message.content.strip()
        # 清理可能的markdown格式
        if content.startswith('```'):
            content = content.split('```')[1]
            if content.startswith('json'):
                content = content[4:]
        
        result = json.loads(content)
        
        critical_list = []
        for req in result.get('critical', []):
            # 标记为关键要求
            if 'japanese' in req.lower() or '日语' in req:
                critical_list.append(f"⚠️ {req} - 关键语言要求！")
            elif 'language' in req.lower() or '语言' in req or 'fluent' in req.lower():
                critical_list.append(f"⚠️ {req}")
            elif 'modelling' in req.lower() or 'modeling' in req.lower() or 'simulation' in req.lower():
                critical_list.append(f"⚠️ {req} - 关键技能要求！")
            elif 'degree' in req.lower() or '学位' in req:
                critical_list.append(f"📋 {req}")
            else:
                critical_list.append(f"⚠️ {req}")
        
        for skill in result.get('soft_skills', []):
            critical_list.append(f"💬 {skill} - 建议添加到简历")
        
        return critical_list
        
    except Exception as e:
        print(f"Critical requirements analysis failed: {e}")
        return identify_critical_requirements_fallback(job_description, resume_text, lang)


def identify_critical_requirements_fallback(job_description: str, resume_text: str, lang: str = 'en') -> List[str]:
    """后备：使用规则匹配识别关键要求"""
    job_lower = job_description.lower()
    resume_lower = resume_text.lower()
    critical = []
    
    # 语言要求
    languages = {
        'japanese': '日语 (Japanese)', 'japan': '日语',
        'mandarin': '普通话 (Mandarin)', 'english': '英语 (English)',
        'danish': '丹麦语 (Danish)', 'chinese': '中文 (Chinese)',
        'german': '德语 (German)', 'french': '法语 (French)'
    }
    for lang_key, lang_name in languages.items():
        if lang_key in job_lower:
            context_start = max(0, job_lower.find(lang_key) - 50)
            context = job_lower[context_start:context_start + 100]
            if any(w in context for w in ['fluent', 'native', 'proficient', 'communicate']):
                if lang_key not in resume_lower:
                    critical.append(f"⚠️ {lang_name} - 关键语言要求！")
    
    # 建模/仿真工具
    if any(kw in job_lower for kw in ['modelling', 'modeling', 'simulation', 'simulations']):
        if 'modelling' not in resume_lower and 'modeling' not in resume_lower and 'simulation' not in resume_lower:
            critical.append("⚠️ 建模与仿真工具 - 关键技能要求！")
    
    # 学位要求
    if 'degree level' in job_lower or 'equivalent' in job_lower:
        critical.append("📋 学位要求 - 请确认学历符合")
    
    # 领导经验
    if any(kw in job_lower for kw in ['leading teams', 'multidisciplinary', 'lead team']):
        critical.append("💬 团队领导经验 - 建议在简历中强调")
    
    # 软技能建议
    soft_skills = {
        'communication': '沟通能力',
        'relationships': '关系建立',
        'stakeholder': '利益相关者管理'
    }
    for skill_key, skill_name in soft_skills.items():
        if skill_key in job_lower:
            if skill_key not in resume_lower:
                critical.append(f"💬 {skill_name} - 建议添加到简历")
    
    return critical


# ===== AI驱动的国内学习资源推荐 =====

def get_chinese_learning_resources_ai(skills: List[str], lang: str = 'zh') -> List[Dict]:
    """
    使用AI根据技能列表动态获取国内学习资源
    
    Args:
        skills: 需要学习的技能列表
        lang: 语言偏好
    
    Returns:
        国内学习资源列表（B站、中国大学MOOC）
    """
    if not skills:
        return []
    
    if not AI_AVAILABLE:
        return get_chinese_learning_resources_fallback(skills)
    
    skills_str = '、'.join(skills[:5])  # 最多5个技能
    
    prompt = f"""你需要为以下技能推荐B站学习资源。

【目标技能】
{skills_str}

【重要规则 - 必须遵守！】
1. 禁止推荐以下平台（课程经常下线/已下线）：
   - 腾讯课堂
   - 网易公开课
   - 学堂在线（xuetangX）
   - 中国大学MOOC（icourse163.org）
   只使用B站搜索链接！

【必须返回的URL格式 - B站搜索链接】
https://search.bilibili.com/all?keyword=技能关键词

例如：
- 财务咨询技能 → https://search.bilibili.com/all?keyword=财务咨询教程
- 客户沟通技能 → https://search.bilibili.com/all?keyword=客户沟通技巧

【返回JSON格式】（只返回JSON，不要其他内容）：
{{
    "chinese_resources": [
        {{
            "skill": "技能中文名",
            "title": "B站学习合集推荐（如：财务咨询系统性学习教程）",
            "provider": "B站",
            "url": "https://search.bilibili.com/all?keyword=技能关键词",
            "duration": "自学时长",
            "difficulty": "初级/中级/高级",
            "type": "视频合集"
        }}
    ]
}}

【规则 - 严格遵守】
1. 所有URL必须使用：https://search.bilibili.com/all?keyword=技能名
2. 禁止返回任何具体课程页面URL（如 icourse163.org/xxx）
3. 每个技能只推荐1个B站搜索链接
4. 技能名称用中文，标题描述要具体
5. 只返回JSON，不要任何解释文字！"""

    try:
        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a helpful assistant that returns ONLY valid JSON. No markdown, no explanations, no text outside the JSON."},
                {"role": "user", "content": prompt}
            ],
            preferred_provider="groq",
            temperature=0.3,
            max_tokens=1200
        )
        
        content = response.choices[0].message.content.strip()
        # 清理可能的markdown包装
        if content.startswith('```'):
            content = content.split('```')[1]
            if content.startswith('json'):
                content = content[4:]
        
        result = json.loads(content)
        chinese_resources = result.get('chinese_resources', [])
        
        # 格式化返回
        formatted = []
        for r in chinese_resources[:5]:
            formatted.append({
                'skill': r.get('skill', ''),
                'title': r.get('title', ''),
                'provider': r.get('provider', ''),
                'url': r.get('url', ''),
                'duration': r.get('duration', ''),
                'difficulty': r.get('difficulty', ''),
                'type': r.get('type', '视频课程'),
                'relevance': f"与「{r.get('skill', '')}」相关的中文学习资源"
            })
        
        return formatted
        
    except Exception as e:
        print(f"AI Chinese learning resources failed: {e}")
        return get_chinese_learning_resources_fallback(skills)


def get_chinese_learning_resources_fallback(skills: List[str]) -> List[Dict]:
    """后备：使用B站搜索链接获取国内学习资源（不包含具体课程页URL）"""
    results = []
    
    # 只使用B站搜索链接，避免课程下线问题
    bilibili_base = "https://search.bilibili.com/all?keyword="
    
    for skill in skills[:5]:
        results.append({
            'skill': skill,
            'title': f'{skill} - B站视频合集',
            'provider': 'B站',
            'url': f'{bilibili_base}{skill}',
            'duration': '视频合集',
            'difficulty': '各水平',
            'type': '视频课程',
            'relevance': f'在B站搜索"{skill}"相关视频'
        })
    
    return results[:5]

def recommend_learning_resources(
    resume_structure: Dict,
    job_description: str,
    skill_gap: Dict,
    lang: str = 'en'
) -> Dict:
    """
    使用AI根据简历结构和技能差距推荐个性化学习资源
    
    Args:
        resume_structure: 解析后的简历结构化数据
        job_description: 职位描述
        skill_gap: 技能差距分析结果
        lang: 语言 (zh/en)
    
    Returns:
        学习资源推荐，包含免费和付费资源，以及国内资源
    """
    if not AI_AVAILABLE:
        return recommend_learning_resources_fallback(skill_gap, lang)
    
    try:
        # 提取简历中的技能和教育背景
        resume_skills = resume_structure.get('skills', {})
        technical_skills = resume_skills.get('technical', [])
        soft_skills = resume_skills.get('soft', [])
        education = resume_structure.get('education', [])
        total_years = resume_structure.get('total_experience_years', 0)
        
        # 获取缺失的技能
        missing_skills = skill_gap.get('missing_skills', [])
        critical_gaps = skill_gap.get('critical_gaps', [])
        matched_skills = skill_gap.get('matched_skills', [])
        
        # 构建上下文
        education_str = ', '.join([
            f"{e.get('degree', '')} {e.get('field', '')} at {e.get('school', '')}"
            for e in education[:2]
        ]) if education else 'Not specified'
        
        # 中文prompt - 优化
        missing_skills_str = '、'.join(missing_skills[:5]) if missing_skills else '无'
        critical_gaps_str = '、'.join(critical_gaps[:3]) if critical_gaps else '无'
        
        prompt_zh = f"""你是一名专业的职业发展顾问。为候选人推荐个性化学习资源，帮助他们弥补技能差距。

【候选人背景】
- 学历：{education_str}
- 工作经验：{total_years}年
- 已掌握技能：{'、'.join(technical_skills[:10]) if technical_skills else '未指定'}

【目标职位】
{skill_gap.get('job_title', '目标职位')}

【职位要求摘要】
{job_description[:1000]}

【需要提升的技能】
缺失技能：{missing_skills_str}
关键差距：{critical_gaps_str}

请返回JSON格式推荐（只返回JSON，不要其他内容）：
{{
    "summary": "用1-2句话说明推荐理由，中文，简洁",
    "free_resources": [
        {{
            "skill": "对应的技能名称（中文）",
            "title": "课程名称（英文原名保留）",
            "type": "课程/视频/文档",
            "provider": "平台名称，如Coursera/edX/Udemy等",
            "url": "使用平台搜索页面或官方学习路径页面URL，例如：https://www.coursera.org/search?query=关键词 或 https://learn.microsoft.com/zh-cn/技能名",
            "duration": "预计学习时长，如8小时、4周",
            "difficulty": "初级/中级/高级",
            "relevance": "一句话说明为什么这个资源有用（中文）"
        }}
    ],
    "paid_resources": [
        {{
            "skill": "对应的技能名称",
            "title": "课程名称",
            "type": "课程/认证",
            "provider": "平台名称",
            "url": "使用平台主页或课程列表页URL",
            "duration": "学习时长",
            "cost": "费用估算",
            "value_proposition": "付费价值说明（中文）"
        }}
    ],
    "quick_wins": ["更新LinkedIn技能标签", "搜索目标公司员工获取 Insights", "准备STAR法则面试案例"]
}}

【重要规则】
1. 所有返回内容必须使用中文
2. 只推荐与职位要求直接相关的技能资源
3. 优先选择知名平台（Coursera、edX、Udemy、LinkedIn Learning、Microsoft Learn）的免费课程
4. URL必须使用平台搜索页或官方学习路径页，不要使用具体课程页面URL
5. 免费资源3-5个，付费资源0-2个
6. quick_wins必须是纯中文的可执行行动建议，禁止包含英文单词或短语
        """
        prompt_en = f"""You are a professional career advisor. Recommend personalized learning resources to help the candidate address skill gaps.

CANDIDATE PROFILE:
- Education: {education_str}
- Experience: {total_years} years
- Current Skills: {', '.join(technical_skills[:10]) if technical_skills else 'Not specified'}

JOB TITLE: {skill_gap.get('job_title', 'Target Position')}

JOB REQUIREMENTS:
{job_description[:1000]}

SKILL GAPS TO ADDRESS:
- Missing: {', '.join(missing_skills[:5]) if missing_skills else 'None'}
- Critical: {', '.join(critical_gaps[:3]) if critical_gaps else 'None'}

Return ONLY valid JSON (no markdown, no explanation):
{{
    "summary": "1-2 sentences explaining the recommendation rationale",
    "free_resources": [
        {{
            "skill": "Skill name",
            "title": "Course name (keep original English name)",
            "type": "course/video/docs",
            "provider": "Platform (Coursera/edX/Udemy/LinkedIn Learning)",
            "url": "ACTUAL accessible course URL (DO NOT make up links)",
            "duration": "Time needed",
            "difficulty": "beginner/intermediate/advanced",
            "relevance": "Why this helps"
        }}
    ],
    "paid_resources": [
        {{
            "skill": "Skill name",
            "title": "Course name",
            "provider": "Platform",
            "url": "Course URL",
            "duration": "Time needed",
            "cost": "Cost estimate",
            "value_proposition": "Why worth it"
        }}
    ],
    "quick_wins": ["Quick action 1", "Quick action 2"]
}}

CRITICAL RULES:
1. Only recommend for skills MENTIONED in job requirements
2. Prioritize well-known platforms (Coursera, edX, Udemy, LinkedIn Learning)
3. ALWAYS use platform search page URLs or official learning path URLs - NEVER use specific course page URLs (they expire)
   - Coursera: https://www.coursera.org/search?query=skill_name
   - Udemy: https://www.udemy.com/topic/skill_name/
   - LinkedIn Learning: https://www.linkedin.com/learning/topics/skill_name
   - Microsoft Learn: https://learn.microsoft.com/en-us/training/
4. Return 3-5 free resources, 0-2 paid resources"""
        
        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a learning resource curator. Return ONLY valid JSON, no markdown, no explanation. NEVER make up URLs - only provide links you are confident exist."},
                {"role": "user", "content": prompt_zh if lang == 'zh' else prompt_en}
            ],
            preferred_provider="groq",
            temperature=0.3,
            max_tokens=2000
        )
        
        content = response.choices[0].message.content.strip()
        if content.startswith('```'):
            content = content.split('```')[1]
            if content.startswith('json'):
                content = content[4:]
        
        result = json.loads(content)
        
        # 获取国内学习资源（AI驱动）
        chinese_resources = get_chinese_learning_resources_ai(missing_skills, lang)
        
        # 构建 recommendations 数组格式（兼容前端）
        recommendations = []
        free_resources = result.get('free_resources', [])[:5]
        paid_resources = result.get('paid_resources', [])[:3]
        quick_wins = result.get('quick_wins', [])
        
        # 按技能分组
        for skill in missing_skills[:5]:
            skill_lower = skill.lower()
            # 找对应的免费资源
            skill_free = [r for r in free_resources if skill_lower in r.get('skill', '').lower()]
            # 找对付费的资源
            skill_paid = [r for r in paid_resources if skill_lower in r.get('skill', '').lower()]
            # 找B站资源
            skill_chinese = [r for r in chinese_resources if skill_lower in r.get('skill', '').lower()]
            
            # 分配资源（确保每个技能都有展示机会）
            if not skill_free:
                skill_free = free_resources[:1] if free_resources else []
            if not skill_paid:
                skill_paid = paid_resources[:1] if paid_resources else []
            
            if skill_free or skill_paid or skill_chinese:
                recommendations.append({
                    'skill': skill,
                    'category': '通用',
                    'free': skill_free,
                    'paid': skill_paid,
                    'chinese': skill_chinese,
                    'ai_generated': True
                })
        
        # 如果没有按技能分组的结果，直接添加资源
        if not recommendations:
            recommendations.append({
                'skill': missing_skills[0] if missing_skills else '通用技能',
                'category': '通用',
                'free': free_resources,
                'paid': paid_resources,
                'chinese': chinese_resources,
                'ai_generated': True
            })
        
        return {
            'success': True,
            'summary': result.get('summary', ''),
            'ai_summary': result.get('summary', ''),
            'free_resources': free_resources,
            'paid_resources': paid_resources,
            'chinese_resources': chinese_resources,
            'quick_wins': quick_wins,
            'skill_gaps_addressed': missing_skills[:5],
            'recommendations': recommendations  # 前端期望的格式
        }
        
    except Exception as e:
        print(f"AI Learning recommendation failed: {e}")
        return recommend_learning_resources_fallback(skill_gap, lang)


def recommend_learning_resources_fallback(skill_gap: Dict, lang: str = 'en') -> Dict:
    """后备：使用硬编码的通用学习资源"""
    missing_skills = skill_gap.get('missing_skills', [])
    
    # 通用免费资源
    generic_free = [
        {
            "skill": "General",
            "title": "Coursera Free Courses",
            "type": "course",
            "provider": "Coursera",
            "url": "https://www.coursera.org/search",
            "duration": "Varies",
            "difficulty": "beginner",
            "relevance": "Wide range of professional courses"
        },
        {
            "skill": "General",
            "title": "LinkedIn Learning",
            "type": "video",
            "provider": "LinkedIn",
            "url": "https://www.linkedin.com/learning/",
            "duration": "Self-paced",
            "difficulty": "all levels",
            "relevance": "Professional development courses"
        },
        {
            "skill": "Technical",
            "title": "MDN Web Docs",
            "type": "documentation",
            "provider": "Mozilla",
            "url": "https://developer.mozilla.org/",
            "duration": "Self-paced",
            "difficulty": "all levels",
            "relevance": "Web development documentation"
        }
    ]
    
    # 构建 recommendations 格式
    recommendations = [
        {
            'skill': skill,
            'category': '通用',
            'free': generic_free,
            'paid': [],
            'chinese': [],
            'ai_generated': False
        }
        for skill in (missing_skills[:3] if missing_skills else ['通用技能'])
    ]
    
    return {
        'success': True,
        'summary': 'Showing general resources. AI-powered recommendations unavailable.',
        'ai_summary': 'Showing general resources. AI-powered recommendations unavailable.',
        'free_resources': generic_free,
        'paid_resources': [],
        'chinese_resources': [],
        'quick_wins': ['Add missing keywords to resume', 'Quantify achievements'],
        'priority_order': missing_skills[:5] if missing_skills else ['Update your skills section'],
        'skill_gaps_addressed': missing_skills[:5],
        'recommendations': recommendations
    }


def extract_required_skills_from_job(job_description: str, lang: str = 'en') -> List[str]:
    """从职位描述中提取所有要求的技术技能"""
    if not AI_AVAILABLE:
        return extract_required_skills_fallback(job_description, lang)
    
    try:
        prompts = {
            'zh': f"""从以下职位描述中提取所有要求的技能和资格。

职位描述：
{job_description[:1500]}

请以JSON数组格式返回技能列表：
["技能1", "技能2", ...]

只返回技能名称，不要其他文字。""",
            'en': f"""Extract ALL required skills and qualifications from this job description.

Job Description:
{job_description[:1500]}

Return JSON array of skills:
["skill 1", "skill 2", ...]

Return ONLY JSON array, no other text."""
        }
        
        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a job analyzer. Return only valid JSON array."},
                {"role": "user", "content": prompts.get(lang, prompts['en'])}
            ],
            preferred_provider="groq",
            temperature=0.1,
            max_tokens=300
        )
        
        content = response.choices[0].message.content.strip()
        if content.startswith('```'):
            content = content.split('```')[1]
            if content.startswith('json'):
                content = content[4:]
        
        return json.loads(content)
        
    except Exception as e:
        print(f"Skill extraction failed: {e}")
        return extract_required_skills_fallback(job_description, lang)


def extract_required_skills_fallback(job_description: str, lang: str = 'en') -> List[str]:
    """后备：使用规则提取技能"""
    job_lower = job_description.lower()
    skills = []
    found = set()
    
    skill_patterns = {
        'Excel': ['excel'],
        'PowerPoint': ['powerpoint', 'power point'],
        'Power BI': ['power bi'],
        'Microsoft Copilot': ['copilot'],
        'Microsoft 365': ['microsoft 365', 'm365'],
        'Azure AI': ['azure ai', 'azure cognitive'],
        'AI Agent': ['ai agent', 'agentic'],
        'Machine Learning': ['machine learning', 'ml/ai', 'machine-learning'],
        'Pre-sales': ['pre-sales', 'presales'],
        'Technical Consulting': ['technical consulting'],
        'Solution Design': ['solution design'],
        'SQL': ['sql', 'mysql', 'postgresql'],
        'Python': ['python'],
        'ERP': ['erp', 'netsuite', 'dynamics', 'sap'],
        'Java': ['java'],
        'JavaScript': ['javascript'],
        'Tableau': ['tableau'],
        '建模': ['modelling', 'modeling'],
        '仿真': ['simulation', 'simulations'],
        '供应链': ['supply chain', 'SCM'],
        '财务': ['finance', 'financial', 'accounting'],
        '咨询': ['consulting', 'advisory'],
        '战略': ['strategy', 'strategic'],
        '项目管理': ['project management'],
        '数据分析': ['data analysis', 'analytics'],
    }
    
    for skill_name, patterns in skill_patterns.items():
        for pattern in patterns:
            if pattern in job_lower and skill_name.lower() not in found:
                skills.append(skill_name)
                found.add(skill_name.lower())
                break
    
    return skills

# AI 配置 - 质量优先：简历用GPT-4o，求职信用GPT-4o-mini
# 首先尝试从 .env 文件加载环境变量
def load_env_file():
    """从 .env 文件加载环境变量"""
    env_path = os.path.join(os.path.dirname(__file__), '.env')
    if os.path.exists(env_path):
        with open(env_path, 'r') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, value = line.split('=', 1)
                    os.environ[key] = value

# 加载 .env 文件
load_env_file()

# ===== 智能 AI 路由系统 =====
# 优先级：Groq（免费）→ OpenAI（备用）→ Fallback

try:
    from openai import OpenAI
    
    # AI 客户端和模型配置
    openai_api_key = os.getenv("OPENAI_API_KEY")
    groq_api_key = os.getenv("GROQ_API_KEY")
    
    # 两个客户端都初始化
    openai_client = None
    groq_client = None
    
    if openai_api_key:
        openai_client = OpenAI(api_key=openai_api_key)
        OPENAI_MODEL = os.getenv("AI_MODEL_RESUME", "gpt-4o-mini")
        print(f"✅ [OpenAI] Ready | Model: {OPENAI_MODEL}")
    
    if groq_api_key:
        groq_client = OpenAI(
            api_key=groq_api_key,
            base_url="https://api.groq.com/openai/v1"
        )
        GROQ_MODEL = "llama-3.3-70b-versatile"
        print(f"✅ [Groq] Ready | Model: {GROQ_MODEL}")
    
    # 状态跟踪
    GROQ_AVAILABLE = True  # 初始假设 Groq 可用
    GROQ_COOLDOWN_UNTIL = 0  # Unix timestamp (0 = 无限制)
    OPENAI_MODEL_CURRENT = OPENAI_MODEL if openai_client else None
    
    # 历史记录
    AI_REQUEST_LOG = []  # 最近的请求记录，用于调试
    
    def smart_ai_request(messages: list, preferred_provider: str = "groq", 
                          temperature: float = 0.1, max_tokens: int = 2000) -> dict:
        """
        智能 AI 请求路由
        
        策略：
        1. 优先尝试 Groq（免费额度）
        2. Groq 失败（429）时自动切换 OpenAI
        3. 记录失败，下次跳过 Groq
        
        Args:
            messages: OpenAI 格式的消息列表
            preferred_provider: "groq" 或 "openai"
            temperature: 温度参数
            max_tokens: 最大 token 数
        
        Returns:
            AI 响应对象
        """
        global GROQ_AVAILABLE, GROQ_COOLDOWN_UNTIL, AI_REQUEST_LOG
        import time
        
        current_provider = "groq" if preferred_provider == "groq" else "openai"
        models = {
            "groq": (groq_client, GROQ_MODEL) if groq_client else None,
            "openai": (openai_client, OPENAI_MODEL_CURRENT) if openai_client else None
        }
        
        tried_groq = False
        last_error = None
        
        # 尝试顺序：Groq → OpenAI
        providers_to_try = []
        if preferred_provider == "groq" and groq_client and GROQ_AVAILABLE:
            if GROQ_COOLDOWN_UNTIL == 0 or time.time() >= GROQ_COOLDOWN_UNTIL:
                providers_to_try = ["groq", "openai"]
            else:
                # Groq 在冷却中，直接用 OpenAI
                providers_to_try = ["openai"]
        else:
            providers_to_try = ["openai"]
        
        for provider in providers_to_try:
            client, model = models[provider]
            if not client:
                continue
            
            try:
                response = client.chat.completions.create(
                    model=model,
                    messages=messages,
                    temperature=temperature,
                    max_tokens=max_tokens
                )
                
                # 成功！记录并返回
                AI_REQUEST_LOG.append({
                    "time": time.time(),
                    "provider": provider,
                    "model": model,
                    "success": True
                })
                if len(AI_REQUEST_LOG) > 20:
                    AI_REQUEST_LOG.pop(0)
                
                return response
                
            except Exception as e:
                error_str = str(e)
                last_error = error_str
                
                if "429" in error_str or "rate_limit" in error_str.lower():
                    # Groq 配额用完
                    if provider == "groq":
                        GROQ_AVAILABLE = False
                        # 从错误信息中提取恢复时间
                        import re
                        match = re.search(r'try again in ([\d]+)h?(\d+)m?(\d+\.\d+)s?', error_str)
                        if match:
                            hours = int(match.group(1)) if match.group(1) else 0
                            mins = int(match.group(2)) if match.group(2) else 0
                            secs = float(match.group(3)) if match.group(3) else 0
                            GROQ_COOLDOWN_UNTIL = time.time() + hours*3600 + mins*60 + secs + 60
                            print(f"⚠️ Groq 配额用完，切换到 OpenAI。预计恢复: {hours}h {mins}m后")
                        else:
                            GROQ_COOLDOWN_UNTIL = time.time() + 3600  # 默认1小时
                            print(f"⚠️ Groq 配额用完，切换到 OpenAI。预计1小时后恢复")
                        continue  # 尝试 OpenAI
                
                # 其他错误，记录并继续尝试下一个 provider
                print(f"⚠️ {provider.upper()} 请求失败: {error_str[:100]}")
                continue
        
        # 全部失败，返回模拟错误
        class FakeError:
            choices = [type('obj', (object,), {'message': type('obj', (object,), {
                'content': f'AI service unavailable. Last error: {last_error}'
            })()})()]
            error = last_error
        return FakeError()
    
    def check_groq_recovery() -> bool:
        """
        检测 Groq 配额是否恢复
        每5分钟自动调用
        """
        global GROQ_AVAILABLE, GROQ_COOLDOWN_UNTIL
        import time
        
        if not groq_client or not groq_api_key:
            return False
        
        if GROQ_COOLDOWN_UNTIL > 0 and time.time() < GROQ_COOLDOWN_UNTIL:
            return False  # 还在冷却中
        
        # 测试 Groq 是否可用
        try:
            test_response = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role": "user", "content": "Hi"}],
                max_tokens=5
            )
            GROQ_AVAILABLE = True
            GROQ_COOLDOWN_UNTIL = 0
            print(f"✅ Groq 配额已恢复！")
            return True
        except Exception as e:
            if "429" in str(e):
                GROQ_AVAILABLE = False
            return False
    
    def get_ai_status() -> dict:
        """获取当前 AI 状态"""
        import time
        return {
            "groq_available": GROQ_AVAILABLE,
            "groq_cooldown_until": GROQ_COOLDOWN_UNTIL if not GROQ_AVAILABLE else 0,
            "groq_cooldown_remaining": max(0, int(GROQ_COOLDOWN_UNTIL - time.time())) if GROQ_COOLDOWN_UNTIL > time.time() else 0,
            "openai_available": openai_client is not None,
            "active_provider": "groq" if GROQ_AVAILABLE and groq_client else "openai"
        }
    
    # 兼容旧代码
    ai_client = groq_client if groq_client else openai_client
    AI_PROVIDER = "groq" if groq_client else ("openai" if openai_client else None)
    AI_MODEL_RESUME = GROQ_MODEL if groq_client else OPENAI_MODEL
    AI_MODEL_COVER = AI_MODEL_RESUME
    AI_MODEL = AI_MODEL_RESUME
    AI_AVAILABLE = groq_client is not None or openai_client is not None
    
    if not groq_client and not openai_client:
        print("⚠️ No AI API key found, using fallback mode")
    elif groq_client and openai_client:
        print(f"✅ [Smart Router] Groq (优先) + OpenAI (备用)")
    elif groq_client:
        print(f"✅ Using Groq API (Free) | Model: {GROQ_MODEL}")
    else:
        print(f"✅ Using OpenAI API | Model: {OPENAI_MODEL}")
        
except Exception as e:
    print(f"⚠️ AI client init failed: {e}, using fallback mode")
    AI_AVAILABLE = False
    AI_PROVIDER = None
    AI_MODEL_RESUME = None
    AI_MODEL_COVER = None
    AI_MODEL = None
    groq_client = None
    openai_client = None
    GROQ_AVAILABLE = False
    
    def smart_ai_request(*args, **kwargs):
        class FakeError:
            choices = [type('obj', (object,), {'message': type('obj', (object,), {'content': 'AI unavailable'})()})()]
        return FakeError()
    
    def check_groq_recovery():
        return False
    
    def get_ai_status():
        return {"groq_available": False, "openai_available": False}

app = FastAPI(title="JobMatchAI Nordic API", version="2.0.0")

# CORS配置 - 明确允许的域名（allow_credentials=True时不能用*）
ALLOWED_ORIGINS = [
    "https://job-match-ai.com",
    "https://jobmatchai.net",
    "https://jobmatchai-4y1.pages.dev",
    "http://localhost:8080",
    "http://localhost:8000",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 学习推荐引擎路由
app.include_router(learning_router)

# === 数据模型 ===
class Job(BaseModel):
    title: str
    company: str
    location: str
    description: str
    url: Optional[str] = ""
    source: str  # linkedin, jobindex, jobnet, etc.
    language: str = "da"  # zh, en, da

class ResumeAnalysis(BaseModel):
    skills: List[str]
    experience_years: int
    strengths: List[str]
    improvements: List[Dict]
    suggested_profile: str

# === 简历解析 ===
# 导入统一 PDF 处理引擎
try:
    from pdf_engine import extract_text_from_pdf, DanishPropertyReportParser, PDFGenerator
    PDF_ENGINE_AVAILABLE = True
    print("✅ PDF Engine loaded (PyMuPDF + pdfplumber + OCR)")
except ImportError as e:
    PDF_ENGINE_AVAILABLE = False
    print(f"⚠️ PDF Engine not available: {e}")

def parse_resume(file_content: bytes, filename: str) -> str:
    """解析简历文件 - PDF / DOCX / TXT
    
    使用统一 PDF 引擎，支持：
    - 文字型 PDF（PyMuPDF 快速提取）
    - 扫描件 PDF（OCR 识别）
    - DOCX、TXT
    """
    text = ""
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith('.pdf'):
            # 优先使用统一 PDF 引擎
            if PDF_ENGINE_AVAILABLE:
                result = extract_text_from_pdf(file_content, use_ocr_fallback=True)
                text = result.get("text", "")
                method = result.get("method", "unknown")
                is_scanned = result.get("is_scanned", False)
                if text:
                    print(f"✅ PDF parsed via {method}, scanned={is_scanned}, chars={len(text)}")
                else:
                    print(f"⚠️ PDF engine returned empty text, method={method}")
            
            # 备用：pdfplumber
            if not text:
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                        for page in pdf.pages:
                            t = page.extract_text()
                            if t:
                                text += t + "\n"
                    if text:
                        print(f"✅ PDF parsed via pdfplumber fallback, chars={len(text)}")
                except Exception as e1:
                    print(f"pdfplumber fallback failed: {e1}")
            
            # 最终备用：PyPDF2
            if not text:
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(io.BytesIO(file_content))
                    for page in reader.pages:
                        t = page.extract_text()
                        if t:
                            text += t + "\n"
                    if text:
                        print(f"✅ PDF parsed via PyPDF2 fallback, chars={len(text)}")
                except Exception as e2:
                    print(f"PyPDF2 also failed: {e2}")
                            
        elif filename_lower.endswith('.docx'):
            try:
                import docx
                document = docx.Document(io.BytesIO(file_content))
                for para in document.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
                for table in document.tables:
                    for row in table.rows:
                        row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_texts:
                            text += " | ".join(row_texts) + "\n"
            except Exception as e:
                print(f"DOCX解析失败: {e}")
        
        elif filename_lower.endswith('.txt'):
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'gb18030']:
                try:
                    text = file_content.decode(encoding)
                    break
                except:
                    continue
                
    except Exception as e:
        print(f"解析错误: {e}")
    
    lines = [l for l in text.splitlines() if l.strip()]
    return "\n".join(lines)

# === 语言检测 ===
def detect_language(text: str) -> str:
    """检测文本语言：zh, en, da"""
    text_lower = text.lower()
    
    # 检测中文字符（最优先）
    zh_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    if zh_chars > 3:
        return 'zh'
    
    # 检测丹麦语特有字符（æ, ø, å）- 最可靠的丹麦语特征
    da_chars = sum(1 for c in ['æ', 'ø', 'å'] if c in text_lower)
    if da_chars >= 1:
        return 'da'
    
    # 检测丹麦语特有词汇（短词如 'en', 'er' 不可靠，容易误判）
    # 用更长的丹麦语词汇来检测
    da_words = ['erfaring', 'kompetencer', 'erfaringer', 'ansvarlig', 'arbejdsopgaver',
                'stilling', 'beskrivelse', 'kvalifikationer', 'uddannelse', 'erhvervserfaring']
    da_word_count = sum(1 for w in da_words if w in text_lower)
    if da_word_count >= 1:
        return 'da'
    
    # 默认英文（英语简历最常见）
    return 'en'

# === 简历结构化解析（AI驱动） ===
def parse_resume_structure(text: str, lang: str = 'en') -> Dict:
    """使用 AI 将简历解析为结构化数据
    
    Args:
        text: 简历文本（原始文本，来自 PDF/DOCX 提取）
        lang: 语言代码 (zh/en/da)
    
    Returns:
        结构化简历数据，包含：
        - personal: 基本信息（姓名、邮箱、电话）
        - education: 教育经历列表
        - experience: 工作经历列表（带年份计算）
        - skills: 技能分类（technical/soft/language）
        - total_experience_years: 总工作年限
    """
    if not AI_AVAILABLE:
        return parse_resume_structure_fallback(text, lang)
    
    prompts = {
        'zh': f"""你是一位专业的简历解析专家。请将以下简历文本解析为结构化JSON数据。

【严格规则】
1. 只从简历文本中提取信息，禁止编造任何内容
2. 教育经历要识别学位类型（博士/硕士/学士/专科）和学校名称
3. 工作经历要识别公司、职位、时间范围，并计算精确年限
4. 技能要分类为 technical（技术技能）、soft（软技能）、language（语言能力）
5. 如果无法从文本中确定某个字段，返回空字符串或空数组
6. 工作经验从最早的工作开始算起
7. 语言检测：根据简历的主要内容语言判断，返回 zh（中文）、en（英文）、da（丹麦文）之一

【输出格式 - 必须严格返回有效JSON】
{{
  "personal": {{
    "name": "姓名（从简历中提取）",
    "email": "邮箱",
    "phone": "电话"
  }},
  "education": [
    {{
      "degree": "学位类型（博士/硕士/学士/专科/其他）",
      "field": "专业",
      "school": "学校名称",
      "year": 年份数字,
      "country": "国家（如果能从简历中识别）"
    }}
  ],
  "experience": [
    {{
      "company": "公司名称",
      "role": "职位/头衔",
      "start": "开始时间（YYYY-MM格式，没有就留空）",
      "end": "结束时间（YYYY-MM或ongoing）",
      "years": 工作年限（数字，保留一位小数）,
      "duties": ["职责1", "职责2"]
    }}
  ],
  "skills": {{
    "technical": ["技术技能1", "技能2"],
    "soft": ["软技能1", "技能2"],
    "language": ["语言能力1（含等级）", "语言能力2"]
  }},
  "total_experience_years": 总工作年限（数字，保留一位小数）,
  "detected_language": "语言代码（zh/en/da，根据简历主要语言判断）"
}}

简历文本：
{text[:5000]}""",
        'en': f"""You are a professional resume parser. Parse the following resume text into structured JSON data.

【Strict Rules】
1. Extract ONLY from the resume text - do NOT fabricate any information
2. Identify degree types: PhD, Master's, Bachelor's, Associate, Diploma, Other
3. Work experience: identify company, role, date range, calculate precise years
4. Skills must be categorized as: technical (tools/languages), soft (interpersonal), language (with proficiency level)
5. If a field cannot be determined from the text, return empty string or empty array
6. Experience should start from the earliest job

【Output Format - Return valid JSON only】
{{
  "personal": {{
    "name": "Full name from resume",
    "email": "Email address",
    "phone": "Phone number"
  }},
  "education": [
    {{
      "degree": "PhD/Master's/Bachelor's/Associate/Diploma/Other",
      "field": "Field of study",
      "school": "University name",
      "year": year_number,
      "country": "Country (if identifiable from resume)"
    }}
  ],
  "experience": [
    {{
      "company": "Company name",
      "role": "Job title",
      "start": "Start date (YYYY-MM, empty if not found)",
      "end": "End date (YYYY-MM or 'ongoing')",
      "years": years_worked (number, one decimal),
      "duties": ["Duty 1", "Duty 2"]
    }}
  ],
  "skills": {{
    "technical": ["Technical skill 1", "Skill 2"],
    "soft": ["Soft skill 1", "Skill 2"],
    "language": ["Language (proficiency)", "Language 2"]
  }},
  "total_experience_years": total_years (number, one decimal)
}}

Resume text:
{text[:5000]}""",
        'da': f"""Du er en professionel CV-parser. Parse den følgende CV-tekst til struktureret JSON-data.

【Strenge regler】
1. Udtræk KUN fra CV-teksten - fabricate IKKE nogen information
2. Identificer gradstyper: PhD, Kandidat, Bachelor, Professionsbachelor, Andet
3. Arbejdserfaring: identificer virksomhed, stilling, datointerval, beregn præcise år
4. Kompetencer skal kategoriseres som: technical (værktøjer/sprog), soft (interpersonlig), language (med niveau)
5. Hvis et felt ikke kan bestemmes fra teksten, returner tom streng eller tom array
6. Erfaring skal starte fra det tidligste job
7. Sprogregistrering: Baseret på CV'ets hovedsproglige indhold, returner zh (kinesisk), en (engelsk) eller da (dansk)

【Output Format - Returner kun gyldig JSON】
{{
  "personal": {{
    "name": "Fulde navn fra CV",
    "email": "E-mailadresse",
    "phone": "Telefonnummer"
  }},
  "education": [
    {{
      "degree": "PhD/Kandidat/Bachelor/Professionsbachelor/Andet",
      "field": "Studieretning",
      "school": "Universitetsnavn",
      "year": årstal,
      "country": "Land (hvis identificerbart)"
    }}
  ],
  "experience": [
    {{
      "company": "Virksomhedsnavn",
      "role": "Stilling",
      "start": "Startdato (YYYY-MM, tom hvis ikke fundet)",
      "end": "Slutdato (YYYY-MM eller 'i gang')",
      "years": arbejdsår (tal, én decimal),
      "duties": ["Pligt 1", "Pligt 2"]
    }}
  ],
  "skills": {{
    "technical": ["Teknisk kompetence 1", "Kompetence 2"],
    "soft": ["Blød kompetence 1", "Kompetence 2"],
    "language": ["Sprog (niveau)", "Sprog 2"]
  }},
  "total_experience_years": samlede_år (tal, én decimal),
  "detected_language": "Sprogkode (zh/en/da, baseret på CV'ets hovedsprog)"
}}

CV tekst:
{text[:5000]}"""
    }
    
    try:
        prompt = prompts.get(lang, prompts['en'])
        
        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a professional resume parser. Return ONLY valid JSON, no markdown, no explanation."},
                {"role": "user", "content": prompt}
            ],
            preferred_provider="groq",
            temperature=0.1,  # 低温度确保一致性
            max_tokens=2500
        )
        
        # 清理响应中的 markdown 代码块
        raw_content = response.choices[0].message.content.strip()
        # 移除 ```json 和 ``` 标记
        if raw_content.startswith('```json'):
            raw_content = raw_content[7:]
        elif raw_content.startswith('```'):
            raw_content = raw_content[3:]
        if raw_content.endswith('```'):
            raw_content = raw_content[:-3]
        raw_content = raw_content.strip()
        
        result = json.loads(raw_content)
        
        # 处理 AI 返回数组或嵌套在 resumes 键中的情况
        if isinstance(result, list):
            result = result[0] if result else {}
        elif 'resumes' in result:
            result = result['resumes'][0] if result['resumes'] else {}
        elif 'resume' in result:
            result = result['resume'] if isinstance(result['resume'], dict) else {}
        
        # 确保有必要的字段
        if 'personal' not in result:
            result['personal'] = {}
        if 'education' not in result:
            result['education'] = []
        if 'experience' not in result:
            result['experience'] = []
        if 'skills' not in result:
            result['skills'] = {'technical': [], 'soft': [], 'language': []}

        # 优先使用 AI 检测的语言（更准确），否则用硬编码结果
        if 'detected_language' in result and result['detected_language'] in ['zh', 'en', 'da']:
            result['detected_language'] = result['detected_language']
        else:
            result['detected_language'] = lang
        
        result['ai_enhanced'] = True
        
        return result
        
    except json.JSONDecodeError as e:
        print(f"JSON parse failed: {e}")
        return parse_resume_structure_fallback(text, lang)
    except Exception as e:
        print(f"AI parse failed: {e}, using fallback")
        return parse_resume_structure_fallback(text, lang)


def parse_resume_structure_fallback(text: str, lang: str = 'en') -> Dict:
    """无 AI 时的简历结构化解析后备方案"""
    import re
    
    text_lower = text.lower()
    
    # 简单提取邮箱和电话
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    phone_pattern = r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\./0-9]{7,}'
    
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    
    # 基本结构（无 AI 时只能返回有限信息）
    return {
        "personal": {
            "name": "",
            "email": emails[0] if emails else "",
            "phone": phones[0] if phones else ""
        },
        "education": [],
        "experience": [],
        "skills": {
            "technical": [],
            "soft": [],
            "language": []
        },
        "total_experience_years": 0,
        "ai_enhanced": False,
        "note": "AI unavailable - limited parsing only. Please upload resume with AI connection for full analysis."
    }


# === AI 增强简历分析 ===
def analyze_resume_with_ai(text: str, lang: str = 'en', job_context: Dict = None) -> Dict:
    """使用 AI 深度分析简历
    
    Args:
        text: 简历文本
        lang: 语言代码 (zh/en/da)
        job_context: 职位上下文，包含 job_title, company, job_description
    """
    if not AI_AVAILABLE:
        return analyze_resume_fallback(text, lang, job_context)
    
    try:
        # 构建基础 prompt
        base_prompt = f"""
请分析以下简历，以JSON格式返回：
{{
  "skills": ["技能1", "技能2"],
  "experience_years": 数字,
  "strengths": ["优势1", "优势2", "优势3"],
  "improvements": [
    {{"type": "weak_verb", "priority": "high/medium/low", "description": "问题描述", "suggestion": "改进建议"}}
  ],
  "suggested_profile": "建议的个人简介（50字以内）",
  "ats_score": 数字（1-100 ATS友好度评分）
}}

简历内容：
{text[:3000]}"""
        
        # 如果有职位上下文，添加职位导向的分析
        if job_context and job_context.get('job_title'):
            job_title = job_context.get('job_title', '')
            job_desc = job_context.get('job_description', '')[:500] if job_context.get('job_description') else ''
            company = job_context.get('company', '')
            
            base_prompt = f"""你是一位专业的职业顾问。请对比分析以下简历与目标职位，然后以JSON格式返回分析结果：

目标职位信息：
- 职位名称：{job_title}
- 公司：{company if company else '未指定'}
- 职位描述：{job_desc if job_desc else '无详细描述'}

简历内容：
{text[:3000]}

请返回JSON格式的分析结果：
{{
  "skills": ["检测到的技能1", "技能2"],
  "experience_years": 数字（工作经验年限）,
  "strengths": ["优势1", "优势2", "优势3"],
  "relevant_experience": ["与职位最相关的经验1", "相关经验2"],
  "job_match_score": 数字（1-100，简历与职位的匹配度）,
  "missing_skills": ["职位要求但简历中缺少的技能1", "缺少的技能2"],
  "improvements": [
    {{"type": "job_relevance", "priority": "high/medium/low", "description": "描述", "suggestion": "针对该职位的改进建议"}}
  ],
  "suggested_profile": "针对该职位优化的个人简介（50字以内，突出最相关的经验和技能）",
  "ats_score": 数字（1-100 ATS友好度评分）,
  "job_specific_tips": ["针对该职位的申请技巧1", "技巧2"]
}}

重点：
1. 突出简历中与目标职位最相关的经验和技能
2. 指出简历中缺失的职位要求技能
3. 提供针对该职位的具体改进建议
4. 评估简历与职位的整体匹配度"""
        
        # 根据语言选择 prompt
        if lang == 'zh':
            prompt = base_prompt
        elif lang == 'da':
            prompt = f"""Du er en professionel karriererådgiver. Analyser følgende CV i forhold til stillingen og returner JSON:

Stilling:
- Titel: {job_context.get('job_title', 'Ikke angivet') if job_context else 'Ikke angivet'}
- Virksomhed: {job_context.get('company', '') if job_context else 'Ikke angivet'}
- Beskrivelse: {(job_context.get('job_description') or '')[:500] or 'Ingen detaljeret beskrivelse'}

CV:
{text[:3000]}

Returner JSON:
{{
  "skills": ["kompetence1", "kompetence2"],
  "experience_years": tal,
  "strengths": ["styrke1", "styrke2", "styrke3"],
  "relevant_experience": ["relevant erfaring 1", "erfaring 2"],
  "job_match_score": tal (1-100),
  "missing_skills": ["manglende kompetence 1", "kompetence 2"],
  "improvements": [
    {{"type": "job_relevance", "priority": "high/medium/low", "description": "problem", "suggestion": "forbedring"}}
  ],
  "suggested_profile": "profil tilpasset stillingen (under 50 ord)",
  "ats_score": tal (1-100),
  "job_specific_tips": ["tips 1", "tips 2"]
}}"""
        else:  # English
            if job_context and job_context.get('job_title'):
                prompt = f"""You are a professional career consultant. Analyze this resume against the target job and return JSON:

Target Job:
- Title: {job_context.get('job_title', 'Not specified')}
- Company: {job_context.get('company', 'Not specified')}
- Description: {(job_context.get('job_description') or '')[:500] or 'No detailed description'}

Resume:
{text[:3000]}

Return JSON:
{{
  "skills": ["skill1", "skill2"],
  "experience_years": number,
  "strengths": ["strength1", "strength2", "strength3"],
  "relevant_experience": ["relevant experience 1", "experience 2"],
  "job_match_score": number (1-100),
  "missing_skills": ["missing skill 1", "skill 2"],
  "improvements": [
    {{"type": "job_relevance", "priority": "high/medium/low", "description": "issue", "suggestion": "improvement"}}
  ],
  "suggested_profile": "job-specific profile (under 50 words, highlighting most relevant experience and skills)",
  "ats_score": number (1-100),
  "job_specific_tips": ["application tip 1", "tip 2"]
}}

Focus on:
1. Highlighting experiences and skills most relevant to the target job
2. Identifying missing skills required by the job
3. Providing specific improvement suggestions for this job
4. Assessing overall resume-job fit"""
            else:
                prompt = f"""Please analyze this resume and return JSON:
{{
  "skills": ["skill1", "skill2"],
  "experience_years": number,
  "strengths": ["strength1", "strength2", "strength3"],
  "improvements": [
    {{"type": "weak_verb", "priority": "high/medium/low", "description": "issue", "suggestion": "improvement"}}
  ],
  "suggested_profile": "suggested profile summary (under 50 words)",
  "ats_score": number (1-100 ATS compatibility score)
}}

Resume:
{text[:3000]}"""
        
        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a professional resume analyzer. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            preferred_provider="groq",
            temperature=0.3,
            max_tokens=2000
        )
        
        result = json.loads(response.choices[0].message.content)
        result['detected_language'] = lang
        result['ai_enhanced'] = True
        result['ai_provider'] = AI_PROVIDER
        
        # 如果有职位上下文，添加相关信息
        if job_context and job_context.get('job_title'):
            result['job_context'] = {
                'job_title': job_context.get('job_title', ''),
                'company': job_context.get('company', ''),
                'job_match_score': result.get('job_match_score', 0)
            }
            
            # AI增强：自动识别关键要求和所需技能
            if job_context.get('job_description'):
                result['critical_requirements'] = identify_critical_requirements(
                    job_context['job_description'], text, lang
                )
                result['required_skills'] = extract_required_skills_from_job(
                    job_context['job_description'], lang
                )
        
        return result
        
    except Exception as e:
        print(f"AI analysis failed: {e}, using fallback")
        return analyze_resume_fallback(text, lang, job_context)

def analyze_resume_fallback(text: str, lang: str = 'en', job_context: Dict = None) -> Dict:
    """本地规则分析（无需AI）
    
    Args:
        text: 简历文本
        lang: 语言代码
        job_context: 职位上下文
    """
    # 技能关键词库（中英丹）
    skills_db = {
        'zh': ['Python', 'Java', 'ERP', 'NetSuite', 'Dynamics', '项目管理', '数据分析', 'SQL', '财务', '供应链', '实施', '顾问', '财务', '供应链', 'AX', 'SAP', 'Oracle'],
        'en': ['Python', 'Java', 'ERP', 'NetSuite', 'Dynamics', 'Project Management', 'Data Analysis', 'SQL', 'Finance', 'Supply Chain', 'Implementation', 'Consultant', 'AX', 'SAP', 'Oracle'],
        'da': ['Python', 'Java', 'ERP', 'NetSuite', 'Dynamics', 'Projektledelse', 'Dataanalyse', 'SQL', 'Finans', 'Forsyningskæde', 'Implementering', 'Konsulent', 'AX', 'SAP', 'Oracle']
    }
    
    weak_verbs = {
        'zh': ['做了', '负责了', '参与了', '协助', '帮助'],
        'en': ['did', 'was responsible for', 'worked on', 'helped with', 'assisted'],
        'da': ['lavede', 'var ansvarlig for', 'arbejdede på', 'hjalp med', 'assisterede']
    }
    
    strong_verbs = {
        'zh': ['主导', '推动', '优化', '实现', '提升', '建立', '设计'],
        'en': ['led', 'implemented', 'optimized', 'achieved', 'improved', 'established', 'designed'],
        'da': ['ledede', 'implementerede', 'optimerede', 'opnåede', 'forbedrede', 'etablerede', 'designede']
    }
    
    text_lower = text.lower()
    detected_skills = [s for s in skills_db.get(lang, skills_db['en']) if s.lower() in text_lower]
    detected_weak = [v for v in weak_verbs.get(lang, weak_verbs['en']) if v in text_lower]
    
    improvements = []
    if detected_weak:
        improvements.append({
            'type': 'weak_verb',
            'priority': 'high',
            'description': {'zh': f'检测到 {len(detected_weak)} 处可强化的动词', 'en': f'Found {len(detected_weak)} weak verbs', 'da': f'Fandt {len(detected_weak)} svage verber'}.get(lang, 'Found weak verbs'),
            'suggestion': {'zh': '使用更强动词如"主导"、"实现"', 'en': 'Use stronger verbs like "led", "achieved"', 'da': 'Brug stærkere verber'}.get(lang, 'Use stronger verbs')
        })
    
    # 计算与职位的匹配度（如果有职位上下文）
    job_match_score = 50  # 默认值
    relevant_experience = []
    missing_skills = []
    job_specific_tips = []
    learning_recommendations = []  # 学习建议
    extracted_job_skills = []  # 职位要求技能
    matched_skills = []  # 匹配的技能
    
    if job_context and job_context.get('job_title'):
        job_title = job_context.get('job_title', '')
        job_desc = job_context.get('job_description', '')
        
        # ===== AI驱动的Skill Gap分析（完全由AI完成，不使用硬编码关键词）=====
        print(f"[AI] Analyzing skill gap for: {job_title}")
        
        skill_gap_result = analyze_skill_gap_ai(
            resume_text=text,
            job_description=job_desc,
            job_title=job_title,
            detected_skills=detected_skills,
            lang=lang
        )
        
        # 使用AI分析结果
        missing_skills = skill_gap_result.get('missing_skills', [])
        matched_skills = skill_gap_result.get('matched_skills', [])
        job_match_score = skill_gap_result.get('score', 50)
        critical_requirements = skill_gap_result.get('critical_gaps', [])
        ai_reasoning = skill_gap_result.get('reasoning', '')
        
        print(f"[AI] Skill gap result: missing={missing_skills}, matched={matched_skills}, score={job_match_score}")
        
        # 限制缺失技能数量
        if len(missing_skills) > 5:
            missing_skills = missing_skills[:5]
        
        # ===== AI驱动的学习推荐 ======
        if AI_AVAILABLE and missing_skills:
            try:
                learning_prompt = {
                    'en': f"""Based on the following skill gaps, generate learning recommendations in JSON format:

Missing Skills: {missing_skills}
Job Title: {job_title}

For each missing skill, generate a recommendation:
[
  {{
    "skill": "skill name",
    "suggestion": "specific learning recommendation (e.g., course name, platform)",
    "level": "beginner/intermediate/advanced",
    "resource_type": "video/course/certification",
    "estimated_time": "e.g., 10 hours"
  }}
]

Rules:
1. Only for skills EXPLICITLY mentioned in the job description
2. Be specific about courses and platforms
3. Return ONLY JSON array""",
                    
                    'zh': f"""基于以下技能差距，生成学习推荐（JSON格式）：

缺失技能：{missing_skills}
职位名称：{job_title}

对每个缺失技能，生成推荐：
[
  {{
    "skill": "技能名",
    "suggestion": "具体学习建议（如课程名、学习平台）",
    "level": "beginner/intermediate/advanced",
    "resource_type": "video/course/certification",
    "estimated_time": "如：10小时"
  }}
]

规则：
1. 只针对职位描述中明确提到的技能
2. 具体说明课程和平台
3. 只返回JSON数组"""
                }.get(lang, '')
                
                resp = smart_ai_request(
                    messages=[
                        {"role": "system", "content": "You are a learning advisor. Return ONLY valid JSON array."},
                        {"role": "user", "content": learning_prompt}
                    ],
                    preferred_provider="groq",
                    temperature=0.3,
                    max_tokens=800
                )
                
                content = resp.choices[0].message.content.strip()
                if content.startswith('```'):
                    content = content.split('```')[1]
                    if content.startswith('json'):
                        content = content[4:]
                
                learning_recommendations = json.loads(content)
                print(f"[AI] Generated {len(learning_recommendations)} learning recommendations")
                
            except Exception as e:
                print(f"Learning recommendation failed: {e}")
                learning_recommendations = []
        
        # 软技能建议（可以添加到简历）
        soft_skill_suggestions = []
        
        # 添加Gap分析改进建议
        improvements.append({
            'type': 'skill_gap',
            'priority': 'high' if missing_skills else 'medium',
            'description': {
                'zh': f'🎯 Skill Gap分析：发现 {len(missing_skills)} 项可提升技能',
                'en': f'🎯 Skill Gap: Found {len(missing_skills)} skills to develop',
                'da': f'🎯 Skill Gap: Fundet {len(missing_skills)} kompetencer at udvikle'
            }.get(lang, f'Skill Gap: {len(missing_skills)} skills'),
            'suggestion': ai_reasoning or {
                'zh': '针对性学习，提升岗位竞争力',
                'en': 'Targeted learning to boost competitiveness',
                'da': 'Målrettet læring for konkurrenceevne'
            }.get(lang, 'Targeted learning')
        })
        
        improvements.append({
            'type': 'job_relevance',
            'priority': 'high' if job_match_score < 60 else 'medium',
            'description': {
                'zh': f'📊 简历-职位匹配度：{job_match_score}%',
                'en': f'📊 Resume-Job Match: {job_match_score}%',
                'da': f'📊 CV-Job Match: {job_match_score}%'
            }.get(lang, f'Match: {job_match_score}%'),
            'suggestion': {
                'zh': '根据职位要求调整简历重点',
                'en': 'Tailor resume to job requirements',
                'da': 'Tilpas CV til jobkrav'
            }.get(lang, 'Tailor to job')
        })
    
    years = re.findall(r'(\d+)\s*(?:年|years?|år)', text_lower)
    exp_years = max([int(y) for y in years] + [0])
    
    # 根据是否有职位上下文生成不同的简介
    if job_context and job_context.get('job_title'):
        profile_templates = {
            'zh': f'拥有{exp_years}年ERP系统实施经验，精通{", ".join(detected_skills[:2]) if detected_skills else "NetSuite和Dynamics"}。',
            'en': f'ERP professional with {exp_years}+ years. Expertise in {", ".join(detected_skills[:2]) if detected_skills else "NetSuite and Dynamics"}.',
            'da': f'ERP-professionel med {exp_years}+ års erfaring. Ekspertise i {", ".join(detected_skills[:2]) if detected_skills else "NetSuite og Dynamics"}.'
        }
    else:
        profile_templates = {
            'zh': f'拥有{exp_years}年ERP系统实施经验，精通NetSuite和Dynamics AX。',
            'en': f'Experienced ERP professional with {exp_years}+ years in NetSuite and Dynamics AX.',
            'da': f'Erfaren ERP-professionel med {exp_years}+ års erfaring i NetSuite og Dynamics AX.'
        }
    
    result = {
        'skills': detected_skills,
        'experience_years': exp_years,
        'strengths': detected_skills[:3] if detected_skills else ['ERP', 'Project Management'],
        'improvements': improvements,
        'suggested_profile': profile_templates.get(lang, profile_templates['en']),
        'ats_score': 70,
        'detected_language': lang,
        'ai_enhanced': False
    }
    
    # 如果有职位上下文，添加相关信息
    if job_context and job_context.get('job_title'):
        result['job_context'] = {
            'job_title': job_context.get('job_title', ''),
            'company': job_context.get('company', ''),
            'job_match_score': job_match_score
        }
        result['job_match_score'] = job_match_score
        result['relevant_experience'] = relevant_experience
        result['missing_skills'] = missing_skills[:5]
        result['job_specific_tips'] = job_specific_tips
        result['learning_recommendations'] = learning_recommendations  # 学习建议
        result['resume_skills'] = detected_skills  # 简历中的技能
        result['required_skills'] = extracted_job_skills  # 职位要求技能
        result['critical_requirements'] = critical_requirements  # 关键要求（日语等）
        result['soft_skill_suggestions'] = soft_skill_suggestions  # 软技能建议
        result['required_skills'] = extracted_job_skills  # 职位要求的技能
        result['matched_skills'] = matched_skills  # 匹配的技能
    
    return result

# 保持向后兼容
analyze_resume = analyze_resume_with_ai

# === AI 增强求职信生成 ===
def generate_cover_letter_with_ai(resume_text: str, job: Dict, lang: str = 'en', resume_highlights: List[Dict] = None, resume_structure: Dict = None) -> str:
    """使用 AI 生成高质量求职信（升级版：使用简历亮点 + 结构化数据防瞎编）
    
    支持 LinkedIn 导入职位的三语生成（中/英/丹）
    resume_highlights: 从精修建议中提取的简历亮点
    resume_structure: 结构化简历数据（用于防止瞎编）
    """
    if not AI_AVAILABLE:
        return generate_cover_letter_fallback(resume_text, job, lang)
    
    try:
        # 职位来源（如果是 LinkedIn，优化 prompt）
        is_linkedin = 'linkedin' in str(job.get('source', '')).lower()
        source_hint = ""
        if is_linkedin:
            source_hint = "\n注意：此职位来自 LinkedIn，求职信应该简练有力，突出核心竞争力，适合快速浏览。"
        
        # 简历亮点（如果有）
        highlights_text = ""
        if resume_highlights and len(resume_highlights) > 0:
            highlights_list = "\n".join([
                f"- {h.get('title', '')}: {h.get('description', '')}" 
                for h in resume_highlights[:3]
            ])
            highlights_text = f"\n\n【简历亮点 - 请在求职信中重点使用】\n{highlights_list}\n"
        
        # 结构化简历数据（用于防止瞎编）
        structure_context = ""
        if resume_structure and resume_structure.get('ai_enhanced'):
            edu_list = resume_structure.get('education', [])
            exp_list = resume_structure.get('experience', [])
            skills_data = resume_structure.get('skills', {})
            
            edu_str = ", ".join([f"{e.get('degree', '')} in {e.get('field', '')} from {e.get('school', '')} ({e.get('year', '')})" for e in edu_list]) if edu_list else "无学历信息"
            exp_str = " | ".join([f"{e.get('company', '')} ({e.get('start', '')}-{e.get('end', '')}): {e.get('role', '')}" for e in exp_list[:3]]) if exp_list else "无工作经历"
            tech_skills = ", ".join(skills_data.get('technical', [])[:10]) if skills_data.get('technical') else "未识别"
            
            structure_context = f"""

【已验证的简历事实 - 求职信中只能使用这些信息，禁止编造！】
学历: {edu_str}
工作经历: {exp_str}
技术技能: {tech_skills}
总工作经验: {resume_structure.get('total_experience_years', 'N/A')}年

⚠️ 严格规则：求职信中的所有案例、数字、公司名称必须来自上述已验证的事实！
如果上述信息中没有某个公司的名字，不要在求职信中提及该公司。
            """
        
        prompts = {
            'zh': f"""写一封专业的求职信（300-400字），打动{job.get('company', '这家公司')}的HR。

{highlights_text}{structure_context}【简历内容 - 请从这些真实信息中选择素材】
{resume_text[:1000]}

【职位信息】
- 公司：{job.get('company', '')}
- 职位：{job.get('title', '')}
- 地点：{job.get('location', '')}
- 要求：{job.get('description', '')[:800]}
{source_hint}

【防瞎编规则 - 绝对遵守！否则求职信会失去可信度】
⚠️ 最最重要规则：
1. 只使用【简历内容】和【已验证的简历事实】中的信息
2. 只能提及简历中明确写出的公司名称
3. 只能使用简历中写出的数字和百分比
4. 只能描述简历中提到的项目和技术

❌ 绝对禁止（这些会让HR觉得你是在撒谎）：
- 编造公司名：如"一家大型零售商"（除非简历中真的这么写）
- 编造数字：如"提高效率40%"、"节省成本15万"（除非简历中真的有）
- 编造项目：如"主导了XX项目"（除非简历中真的提到了）
- 编造职位：如"曾任技术总监"（除非简历中真的这么写）

✅ 正确做法：
- 直接引用简历中的真实经历
- 用"在我担任[简历中的真实职位]期间，我[简历中提到的真实行动]"格式
- 如果简历中没有具体数字，就不要提数字，只描述做了什么

1. 开头（前2句）：
   ✅ 好："在[简历中的真实公司名]担任[简历中的真实职位]期间，我..."
   ❌ 差："在一家知名企业，我让效率提升50%"（除非简历里真的这么写）
2. 中间（2-3段）：
   - 只描述简历中提到的真实经历
   - 如果简历说"负责数据分析"，就说"负责数据分析"
   - 如果简历没写数字，就不要说"提升了30%"
3. 结尾：
   ✅ 好："我对[公司名]的[职位相关业务]很感兴趣，希望能进一步讨论"
   ❌ 差："相信我的经验能为贵公司创造巨大价值"（太泛）

【北欧职场风格】
- 直接、真诚、不过度自夸
- 强调实际贡献而非夸大
- 数据要真实，没有就不写

直接返回求职信正文，不要JSON，不要格式符号。""",
            'en': f"""Write a compelling cover letter (300-400 words) that will impress {job.get('company', 'this company')}'s HR.

{highlights_text}【Resume Content - Use ONLY these REAL experiences】
{resume_text[:1000]}

{structure_context}

【Job Details】
- Company: {job.get('company', '')}
- Position: {job.get('title', '')}
- Location: {job.get('location', '')}
- Requirements: {job.get('description', '')[:800]}
{source_hint}

【ANTI-FABRICATION RULES - MANDATORY!】
⚠️ MOST CRITICAL RULE:
1. Use ONLY information from the resume content and verified facts above
2. Only mention company names that are EXPLICITLY in the resume
3. Only use numbers/percentages that are EXPLICITLY in the resume
4. Only describe projects/technologies that are in the resume

❌ ABSOLUTELY FORBIDDEN (HR will think you're lying):
- NEVER invent ANY percentage: 30%, 40%, 50%, etc. - UNLESS the word "30%" (or similar) is literally in the resume text
- NEVER invent metrics: "increased by X", "reduced by Y", "saved Z amount" - UNLESS these numbers are in the resume
- NEVER use company names not in the resume
- NEVER describe projects not mentioned in the resume

✅ COPY-PASTE RULE:
- Your job is to REWORD, not INVENT
- If resume says "Improved reporting efficiency" - use those words or similar, but DO NOT add "by 30%"
- If resume says "worked on data projects" - say "worked on data projects", not "led data projects that increased efficiency by 25%"

⚠️ EXAMPLE - WRONG:
Resume: "Improved reporting efficiency"
Letter: "I improved reporting efficiency by 30%" ← ❌ WRONG! "30%" not in resume!

⚠️ EXAMPLE - CORRECT:
Resume: "Improved reporting efficiency"
Letter: "I improved reporting efficiency" ← ✅ CORRECT! No invented numbers

1. Opening (first 2 sentences):
   ✅ Good: "In my role as a Data Analyst at ABC Company, I improved reporting efficiency."
   ❌ Bad: "I improved reporting efficiency by 30%." (unless 30% is in resume)
2. Body (2-3 paragraphs):
   - Describe using the SAME LEVEL of detail as the resume
   - If resume is vague, your letter should be vague (but professional)
   - Never add specificity that isn't in the resume
3. Closing:
   ✅ Good: "I'm particularly interested in [company]'s work, and would welcome the opportunity to discuss how my experience aligns with the role."
   ❌ Bad: "I will bring immediate value to your team" (too boastful without evidence)

【Nordic Workplace Style】
- Direct, genuine, not over-selling
- Focus on actual contributions, not exaggerated claims
- Real data speaks - if no numbers exist, don't invent them

Return ONLY the cover letter text, no JSON, no formatting symbols.""",
            'da': f"""Skriv en professionel ansøgning (300-400 ord), der vil imponere {job.get('company', 'denne virksomhed')}s HR.

{highlights_text}【CV Indhold - Brug KUN disse VIRKELIGE erfaringer som materiale】
{resume_text[:1000]}

【Jobdetaljer】
- Virksomhed: {job.get('company', '')}
- Stilling: {job.get('title', '')}
- Lokation: {job.get('location', '')}
- Krav: {job.get('description', '')[:800]}
{source_hint}

【Skrivekrav - VIGTIGT! Følg strengt!】
⚠️ MEST VIGTIGE REGLER: Alle eksempler, tal og resultater SKAL komme fra CV-indholdet ovenfor. FABRIKER ALDRIG!
❌ FORBUDT: At opfinde virksomhedsnavne (som "en stor detailhandel"), at opfinde tal (som "15% stigning"), at opfinde projekter
✅ KORREKT: Vælg virkelige erfaringer fra CV'et og omstrukturer dem med "Udfordring → Handling → Resultat"

1. Indledning (første 2 sætninger): Fang opmærksomhed straks
   ✅ Godt: "I min [virkelig jobtitel fra CV], lykkedes det mig at [virkelig handling fra CV]..."
   ❌ Dårligt: "Hos Rolex reducerede jeg projektleveringstiden med 40%" (medmindre det virkelig står i CV'et)
2. Hoveddel (2-3 afsnit): Baseret på VIRKELIGE CV-erfaringer
   - Vælg virkelige projekter eller resultater fra CV'et
   - Brug "Udfordring → Handling → Resultat" struktur
   - Hvis CV'et ikke har specifikke tal, så tilføj ikke tal
3. Afslutning: Specifik call to action
   ✅ Godt: "Jeg vil meget gerne diskutere, hvordan jeg kan hjælpe {job.get('company', 'jeres virksomhed')} med at nå [specifikt mål fra jobbeskrivelsen]"
   ❌ Dårligt: "Jeg ser frem til at høre fra jer"

【Nordisk Arbejdspladsstil】
- Direkte, ægte, ikke over-sælge
- Fremhæv teamwork og kontinuerlig læring
- Lad virkelige data tale (skal være autentiske data fra CV)

Returner KUN ansøgningsteksten, ingen JSON."""
        }
        
        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a professional career coach who writes compelling, personalized cover letters. You understand Nordic/Danish work culture — direct communication, value-focused, not overly boastful. You specialize in writing cover letters that get interviews."},
                {"role": "user", "content": prompts.get(lang, prompts['en'])}
            ],
            preferred_provider="groq",
            temperature=0.7,
            max_tokens=1200
        )
        
        letter = response.choices[0].message.content.strip()
        return letter
        
    except Exception as e:
        print(f"AI cover letter failed: {e}, using fallback")
        return generate_cover_letter_fallback(resume_text, job, lang)

def generate_cover_letter_fallback(resume_text: str, job: Dict, lang: str = 'en') -> str:
    """模板生成（无需AI）— 针对北欧市场优化"""
    # 提取姓名（从简历第一行或默认）
    name_line = resume_text.strip().split('\n')[0] if resume_text else ""
    # 尝试提取英文姓名
    import re as _re
    name_match = _re.search(r'^([A-Z][a-z]+\s+[A-Z][a-z]+)', name_line)
    name = name_match.group(1) if name_match else "[Dit Navn]"

    templates = {
        'zh': f"""尊敬的招聘经理：

您好！我在{job.get('source', '招聘网站')}上看到贵公司{job.get('company', '')}的{job.get('title', '')}职位，对此非常感兴趣。

基于我的简历，我拥有丰富的相关经验，相信能为贵公司带来价值。

职位要求：
{job.get('description', '')[:500]}...

期待有机会与您进一步交流。

此致
敬礼！

{[name]}""",
        'en': f"""Dear Hiring Manager,

I am writing to express my strong interest in the {job.get('title', '')} position at {job.get('company', '')}.

Based on my experience outlined in my resume, I am confident that I can bring significant value to your team.

Job Requirements:
{job.get('description', '')[:500]}...

I look forward to discussing how my skills align with your needs.

Best regards,

{name}""",
        'da': f"""Kære Rekrutteringsansvarlig,

Jeg skriver for at udtrykke min interesse for stillingen som {job.get('title', '')} hos {job.get('company', '')}.

Baseret på min erfaring beskrevet i mit CV er jeg overbevist om, at jeg kan bidrage med reel værdi til jeres team fra første dag.

Jobkrav:
{job.get('description', '')[:500]}...

Jeg vil meget gerne til en samtale, hvor vi kan drøfte, hvordan mine kompetencer matcher jeres behov.

Med venlig hilsen,

{name}"""
    }
    return templates.get(lang, templates['en'])

# 保持向后兼容
generate_cover_letter = generate_cover_letter_with_ai

# === 职位搜索 API ===
import requests

# Adzuna API 配置 (支持 UK, AU, CA, FR, NL, BE)
ADZUNA_APP_ID = "690f8e34"
ADZUNA_APP_KEY = "9e5d7db533450288d6780344c1c160ba"
ADZUNA_COUNTRIES = {
    "gb": "英国",
    "au": "澳大利亚",
    "ca": "加拿大",
    "fr": "法国",
    "nl": "荷兰",
    "be": "比利时"
}

def search_adzuna(keyword: str, country: str = "gb", location: str = "", limit: int = 10) -> List[Dict]:
    """搜索国际职位 - Adzuna API"""
    url = f"https://api.adzuna.com/v1/api/jobs/{country}/search/1"
    params = {
        "app_id": ADZUNA_APP_ID,
        "app_key": ADZUNA_APP_KEY,
        "what": keyword,
        "results_per_page": min(limit, 50),
        "content-type": "application/json"
    }
    if location:
        params["where"] = location
    
    try:
        r = requests.get(url, params=params, timeout=15)
        if r.status_code != 200:
            return []
        
        data = r.json()
        jobs = []
        country_name = ADZUNA_COUNTRIES.get(country, country.upper())
        
        if "results" in data:
            for j in data["results"]:
                company_val = j.get("company")
                if isinstance(company_val, dict):
                    company_name = company_val.get("display_name", "N/A")
                else:
                    company_name = str(company_val) if company_val else "N/A"
                
                location_val = j.get("location")
                if isinstance(location_val, dict):
                    location_name = location_val.get("display_name", "")
                else:
                    location_name = str(location_val) if location_val else ""
                
                # 提取薪资范围
                salary_min = j.get("salary_min")
                salary_max = j.get("salary_max")
                salary_range = ""
                if salary_min or salary_max:
                    currency = "DKK" if country == "dk" else "GBP" if country == "gb" else "EUR"
                    if salary_min and salary_max:
                        # 转换为年薪
                        if salary_min < 10000:
                            salary_min *= 12
                            salary_max *= 12
                        salary_range = f"{currency} {salary_min/1000:.0f}K-{salary_max/1000:.0f}K"
                    elif salary_min:
                        salary_range = f"From {currency} {salary_min/1000:.0f}K"
                
                jobs.append({
                    "title": j.get("title", ""),
                    "company": company_name,
                    "location": location_name,
                    "description": j.get("description", "")[:500] if j.get("description") else "",
                    "url": j.get("redirect_url", ""),
                    "date": j.get("created", "")[:10] if j.get("created") else "",
                    "source": f"🇦🇹 Adzuna-{country_name}",
                    "salary_range": salary_range,
                    "language": "en"
                })
        return jobs
    except Exception as e:
        print(f"Adzuna API Error: {e}")
        return []

# 德国 Arbeitsagentur API v2
ARBEITSAGENTUR_KEY = "jobboerse-jobsuche"
ARBEITSAGENTUR_BASE = "https://rest.arbeitsagentur.de/jobboerse/jobsuche-service"

def search_germany(keyword: str, location: str = "", limit: int = 10) -> List[Dict]:
    """搜索德国职位 - Arbeitsagentur API v2"""
    url = f"{ARBEITSAGENTUR_BASE}/pc/v4/jobs"
    params = {
        "was": keyword,
        "page": 1,
        "size": min(limit, 50)
    }
    if location:
        params["wo"] = location
    else:
        params["wo"] = "Deutschland"
    
    headers = {
        "X-API-Key": ARBEITSAGENTUR_KEY,
        "User-Agent": "JobMatchAI/1.0"
    }
    
    try:
        r = requests.get(url, params=params, headers=headers, timeout=15)
        if r.status_code != 200:
            print(f"Arbeitsagentur API Error: {r.status_code} - {r.text[:200]}")
            return []
        
        data = r.json()
        jobs = []
        
        # 新版 API 返回格式: 直接是 stellenangebote 数组
        for j in data.get("stellenangebote", [])[:limit]:
            if not isinstance(j, dict):
                continue
            arbeitsort = j.get("arbeitsort", {})
            if isinstance(arbeitsort, dict):
                location = arbeitsort.get("ort", "")
            else:
                location = str(arbeitsort)
            
            jobs.append({
                "title": j.get("titel", ""),
                "company": j.get("arbeitgeber", "N/A"),
                "location": location,
                "description": j.get("beruf", ""),
                "url": j.get("externeUrl", f"https://jobboerse.arbeitsagentur.de/prod/jobboard/pc/v4/jobdetails/{j.get('refnr', '')}"),
                "date": j.get("aktuelleVeroeffentlichungsdatum", "")[:10] if j.get("aktuelleVeroeffentlichungsdatum") else "",
                "source": "🇩🇪 Arbeitsagentur",
                "language": "de"
            })
        return jobs
    except Exception as e:
        print(f"Arbeitsagentur API Error: {e}")
        return []

# 瑞典 Jobtechdev API
JOBTECHDEV_BASE = "https://jobtechdev.se/api/v1"

def search_sweden(keyword: str, location: str = "", limit: int = 10) -> List[Dict]:
    """搜索瑞典职位 - Jobtechdev API"""
    url = f"{JOBTECHDEV_BASE}/search"
    params = {
        "q": keyword,
        "limit": min(limit, 50)
    }
    if location:
        params["place"] = location
    
    try:
        r = requests.get(url, params=params, timeout=15)
        if r.status_code != 200:
            return []
        
        data = r.json()
        jobs = []
        
        for j in data.get("hits", {}).get("hits", [])[:limit]:
            src = j.get("_source", {})
            workplace = src.get("workplace_address", {}) or {}
            jobs.append({
                "title": src.get("headline", ""),
                "company": src.get("employer", {}).get("name", "N/A") if isinstance(src.get("employer"), dict) else "N/A",
                "location": workplace.get("city", "") if isinstance(workplace, dict) else "",
                "description": src.get("description", {}).get("text", "")[:500] if isinstance(src.get("description"), dict) else "",
                "url": j.get("_source", {}).get("webpage_url", ""),
                "date": src.get("publication_date", "")[:10] if src.get("publication_date") else "",
                "source": "🇸🇪 Jobtechdev",
                "language": "sv"
            })
        return jobs
    except Exception as e:
        print(f"Jobtechdev API Error: {e}")
        return []

# 丹麦 Jobinsats API
JOBINSATS_API_KEY = os.environ.get("JOBINSATS_API_KEY", "00e043b3fc3a0d9ab5eb956ed644f113c3856175fc96fd54")
JOBINSATS_BASE = "https://api.jobindsats.dk"

def search_denmark(keyword: str, location: str = "", limit: int = 10) -> List[Dict]:
    """搜索丹麦职位 - Jobinsats API"""
    url = f"{JOBINSATS_BASE}/stillingsopslag"
    params = {
        "q": keyword,
        "size": min(limit, 50)
    }
    if location:
        params["place"] = location
    
    headers = {
        "Authorization": f"Bearer {JOBINSATS_API_KEY}",
        "Accept": "application/json"
    }
    
    try:
        r = requests.get(url, params=params, headers=headers, timeout=15)
        if r.status_code != 200:
            print(f"Jobinsats API Error: {r.status_code} - {r.text[:200]}")
            return []
        
        data = r.json()
        jobs = []
        
        # Jobinsats 返回格式
        for j in data.get("stillingsopslag", [])[:limit]:
            jobs.append({
                "title": j.get("titel", ""),
                "company": j.get("arbejdsgiver", {}).get("navn", "N/A") if isinstance(j.get("arbejdsgiver"), dict) else str(j.get("arbejdsgiver", "N/A")),
                "location": j.get("arbejdssted", {}).get("adresse", "") if isinstance(j.get("arbejdssted"), dict) else str(j.get("arbejdssted", "")),
                "description": j.get("bruttoarbejdstid", ""),
                "url": j.get("self", ""),
                "date": j.get("publiceringsdato", "")[:10] if j.get("publiceringsdato") else "",
                "source": "🇩🇰 Jobinsats",
                "language": "da"
            })
        return jobs
    except Exception as e:
        print(f"Jobinsats API Error: {e}")
        return []

# 美国 USAJOBS API
USAJOBS_API_KEY = os.environ.get("USAJOBS_API_KEY", "UiuR7SrGKS0sj3eFFKOFqnACxbc4+oUjPuvLRiY38nU=")
USAJOBS_BASE = "https://data.usajobs.gov/api/search"

def search_usa(keyword: str, location: str = "", limit: int = 10) -> List[Dict]:
    """搜索美国职位 - USAJOBS API"""
    url = USAJOBS_BASE
    params = {
        "Keyword": keyword,
        "ResultsPerPage": min(limit, 25)
    }
    if location:
        params["LocationName"] = location
    
    headers = {
        "Authorization-Key": USAJOBS_API_KEY,
        "User-Agent": "JobMatchAI/1.0 (wei.li@outlook.dk)"
    }
    
    try:
        r = requests.get(url, params=params, headers=headers, timeout=15)
        if r.status_code != 200:
            print(f"USAJOBS API Error: {r.status_code} - {r.text[:200]}")
            return []
        
        data = r.json()
        jobs = []
        
        # 检查数据结构
        search_result = data.get("SearchResult") if isinstance(data, dict) else {}
        if isinstance(search_result, dict):
            items = search_result.get("SearchResultItems", [])
        else:
            items = []
        
        for j in items[:limit]:
            if not isinstance(j, dict):
                continue
            src = j.get("MatchedObjectDescriptor", {})
            if not isinstance(src, dict):
                continue
            jobs.append({
                "title": src.get("Title", ""),
                "company": src.get("OrganizationName", "N/A"),
                "location": src.get("LocationName", ""),
                "description": src.get("UserArea", {}).get("Details", "")[:500] if isinstance(src.get("UserArea"), dict) and src.get("UserArea", {}).get("Details") else "",
                "url": src.get("PositionURI", ""),
                "date": src.get("PublicationStartDate", "")[:10] if src.get("PublicationStartDate") else "",
                "source": "🇺🇸 USAJOBS",
                "language": "en"
            })
        return jobs
    except Exception as e:
        print(f"USAJOBS API Error: {e}")
        return []

# === 中国职位搜索 ===
# 注意：中国主流招聘平台（前程无忧、智联、Boss直聘）均无公开免费API
# 搜索中国职位请用户粘贴职位文本，详见 paste 模式
# 中国职位需要用户手动粘贴，不支持自动搜索

# === 启动时初始化IP地理数据库 ===
import ip_geo
ip_geo.init()  # 使用内置的中国IP段数据库
print(f"✅ IP地理数据库已加载（{len(ip_geo._CHINA_IP_RANGES)} 个IP段）")


@app.get("/")
def read_root(request: Request):
    """根据客户端IP地理位置返回前端页面（大陆IP→中文版，其他→英文版）"""
    # 获取客户端IP（考虑代理）
    client_ip = request.headers.get("x-forwarded-for", "").split(",")[0].strip()
    if not client_ip:
        client_ip = request.client.host if request.client else "127.0.0.1"
    
    # 根据IP地理位置判断
    is_china = ip_geo.is_china_ip(client_ip) if hasattr(ip_geo, '_CHINA_IP_RANGES') and ip_geo._CHINA_IP_RANGES else False
    
    if is_china:
        # 中国大陆IP → 中文版
        frontend_path = os.path.join(os.path.dirname(__file__), "frontend", "index-zh.html")
    else:
        # 其他地区IP（丹麦、欧洲等）→ 英文版（目前内容也是中文）
        frontend_path = os.path.join(os.path.dirname(__file__), "frontend", "index-en.html")
    
    if os.path.exists(frontend_path):
        return FileResponse(frontend_path)
    return {"message": "JobMatchAI Nordic API", "version": "2.0.0"}

@app.get("/beta")
def read_beta_page():
    """返回Beta测试引导页面"""
    beta_path = os.path.join(os.path.dirname(__file__), "frontend", "beta.html")
    if os.path.exists(beta_path):
        return FileResponse(beta_path)
    return {"error": "Beta page not found"}

@app.get("/en")
def read_english_page():
    """返回英文版页面"""
    en_path = os.path.join(os.path.dirname(__file__), "frontend", "index-en.html")
    if os.path.exists(en_path):
        return FileResponse(en_path)
    return {"error": "English page not found"}

@app.get("/v3")
def read_v3_page():
    """返回V3测试版页面"""
    v3_path = os.path.join(os.path.dirname(__file__), "frontend", "index-v3.html")
    if os.path.exists(v3_path):
        return FileResponse(v3_path)
    return {"error": "V3 page not found"}

@app.get("/health")
def health_check():
    """健康检查"""
    return {"status": "ok", "service": "JobMatchAI Nordic"}

@app.post("/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """上传并解析简历（升级版：AI结构化解析）"""
    try:
        content = await file.read()
        text = parse_resume(content, file.filename)
        lang = detect_language(text)
        
        # 使用 AI 进行结构化解析
        parsed_structure = parse_resume_structure(text, lang)
        
        return {
            "success": True,
            "filename": file.filename,
            "text": text,
            "detected_language": lang,
            "language_name": {"zh": "中文", "en": "English", "da": "Dansk"}.get(lang, "English"),
            # 结构化数据（新功能）
            "structure": {
                "personal": parsed_structure.get("personal", {}),
                "education": parsed_structure.get("education", []),
                "experience": parsed_structure.get("experience", []),
                "skills": parsed_structure.get("skills", {}),
                "total_experience_years": parsed_structure.get("total_experience_years", 0),
                "ai_enhanced": parsed_structure.get("ai_enhanced", False)
            }
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/upload-job-document")
async def upload_job_document(file: UploadFile = File(...)):
    """上传并解析职位文档（PDF/DOCX/TXT）"""
    try:
        content = await file.read()
        text = parse_resume(content, file.filename)
        lang = detect_language(text)
        
        # 使用 AI 提取职位信息
        job_info = extract_job_from_text(text, lang)
        
        return {
            "success": True,
            "filename": file.filename,
            "text": text,
            "detected_language": lang,
            "language_name": {"zh": "中文", "en": "English", "da": "Dansk"}.get(lang, "English"),
            "job_info": job_info
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def extract_job_from_text(text: str, lang: str = 'en') -> Dict:
    """从职位文档文本中提取结构化职位信息"""
    if not AI_AVAILABLE:
        # Fallback: 返回原始文本作为描述
        return {
            "title": "",
            "company": "",
            "description": text[:2000],
            "requirements": [],
            "location": "",
            "success": True,
            "source": "Document Upload"
        }
    
    try:
        prompt = f"""You are a job description parser. Extract structured information from the following job posting text.

Text:
{text[:4000]}

Return JSON format:
{{
    "title": "Job title",
    "company": "Company name",
    "description": "Brief job description (max 500 chars)",
    "requirements": ["requirement 1", "requirement 2"],
    "location": "Job location",
    "success": true
}}

If information is not found, use empty string or empty array."""

        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a job description parser. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            preferred_provider="groq",
            temperature=0.3,
            max_tokens=1500
        )
        
        content = response.choices[0].message.content.strip()
        if content.startswith('```'):
            content = content.split('```')[1]
            if content.startswith('json'):
                content = content[4:]
        
        job_info = json.loads(content)
        job_info['success'] = True
        job_info['source'] = 'Document Upload'
        return job_info
        
    except Exception as e:
        print(f"Job extraction failed: {e}")
        return {
            "title": "",
            "company": "",
            "description": text[:2000],
            "requirements": [],
            "location": "",
            "success": True,
            "source": "Document Upload"
        }

@app.post("/analyze-resume")
async def analyze_resume_endpoint(
    resume_text: str = Form(...),
    language: str = Form("auto"),
    job_title: str = Form(""),
    job_description: str = Form(""),
    company: str = Form("")
):
    """分析简历并返回改进建议
    
    可选参数:
    - job_title: 目标职位名称
    - job_description: 职位描述
    - company: 公司名称
    
    如果提供职位上下文，将返回针对该职位的匹配分析和改进建议
    """
    try:
        if language == "auto":
            lang = detect_language(resume_text)
        else:
            lang = language
        
        # 构建职位上下文
        job_context = None
        if job_title or job_description:
            job_context = {
                'job_title': job_title,
                'job_description': job_description,
                'company': company
            }
        
        analysis = analyze_resume(resume_text, lang, job_context)
        
        return {
            "success": True,
            "analysis": analysis,
            "detected_language": lang
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/generate-cover-letter")
async def generate_cover_letter_endpoint(
    resume_text: str = Form(...),
    job_title: str = Form(...),
    company: str = Form(...),
    job_description: str = Form(...),
    language: str = Form("da"),
    resume_highlights: str = Form(""),  # JSON string of highlights from polish
    resume_structure: str = Form("")  # JSON string of structured resume data
):
    """生成求职信（升级版：支持简历亮点 + 结构化数据防瞎编）"""
    try:
        job = {
            "title": job_title,
            "company": company,
            "description": job_description,
            "source": "User Input"
        }
        
        # 解析简历亮点
        highlights = []
        if resume_highlights:
            try:
                highlights = json.loads(resume_highlights)
            except:
                pass
        
        # 解析结构化简历数据
        structure_data = None
        if resume_structure:
            try:
                structure_data = json.loads(resume_structure)
            except:
                pass
        
        # 如果没有提供亮点，先调用精修获取
        if not highlights and AI_AVAILABLE:
            job_context = {
                "job_title": job_title,
                "job_description": job_description,
                "company": company
            }
            lang = language if language != "auto" else detect_language(resume_text)
            polish_result = generate_polish_suggestions(resume_text, lang, job_context, structure_data)
            if isinstance(polish_result, dict):
                highlights = polish_result.get('resume_highlights', [])
        
        cover_letter = generate_cover_letter_with_ai(resume_text, job, language, highlights, structure_data)
        
        return {
            "success": True,
            "cover_letter": cover_letter,
            "language": language,
            "highlights_used": len(highlights) > 0
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/detect-job-language")
def detect_job_language_endpoint(text: str):
    """检测职位描述语言"""
    lang = detect_language(text)
    return {
        "language": lang,
        "language_name": {"zh": "中文", "en": "English", "da": "Dansk"}.get(lang, "English")
    }


# === 智能职位匹配 API（新增）===

from smart_matcher import UserProfileManager, SmartJobMatcher, AutoJobProcessor
from email_reader import JobEmailReader
from linkedin_importer import linkedin_importer, detect_job_language as linkedin_detect_language

# 初始化管理器
profile_manager = UserProfileManager()
job_matcher = SmartJobMatcher(profile_manager)
auto_processor = AutoJobProcessor(profile_manager, job_matcher)

class EmailConfig(BaseModel):
    email: str
    password: str
    imap_server: Optional[str] = None

class JobAction(BaseModel):
    user_id: str
    job: Dict
    action: str  # 'view', 'save', 'apply', 'ignore'

@app.post("/user-profile/create")
async def create_user_profile(user_id: str = Form(...), resume_text: str = Form(...)):
    """从简历创建用户画像"""
    try:
        profile = profile_manager.create_profile(user_id, resume_text)
        return {
            "success": True,
            "profile": {
                "user_id": profile.user_id,
                "skills": profile.skills,
                "total_years": profile.total_years,
                "industries": profile.industries,
                "roles": profile.roles
            }
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/user-profile/{user_id}")
async def get_user_profile(user_id: str):
    """获取用户画像"""
    profile = profile_manager.get_profile(user_id)
    if not profile:
        return {"success": False, "error": "Profile not found"}
    
    return {
        "success": True,
        "profile": profile.to_dict()
    }

@app.post("/user-profile/update-from-resume")
async def update_profile_from_resume(user_id: str = Form(...), resume_text: str = Form(...)):
    """根据更新的简历刷新用户画像"""
    try:
        profile = profile_manager.update_profile_from_resume(user_id, resume_text)
        return {
            "success": True,
            "profile": {
                "user_id": profile.user_id,
                "skills": profile.skills,
                "total_years": profile.total_years,
                "updated_at": profile.updated_at
            }
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

# ========== 用户粘性增强 API ==========

@app.get("/user-profile/{user_id}/summary")
async def get_profile_summary(user_id: str):
    """获取用户画像摘要"""
    try:
        summary = profile_manager.get_profile_summary(user_id)
        return {"success": True, "summary": summary}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/user-profile/resume-history")
async def add_resume_history(
    user_id: str = Form(...),
    resume_text: str = Form(...),
    polish_summary: str = Form(...),
    job_context: str = Form(...),
    match_score: float = Form(...),
    skill_gaps: str = Form(...)  # JSON string
):
    """记录简历精修历史"""
    try:
        import json
        gaps = json.loads(skill_gaps) if skill_gaps else []
        profile = profile_manager.add_resume_history(
            user_id, resume_text, polish_summary, 
            job_context, match_score, gaps
        )
        return {"success": True, "version": len(profile.resume_history)}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/user-profile/skill-gap")
async def add_skill_gap(
    user_id: str = Form(...),
    job_title: str = Form(...),
    job_description: str = Form(...),
    resume_skills: str = Form(...),  # JSON string
    required_skills: str = Form(...),  # JSON string
    missing_skills: str = Form(...),  # JSON string
    match_score: float = Form(...),
    learning_recommendations: str = Form(...)  # JSON string
):
    """记录Skill Gap分析"""
    try:
        import json
        resume_skills_list = json.loads(resume_skills) if resume_skills else []
        required_skills_list = json.loads(required_skills) if required_skills else []
        missing_skills_list = json.loads(missing_skills) if missing_skills else []
        learning_recs = json.loads(learning_recommendations) if learning_recommendations else []
        
        profile = profile_manager.add_skill_gap(
            user_id, job_title, job_description,
            resume_skills_list, required_skills_list,
            missing_skills_list, match_score, learning_recs
        )
        return {"success": True, "gap_count": len(profile.skill_gap_history)}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/user-profile/application")
async def add_application(
    user_id: str = Form(...),
    job_title: str = Form(...),
    company: str = Form(...),
    cover_letter: str = Form(...),
    application_status: str = Form("sent")
):
    """记录申请历史"""
    try:
        profile = profile_manager.add_application(
            user_id, job_title, company, cover_letter, application_status
        )
        return {"success": True, "application_count": len(profile.application_history)}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/user-profile/learning")
async def add_learning_record(
    user_id: str = Form(...),
    skill: str = Form(...),
    resource_title: str = Form(...),
    resource_url: str = Form(...),
    status: str = Form("viewed"),
    completion_percent: int = Form(0)
):
    """记录学习资源浏览"""
    try:
        profile = profile_manager.add_learning_record(
            user_id, skill, resource_title, resource_url,
            status, completion_percent
        )
        return {"success": True, "learning_count": len(profile.learning_history)}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/user-profile/{user_id}/gap-trend")
async def get_skill_gap_trend(user_id: str):
    """获取Skill Gap改进趋势"""
    try:
        profile = profile_manager.get_profile(user_id)
        if not profile:
            return {"success": False, "error": "Profile not found"}
        
        # 计算Gap改进趋势
        if len(profile.skill_gap_history) < 2:
            return {"success": True, "trend": [], "message": "Need more data"}
        
        recent = profile.skill_gap_history[-5:]  # 最近5次
        trend = []
        for i, gap in enumerate(recent):
            # 计算Gap减少情况
            previous_gap = recent[i-1] if i > 0 else None
            gap_count = len(gap.get('missing_skills', []))
            
            trend.append({
                'timestamp': gap['timestamp'],
                'job_title': gap['job_title'],
                'match_score': gap['match_score'],
                'missing_skills_count': gap_count,
                'improved': previous_gap and gap_count < len(previous_gap.get('missing_skills', []))
            })
        
        return {"success": True, "trend": trend}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/jobs/fetch-from-email")
async def fetch_jobs_from_email(config: EmailConfig):
    """从用户邮箱读取招聘邮件"""
    try:
        reader = JobEmailReader(
            email_address=config.email,
            password=config.password,
            imap_server=config.imap_server
        )
        
        jobs = reader.fetch_job_emails(limit=50)
        reader.disconnect()
        
        return {
            "success": True,
            "count": len(jobs),
            "jobs": jobs
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/jobs/match")
async def match_jobs_endpoint(
    user_id: str = Form(...),
    jobs: str = Form(...)  # JSON字符串
):
    """智能匹配职位"""
    try:
        jobs_list = json.loads(jobs)
        matched_jobs = job_matcher.filter_and_rank_jobs(user_id, jobs_list, min_score=60.0)
        
        return {
            "success": True,
            "total": len(jobs_list),
            "matched": len(matched_jobs),
            "jobs": matched_jobs
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/jobs/process-email-batch")
async def process_email_batch(
    user_id: str = Form(...),
    email_config: str = Form(...)  # JSON字符串
):
    """一键处理：读取邮件 → 智能筛选 → 返回匹配职位"""
    try:
        config = json.loads(email_config)
        
        # 1. 读取邮件
        reader = JobEmailReader(
            email_address=config['email'],
            password=config['password'],
            imap_server=config.get('imap_server')
        )
        jobs = reader.fetch_job_emails(limit=50)
        reader.disconnect()
        
        # 2. 智能筛选
        result = auto_processor.process_email_jobs(user_id, jobs)
        
        # 3. 生成每日摘要
        digest = auto_processor.generate_daily_digest(user_id, jobs)
        
        return {
            "success": True,
            "summary": result['summary'],
            "high_match": result['high_match'],
            "medium_match": result['medium_match'],
            "digest": digest
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/jobs/action")
async def record_job_action(action: JobAction):
    """记录用户对职位的操作（用于学习偏好）"""
    try:
        profile_manager.learn_from_job_action(action.user_id, action.job, action.action)
        return {"success": True, "message": f"Recorded {action.action} action"}
    except Exception as e:
        return {"success": False, "error": str(e)}

# === 职位搜索 API ===
@app.get("/jobs/search")
async def search_jobs(
    keyword: str = "",
    country: str = "all",
    location: str = "",
    limit: int = 10
):
    """统一职位搜索 API
    
    支持的 country 参数:
    - all: 搜索所有来源（英国、澳大利亚、加拿大、法国、荷兰、比利时、美国、德国）
    - gb: 英国 (Adzuna)
    - au: 澳大利亚 (Adzuna)
    - ca: 加拿大 (Adzuna)
    - fr: 法国 (Adzuna)
    - nl: 荷兰 (Adzuna)
    - be: 比利时 (Adzuna)
    - us: 美国 (Adzuna)
    - de: 德国 (Arbeitsagentur)
    """
    all_jobs = []
    
    try:
        if country == "all":
            # 搜索所有来源
            all_jobs.extend(search_adzuna(keyword, "gb", location, limit))
            all_jobs.extend(search_adzuna(keyword, "us", location, limit))
            all_jobs.extend(search_adzuna(keyword, "au", location, limit))
            all_jobs.extend(search_adzuna(keyword, "ca", location, limit))
            all_jobs.extend(search_adzuna(keyword, "fr", location, limit))
            all_jobs.extend(search_adzuna(keyword, "nl", location, limit))
            all_jobs.extend(search_adzuna(keyword, "be", location, limit))
            all_jobs.extend(search_germany(keyword, location, limit))
        elif country in ["gb", "au", "ca", "fr", "nl", "be", "us"]:
            all_jobs.extend(search_adzuna(keyword, country, location, limit))
        elif country == "de":
            all_jobs.extend(search_germany(keyword, location, limit))
        else:
            return {"success": False, "error": f"不支持的国家: {country}"}
        
        return {
            "success": True,
            "count": len(all_jobs),
            "jobs": all_jobs[:limit * 3] if country == "all" else all_jobs
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/jobs/sources")
def get_job_sources():
    """获取支持的职位来源"""
    return {
        "success": True,
        "sources": [
            {"id": "all", "name": "全部来源", "flag": "🌐"},
            {"id": "gb", "name": "英国", "flag": "🇬🇧", "api": "Adzuna"},
            {"id": "au", "name": "澳大利亚", "flag": "🇦🇺", "api": "Adzuna"},
            {"id": "ca", "name": "加拿大", "flag": "🇨🇦", "api": "Adzuna"},
            {"id": "fr", "name": "法国", "flag": "🇫🇷", "api": "Adzuna"},
            {"id": "nl", "name": "荷兰", "flag": "🇳🇱", "api": "Adzuna"},
            {"id": "be", "name": "比利时", "flag": "🇧🇪", "api": "Adzuna"},
            {"id": "de", "name": "德国", "flag": "🇩🇪", "api": "Arbeitsagentur"},
            {"id": "se", "name": "瑞典", "flag": "🇸🇪", "api": "Jobtechdev"},
            {"id": "dk", "name": "丹麦", "flag": "🇩🇰", "api": "Jobinsats"},
            {"id": "us", "name": "美国", "flag": "🇺🇸", "api": "USAJOBS"},
        ]
    }


# === 智能简历精修 API（行为数据挖掘）===

from resume_enhancer import (
    BehaviorDataCollector, 
    ResumeEnhancer, 
    DynamicResumeOptimizer,
    HiddenSkill,
    ForgottenExperience
)
from resume_extractor import (
    ResumeExtractor,
    ResumeJobMatcher,
    BatchJobMatcher,
    ResumeProfile,
    MatchResult
)

# 初始化
resume_enhancer = ResumeEnhancer(ai_client) if AI_AVAILABLE else ResumeEnhancer()
dynamic_optimizer = DynamicResumeOptimizer(ai_client) if AI_AVAILABLE else DynamicResumeOptimizer()
resume_extractor = ResumeExtractor()
job_matcher = ResumeJobMatcher()
batch_matcher = BatchJobMatcher()


@app.post("/enhance/collect-behavior")
async def collect_behavior_data(
    user_id: str = Form(...),
    emails: str = Form("[]"),  # JSON字符串
    calendar_events: str = Form("[]"),  # JSON字符串
    documents: str = Form("[]")  # JSON字符串
):
    """收集用户行为数据（邮件、日历、文档）
    
    用于挖掘用户的隐性能力和可能被遗忘的经历
    """
    try:
        import json
        
        # 解析JSON
        emails_list = json.loads(emails) if emails else []
        events_list = json.loads(calendar_events) if calendar_events else []
        docs_list = json.loads(documents) if documents else []
        
        # 收集数据
        collector = BehaviorDataCollector()
        email_data = collector.collect_from_email(emails_list)
        calendar_data = collector.collect_from_calendar(events_list)
        doc_data = collector.collect_from_documents(docs_list)
        
        return {
            "success": True,
            "collected": {
                "email_count": len(emails_list),
                "calendar_events_count": len(events_list),
                "documents_count": len(docs_list)
            },
            "email_data": email_data,
            "calendar_data": calendar_data,
            "doc_data": doc_data
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/enhance/resume")
async def enhance_resume(
    user_id: str = Form(...),
    resume_text: str = Form(...),
    resume_skills: str = Form("[]"),  # JSON字符串
    emails: str = Form("[]"),
    calendar_events: str = Form("[]"),
    documents: str = Form("[]")
):
    """智能简历精修 - 挖掘隐性能力 + 发现遗忘经历
    
    通过分析用户的邮件、日历、文档等行为数据，
    发现简历中可能未体现的隐性能力和可能被遗忘的项目/培训/证书。
    """
    try:
        import json
        
        # 解析数据
        skills_list = json.loads(resume_skills) if resume_skills else []
        emails_list = json.loads(emails) if emails else []
        events_list = json.loads(calendar_events) if calendar_events else []
        docs_list = json.loads(documents) if documents else []
        
        # 收集行为数据
        collector = BehaviorDataCollector()
        email_data = collector.collect_from_email(emails_list)
        calendar_data = collector.collect_from_calendar(events_list)
        doc_data = collector.collect_from_documents(docs_list)
        
        # 分析并生成增强建议
        enhancement = resume_enhancer.analyze_and_enhance(
            user_id=user_id,
            resume_text=resume_text,
            resume_skills=skills_list,
            email_data=email_data,
            calendar_data=calendar_data,
            doc_data=doc_data
        )
        
        return {
            "success": True,
            "hidden_skills": [s.to_dict() for s in enhancement.hidden_skills],
            "forgotten_experiences": [e.to_dict() for e in enhancement.forgotten_experiences],
            "keyword_additions": enhancement.keyword_additions,
            "keyword_weights": enhancement.keyword_weights,
            "suggestions": enhancement.suggestions,
            "ai_insights": enhancement.ai_insights
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/enhance/dynamic-optimize")
async def dynamic_optimize_resume(
    user_id: str = Form(...),
    resume_text: str = Form(...),
    resume_skills: str = Form("[]"),
    job_title: str = Form(...),
    job_company: str = Form(""),
    job_description: str = Form("")
):
    """动态简历优化 - 针对特定职位优化简历
    
    根据目标职位的要求，动态调整简历关键词和表述方式，
    最大化简历与职位的匹配度。
    """
    try:
        import json
        
        skills_list = json.loads(resume_skills) if resume_skills else []
        
        target_job = {
            "title": job_title,
            "company": job_company,
            "description": job_description
        }
        
        # 执行动态优化
        result = dynamic_optimizer.optimize_for_job(
            user_id=user_id,
            resume_text=resume_text,
            resume_skills=skills_list,
            target_job=target_job
        )
        
        return {
            "success": True,
            **result
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/enhance/sources")
def get_behavior_sources():
    """获取支持的行为数据来源"""
    return {
        "success": True,
        "sources": [
            {
                "id": "email",
                "name": "邮件",
                "description": "分析邮件主题和发件人，挖掘项目经验和协作技能",
                "icon": "📧",
                "example": "从'Re: ERP Migration Kickoff'中发现ERP迁移项目经验"
            },
            {
                "id": "calendar",
                "name": "日历",
                "description": "分析日历事件，发现培训和会议参与情况",
                "icon": "📅",
                "example": "从'SAP Certification Workshop'中发现可能的证书"
            },
            {
                "id": "documents",
                "name": "文档/笔记",
                "description": "分析文档标题和笔记内容，发现专业技能",
                "icon": "📄",
                "example": "从技术文档标题中发现PowerBI、Python等技能"
            },
            {
                "id": "job_views",
                "name": "浏览记录",
                "description": "分析浏览过的职位，发现用户的真实兴趣",
                "icon": "👀",
                "example": "通过多次浏览D365职位推断用户对微软生态的偏好"
            }
        ],
        "privacy_note": "所有数据仅在本地处理，不会上传到任何服务器"
    }


# === 简历精准抓取 API ===

@app.post("/extract/resume")
async def extract_resume(resume_text: str = Form(...)):
    """精准抓取简历结构化信息
    
    从简历文本中提取：
    - 基本信息（姓名、邮箱、电话、地点）
    - 工作经历（公司、职位、时间、描述、技能）
    - 教育背景
    - 技能列表
    - 语言能力
    - 证书
    """
    try:
        profile = resume_extractor.extract(resume_text)
        return {
            "success": True,
            "profile": profile.to_dict()
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/extract/match")
async def match_resume_to_job(
    resume_text: str = Form(...),
    job: str = Form(...)  # JSON字符串
):
    """简历与单个职位精准匹配
    
    返回：
    - 综合匹配分数
    - 技能匹配率
    - 经验匹配率
    - 匹配/缺失技能列表
    - 能力差距分析
    - 简历优化建议
    """
    try:
        job_dict = json.loads(job)
        match_result = job_matcher.match(resume_text, job_dict)
        return {
            "success": True,
            "match": match_result.to_dict(),
            "job_title": job_dict.get("title", ""),
            "company": job_dict.get("company", "")
        }
    except Exception as e:
        return {"success": False, "error": str(e)}




@app.post("/extract/batch-match")
async def batch_match_jobs(
    resume_text: str = Form(...),
    jobs: str = Form(...),  # JSON字符串
    min_score: float = Form(50.0)
):
    """批量匹配简历与多个职位
    
    对所有职位按匹配度排序，返回排名结果。
    """
    try:
        jobs_list = json.loads(jobs)
        ranked_jobs = batch_matcher.rank_jobs(resume_text, jobs_list, min_score)
        
        return {
            "success": True,
            "total_jobs": len(jobs_list),
            "matched_jobs": len(ranked_jobs),
            "ranked_jobs": ranked_jobs
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/extract/skill-gap")
async def analyze_skill_gap(
    resume_text: str = Form(...),
    job_description: str = Form(...)
):
    """分析简历与职位描述之间的技能差距
    
    比对简历中已有的技能与职位要求，识别：
    - 完全匹配的技能
    - 部分匹配的技能
    - 完全缺失的技能
    """
    try:
        profile = resume_extractor.extract(resume_text)
        job_dict = {"title": "", "description": job_description, "company": ""}
        match_result = job_matcher.match(resume_text, job_dict)
        
        return {
            "success": True,
            "resume_skills": profile.skills,
            "matched_skills": match_result.matched_skills,
            "partial_skills": match_result.partial_skills,
            "missing_skills": match_result.missing_skills,
            "skill_match_rate": match_result.skill_match_rate,
            "experience_years": profile.total_years,
            "work_experiences": [e.to_dict() for e in profile.work_experiences]
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# === AI学习资源推荐 API ===
@app.post("/api/learning/recommend")
async def get_learning_recommendations(
    resume_structure: Dict = None,
    job_description: str = "",
    skill_gap: Dict = None,
    lang: str = "en",
    # 支持前端简化的参数格式
    missing_skills: List[str] = None,
    job_title: str = "",
    language: str = "en"
):
    """根据简历结构和技能差距推荐个性化学习资源
    
    Args:
        resume_structure: 解析后的简历结构化数据（完整格式）
        job_description: 职位描述
        skill_gap: 技能差距分析结果（完整格式）
        lang: 语言 (zh/en)
        # 以下为简化的前端参数
        missing_skills: 缺失技能列表（简化格式）
        job_title: 职位名称（简化格式）
        language: 语言（简化格式）
    """
    try:
        # 处理简化的前端参数格式
        if not resume_structure:
            # 前端使用简化参数，构建 skill_gap
            if missing_skills is not None:
                resume_structure = {}  # 空简历结构，使用技能差距驱动
                skill_gap = {
                    "missing_skills": missing_skills,
                    "critical_gaps": [],
                    "matched_skills": [],
                    "job_title": job_title,
                    "reasoning": ""
                }
                # 使用简化参数中的 language
                if language:
                    lang = language
            else:
                return {"success": False, "error": "Missing required parameters"}
        
        if not skill_gap:
            skill_gap = {
                "missing_skills": [],
                "critical_gaps": [],
                "matched_skills": [],
                "score": 50,
                "reasoning": ""
            }
        
        result = recommend_learning_resources(
            resume_structure=resume_structure,
            job_description=job_description,
            skill_gap=skill_gap,
            lang=lang
        )
        
        return result
        
    except Exception as e:
        print(f"Learning recommendation error: {e}")
        return {"success": False, "error": str(e)}


# === 验证学习资源链接有效性 ===
@app.post("/api/learning/validate-links")
async def validate_learning_links(urls: List[str] = None):
    """验证学习资源链接的有效性
    
    Args:
        urls: 要验证的链接列表
    """
    import asyncio
    import httpx
    
    if not urls:
        return {"success": False, "error": "No URLs provided"}
    
    async def check_url(url: str) -> Dict:
        """检查单个URL的有效性"""
        try:
            async with httpx.AsyncClient(timeout=10.0) as client:
                response = await client.get(url, follow_redirects=True)
                if response.status_code == 200:
                    return {"url": url, "valid": True, "status": 200}
                elif response.status_code == 404:
                    return {"url": url, "valid": False, "status": 404, "error": "页面不存在"}
                else:
                    return {"url": url, "valid": False, "status": response.status_code, "error": f"HTTP {response.status_code}"}
        except httpx.TimeoutException:
            return {"url": url, "valid": False, "error": "请求超时"}
        except httpx.RequestError as e:
            return {"url": url, "valid": False, "error": "无法访问"}
        except Exception as e:
            return {"url": url, "valid": False, "error": str(e)}
    
    # 并发验证所有链接
    results = await asyncio.gather(*[check_url(url) for url in urls[:10]])  # 最多验证10个
    
    valid_count = sum(1 for r in results if r.get('valid', False))
    
    return {
        "success": True,
        "total": len(results),
        "valid": valid_count,
        "invalid": len(results) - valid_count,
        "results": results,
        "recommendation": f"{(valid_count / len(results) * 100):.0f}%链接有效" if results else "无链接可验证"
    }


# === 求职信质量测试 API ===

class CoverLetterQualityChecker:
    """求职信质量检查器"""

    def __init__(self):
        pass

    def check(self, cover_letter: str, job: Dict, resume_text: str = "") -> Dict:
        """检查求职信质量"""
        scores = {}
        issues = []
        suggestions = []

        # 1. 长度检查
        word_count = len(cover_letter.split())
        if 200 <= word_count <= 400:
            scores['length'] = 100
        elif 150 <= word_count < 200 or 400 < word_count <= 500:
            scores['length'] = 70
        else:
            scores['length'] = 40
            issues.append(f"字数 {word_count} 字，理想范围 200-400 词")
            suggestions.append("将求职信长度调整到 200-400 词之间")

        # 2. 结构检查
        structure_keywords = {
            'opening': ['dear', 'dearest', 'hello', '尊敬的', 'kære'],
            'closing': ['regards', 'sincerely', 'best', '致敬', 'med venlig hilsen']
        }
        structure_score = 0
        for section, keywords in structure_keywords.items():
            if any(kw in cover_letter.lower() for kw in keywords):
                structure_score += 50
        scores['structure'] = structure_score
        if structure_score < 100:
            issues.append("求职信缺少开头或结尾称呼")
            suggestions.append("确保包含正式开头（Dear...）和结尾（Best regards...）")

        # 3. 关键词匹配检查
        if job.get('description'):
            job_keywords = set(job['description'].lower().split())
            cover_lower = cover_letter.lower()
            matched = sum(1 for kw in job_keywords if kw in cover_lower and len(kw) > 3)
            total = sum(1 for kw in job_keywords if len(kw) > 3)
            keyword_rate = (matched / max(total, 1)) * 100
            scores['keyword_match'] = min(keyword_rate * 2, 100)
            if keyword_rate < 30:
                suggestions.append("求职信中职位关键词覆盖率较低，建议更多引用职位描述中的关键要求")

        # 4. ATS友好度检查
        ats_score = 100
        bad_patterns = ['***', '████', '[REDACTED]', 'confidential']
        for pat in bad_patterns:
            if pat in cover_letter:
                ats_score -= 30
        if not re.search(r'[a-zA-Z]', cover_letter):
            ats_score = 0
        scores['ats_friendly'] = max(ats_score, 0)

        # 5. 强动词检查
        strong_verbs = ['achieved', 'led', 'implemented', 'improved', 'delivered', 
                        'increased', 'reduced', 'optimized', '主导', '实现', '提升']
        weak_verbs = ['did', 'worked', 'helped', '参与', '做了']
        
        strong_count = sum(1 for v in strong_verbs if v in cover_letter.lower())
        weak_count = sum(1 for v in weak_verbs if v in cover_letter.lower())
        
        if strong_count >= 3 and weak_count <= 1:
            scores['action_verbs'] = 100
        elif strong_count >= 1:
            scores['action_verbs'] = 60
        else:
            scores['action_verbs'] = 30
            suggestions.append("使用更有力的动词（Led, Achieved, Implemented 等）替代弱动词")

        # 6. 个性化检查（是否提到公司名）
        company_mentioned = job.get('company', '') and \
                             job['company'].lower() in cover_letter.lower()
        scores['personalization'] = 100 if company_mentioned else 50
        if not company_mentioned:
            suggestions.append(f"建议在求职信中提到公司名称：{job.get('company', '目标公司')}")

        # 7. 简历匹配度
        resume_match = 50
        if resume_text:
            resume_words = set(resume_text.lower().split())
            cover_words = set(cover_letter.lower().split())
            overlap = len(resume_words & cover_words) / max(len(resume_words), 1)
            resume_match = overlap * 100
        scores['resume_relevance'] = resume_match

        # 综合评分 - 申请成功率视角
        # 核心：简历-职位匹配度 和 关键词匹配率 是申请成功的关键
        weights = {
            'length': 0.10,           # 长度（降低权重）
            'structure': 0.10,       # 结构（降低权重）
            'keyword_match': 0.30,   # 关键词匹配（提高！说明理解职位要求）
            'ats_friendly': 0.10,    # ATS友好（降低权重）
            'action_verbs': 0.10,    # 行动词（降低权重）
            'personalization': 0.10, # 个性化
            'resume_relevance': 0.20 # 简历-职位关联（提高！核心指标）
        }
        overall = sum(scores[k] * weights[k] for k in weights)

        # 评分等级 - 申请成功率视角
        if overall >= 80:
            grade = "🌟 高度推荐"
        elif overall >= 65:
            grade = "✅ 推荐申请"
        elif overall >= 50:
            grade = "⚠️ 谨慎申请"
        else:
            grade = "❌ 不建议申请"

        return {
            "overall_score": round(overall, 1),
            "grade": grade,
            "scores": {k: round(v, 1) for k, v in scores.items()},
            "issues": issues,
            "suggestions": suggestions,
            "word_count": word_count,
            "company_mentioned": company_mentioned
        }


# 初始化求职信检查器
cover_letter_checker = CoverLetterQualityChecker()


@app.post("/cover-letter/check")
async def check_cover_letter(
    cover_letter: str = Form(...),
    job_title: str = Form(""),
    company: str = Form(""),
    job_description: str = Form(""),
    resume_text: str = Form("")
):
    """评估申请成功率
    
    检查维度（申请成功率视角）：
    - 关键词匹配率 (30%) - 理解职位要求的程度
    - 简历-职位关联度 (20%) - 核心！简历能力是否匹配职位
    - 长度、结构、ATS友好度、行动词、个性化 (各10%)
    """
    try:
        job = {
            "title": job_title,
            "company": company,
            "description": job_description
        }
        result = cover_letter_checker.check(cover_letter, job, resume_text)
        
        return {
            "success": True,
            **result
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/cover-letter/check-with-job")
async def check_cover_letter_with_job(
    cover_letter: str = Form(...),
    saved_jobs: str = Form("[]")  # JSON字符串
):
    """根据已保存的职位测试求职信质量
    
    saved_jobs 为包含职位信息的JSON数组
    """
    try:
        jobs_list = json.loads(saved_jobs)
        results = []
        
        for job in jobs_list:
            result = cover_letter_checker.check(cover_letter, job)
            results.append({
                "job_title": job.get("title", ""),
                "company": job.get("company", ""),
                **result
            })
        
        # 按质量分数排序
        results.sort(key=lambda x: x.get("overall_score", 0), reverse=True)
        
        return {
            "success": True,
            "total_jobs": len(jobs_list),
            "results": results
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# === 企业端龙虾接口 API ===

class LobsterCompanyAPI:
    """
    企业端龙虾接口
    
    龙虾(Lobster)是欧洲企业对B2B SaaS平台的昵称，
    这里指为企业HR/招聘团队提供的接口服务：
    - 接收投递的职位申请
    - 简历自动筛选
    - 候选人匹配排序
    - 申请状态管理
    """

    def __init__(self):
        # 企业注册表（内存存储，可扩展到数据库）
        self._companies: Dict[str, Dict] = {}
        self._job_postings: Dict[str, List[Dict]] = {}
        self._applications: Dict[str, List[Dict]] = {}

    def register_company(self, company_id: str, company_name: str, 
                         api_key: str, industry: str = "") -> Dict:
        """注册企业"""
        self._companies[company_id] = {
            "company_id": company_id,
            "name": company_name,
            "api_key": api_key,
            "industry": industry,
            "registered_at": "2026-04-01",
            "status": "active"
        }
        self._job_postings[company_id] = []
        self._applications[company_id] = []
        return self._companies[company_id]

    def post_job(self, company_id: str, job: Dict) -> Dict:
        """发布职位"""
        job_id = f"{company_id}_{len(self._job_postings.get(company_id, [])) + 1}"
        job_entry = {
            "job_id": job_id,
            "company_id": company_id,
            "title": job.get("title", ""),
            "description": job.get("description", ""),
            "requirements": job.get("requirements", ""),
            "location": job.get("location", ""),
            "salary_range": job.get("salary_range", ""),
            "status": "active",
            "posted_at": "2026-04-01",
            "applications_count": 0
        }
        if company_id not in self._job_postings:
            self._job_postings[company_id] = []
        self._job_postings[company_id].append(job_entry)
        return job_entry

    def submit_application(self, company_id: str, job_id: str, 
                           applicant: Dict) -> Dict:
        """接收职位申请"""
        app_id = f"APP_{company_id}_{len(self._applications.get(company_id, [])) + 1}"
        application = {
            "application_id": app_id,
            "job_id": job_id,
            "company_id": company_id,
            "applicant_name": applicant.get("name", ""),
            "applicant_email": applicant.get("email", ""),
            "resume_text": applicant.get("resume_text", ""),
            "cover_letter": applicant.get("cover_letter", ""),
            "status": "new",  # new, screening, interview, offer, rejected
            "submitted_at": "2026-04-01",
            "match_score": 0
        }
        
        # 自动计算匹配分
        if applicant.get("resume_text") and job_id:
            job = self._find_job(company_id, job_id)
            if job:
                match_res = job_matcher.match(applicant["resume_text"], job)
                application["match_score"] = match_res.overall_score
        
        if company_id not in self._applications:
            self._applications[company_id] = []
        self._applications[company_id].append(application)
        
        # 更新职位申请计数
        for j in self._job_postings.get(company_id, []):
            if j["job_id"] == job_id:
                j["applications_count"] += 1
                break
        
        return application

    def _find_job(self, company_id: str, job_id: str) -> Optional[Dict]:
        """查找职位"""
        for job in self._job_postings.get(company_id, []):
            if job["job_id"] == job_id:
                return job
        return None

    def get_candidates(self, company_id: str, job_id: str, 
                       min_score: float = 0) -> List[Dict]:
        """获取候选人列表（按匹配度排序）"""
        apps = self._applications.get(company_id, [])
        candidates = [a for a in apps if a["job_id"] == job_id]
        candidates.sort(key=lambda x: x.get("match_score", 0), reverse=True)
        return [c for c in candidates if c["match_score"] >= min_score]

    def update_application_status(self, company_id: str, application_id: str,
                                   new_status: str) -> Dict:
        """更新申请状态"""
        for app in self._applications.get(company_id, []):
            if app["application_id"] == application_id:
                app["status"] = new_status
                return app
        return {}


# 初始化龙虾接口
lobster_api = LobsterCompanyAPI()

# 预注册演示企业
DEMO_COMPANY_ID = "demo_company_001"
lobster_api.register_company(
    company_id=DEMO_COMPANY_ID,
    company_name="Nordic ERP Solutions A/S",
    api_key="lobster_demo_key_2026",
    industry="IT Consulting"
)


@app.post("/lobster/register-company")
async def lobster_register_company(
    company_id: str = Form(...),
    company_name: str = Form(...),
    api_key: str = Form(...),
    industry: str = Form("")
):
    """注册企业（龙虾接口）"""
    try:
        company = lobster_api.register_company(company_id, company_name, api_key, industry)
        return {"success": True, "company": company}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/lobster/post-job")
async def lobster_post_job(
    company_id: str = Form(...),
    job: str = Form(...)  # JSON字符串
):
    """发布职位（龙虾接口）"""
    try:
        job_dict = json.loads(job)
        job_entry = lobster_api.post_job(company_id, job_dict)
        return {"success": True, "job": job_entry}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/lobster/submit-application")
async def lobster_submit_application(
    company_id: str = Form(...),
    job_id: str = Form(...),
    applicant: str = Form(...)  # JSON字符串
):
    """提交职位申请（龙虾接口）
    
    自动：
    1. 接收简历和求职信
    2. 计算与职位的匹配分
    3. 按匹配度排序候选人
    """
    try:
        applicant_dict = json.loads(applicant)
        application = lobster_api.submit_application(company_id, job_id, applicant_dict)
        return {
            "success": True,
            "application": application,
            "message": f"申请已提交！匹配度评分：{application['match_score']}%"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/lobster/candidates/{company_id}/{job_id}")
async def lobster_get_candidates(
    company_id: str,
    job_id: str,
    min_score: float = 0
):
    """获取职位候选人列表（按匹配度排序）"""
    try:
        candidates = lobster_api.get_candidates(company_id, job_id, min_score)
        return {
            "success": True,
            "job_id": job_id,
            "total": len(candidates),
            "candidates": candidates
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/lobster/jobs/{company_id}")
async def lobster_get_jobs(company_id: str):
    """获取企业的所有职位"""
    jobs = lobster_api._job_postings.get(company_id, [])
    return {"success": True, "jobs": jobs}


# === Stripe 付费系统 API ===

import stripe
from datetime import datetime

# Stripe 配置
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY")
STRIPE_PUBLISHABLE_KEY = os.getenv("STRIPE_PUBLISHABLE_KEY", "")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET", "")

stripe_client = None
if STRIPE_SECRET_KEY:
    try:
        stripe.api_key = STRIPE_SECRET_KEY
        stripe_client = stripe
        print(f"✅ Stripe 已连接")
    except Exception as e:
        print(f"⚠️ Stripe 初始化失败: {e}")

# 订阅套餐配置 - Pro 产品（2026-04-11 更新价格）
SUBSCRIPTION_PLANS = {
    "monthly": {
        "name": "Pro 月卡",
        "price_id": "price_1TKvh5ETBa7HTDFvGm4okWr8",  # 月付
        "price": 59,  # DKK / 月
        "price_cny": 59,  # RMB / 月
        "currency": "dkk",
        "currency_cny": "cny",
        "interval": "month",
        "features": [
            "✅ AI智能简历分析",
            "✅ 无限职位搜索",
            "✅ 无限求职信生成",
            "✅ ATS关键词优化",
            "✅ 简历模版",
            "✅ 申请追踪"
        ],
        "tagline": "💡 灵活订阅，随时取消"
    },
    "quarterly": {
        "name": "Pro 季卡",
        "price_id": "price_1TKvh5ETBa7HTDFvTNrSqL79",  # 季付
        "price": 159,  # DKK / 季
        "price_cny": 159,  # RMB / 季
        "currency": "dkk",
        "currency_cny": "cny",
        "interval": "quarter",
        "features": [
            "✅ AI智能简历分析",
            "✅ 无限职位搜索",
            "✅ 无限求职信生成",
            "✅ ATS关键词优化",
            "✅ 简历模版",
            "✅ 申请追踪",
            "✅ 优先客户支持"
        ],
        "tagline": "💡 比月付省 ¥18"
    },
    "yearly": {
        "name": "Pro 年卡",
        "price_id": "price_1TKkGmETBa7HTDFva5uzb9wo",  # 年付
        "price": 499,  # DKK / 年
        "price_cny": 499,  # RMB / 年
        "currency": "dkk",
        "currency_cny": "cny",
        "interval": "year",
        "features": [
            "✅ AI智能简历分析",
            "✅ 无限职位搜索",
            "✅ 无限求职信生成",
            "✅ ATS关键词优化",
            "✅ 简历模版",
            "✅ 申请追踪",
            "✅ 优先客户支持",
            "✅ 简历与职位匹配度评估"
        ],
        "tagline": "💡 Best Value ⭐ 节省 ¥209"
    }
}

# 用户订阅状态（内存存储，生产环境应使用数据库）
user_subscriptions: Dict[str, Dict] = {}


class StripeManager:
    """Stripe 订阅管理器"""

    def __init__(self):
        self.plans = SUBSCRIPTION_PLANS

    def is_china_ip(self, client_ip: str) -> bool:
        """简单判断是否为中国大陆IP（生产环境建议用专业的IP库）"""
        if not client_ip:
            return False
        # 常见中国IP段（简化版，生产环境建议用 ip2location 或 ipapi）
        china_ip_prefixes = ['36.', '42.', '58.', '59.', '60.', '61.', '101.', '103.', '106.', '110.', '111.', '112.', '113.', '114.', '115.', '116.', '117.', '118.', '119.', '120.', '121.', '122.', '123.', '124.', '125.', '140.', '175.', '180.', '182.', '183.', '202.', '203.', '210.', '211.', '218.', '220.', '221.', '222.', '223.']
        return any(client_ip.startswith(prefix) for prefix in china_ip_prefixes)

    def get_plans(self, client_ip: str = None) -> List[Dict]:
        """获取所有订阅套餐（根据IP分流显示不同货币）"""
        is_china = self.is_china_ip(client_ip) if client_ip else False
        result = []
        for plan_id, plan in self.plans.items():
            # 根据IP选择货币和价格
            if is_china:
                price = plan.get("price_cny", plan["price"])
                currency = plan.get("currency_cny", "cny")
            else:
                price = plan["price"]
                currency = plan["currency"]

            result.append({
                "id": plan_id,
                "name": plan["name"],
                "price": price,
                "currency": currency,
                "interval": plan["interval"],
                "features": plan["features"],
                "savings": self._calculate_savings(plan_id, is_china)
            })
        return result

    def _calculate_savings(self, plan_id: str, is_china: bool = False) -> str:
        """计算节省金额"""
        currency = "¥" if is_china else "DKK "
        monthly_total = 59 * 12  # 基于月付价格 59 DKK
        plan = self.plans.get(plan_id, {})

        if is_china:
            yearly_price = plan.get("price_cny", plan["price"])
        else:
            yearly_price = plan["price"]

        if plan_id == "yearly" and yearly_price > 0:
            saved = monthly_total - yearly_price
            return f"省 {currency}{saved}/年"
        elif plan_id == "quarterly" and yearly_price > 0:
            saved = monthly_total - yearly_price * 4
            return f"省 {currency}{saved}/年"
        return ""

    def create_checkout_session(self, plan_id: str, user_id: str,
                                 success_url: str = "/?payment=success",
                                 cancel_url: str = "/?payment=cancelled",
                                 client_ip: str = None) -> Dict:
        """创建 Stripe Checkout Session（支持IP分流）"""
        if not stripe_client:
            # Demo 模式：模拟支付流程
            return self._create_demo_checkout(plan_id, user_id)

        plan = self.plans.get(plan_id)
        if not plan:
            return {"success": False, "error": "Invalid plan"}

        # 根据IP选择货币和价格
        is_china = self.is_china_ip(client_ip) if client_ip else False
        if is_china:
            currency = plan.get("currency_cny", "cny")
            price = plan.get("price_cny", plan["price"])
        else:
            currency = plan["currency"]
            price = plan["price"]

        try:
            session = stripe_client.checkout.Session.create(
                payment_method_types=['card'],
                line_items=[{
                    'price_data': {
                        'currency': currency,
                        'product_data': {
                            'name': f"JobMatchAI {plan['name']}",
                            'description': f"Access to all JobMatchAI features - {plan['interval']}"
                        },
                        'unit_amount': price * 100,  # Stripe 使用分
                    },
                    'quantity': 1,
                }],
                mode='payment',
                success_url=success_url,
                cancel_url=cancel_url,
                metadata={
                    "user_id": user_id,
                    "plan_id": plan_id,
                    "is_china": str(is_china)
                },
                customer_email=None
            )
            return {
                "success": True,
                "session_id": session.id,
                "checkout_url": session.url,
                "currency": currency,
                "price": price
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def _create_demo_checkout(self, plan_id: str, user_id: str) -> Dict:
        """创建演示用 checkout（无真实 Stripe key）"""
        plan = self.plans.get(plan_id, {})
        demo_session_id = f"demo_session_{user_id}_{plan_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
        return {
            "success": True,
            "demo_mode": True,
            "session_id": demo_session_id,
            "checkout_url": f"/#/payment/demo?plan={plan_id}&session={demo_session_id}",
            "message": "⚠️ Demo模式：配置 STRIPE_SECRET_KEY 后启用真实支付"
        }

    def verify_subscription(self, user_id: str) -> Dict:
        """验证用户订阅状态"""
        sub = user_subscriptions.get(user_id, {})
        plan_id = sub.get("plan_id", "")
        plan = self.plans.get(plan_id, {})

        # 检查是否过期
        expires_at = sub.get("expires_at", "")
        is_active = False
        if expires_at:
            try:
                exp_date = datetime.fromisoformat(expires_at)
                is_active = exp_date > datetime.now()
            except:
                is_active = False

        return {
            "user_id": user_id,
            "is_active": is_active,
            "plan_id": plan_id,
            "plan_name": plan.get("name", "免费版"),
            "expires_at": expires_at,
            "features": plan.get("features", []) if is_active else []
        }

    def activate_subscription(self, user_id: str, plan_id: str) -> Dict:
        """激活订阅（支付成功后调用）"""
        plan = self.plans.get(plan_id)
        if not plan:
            return {"success": False, "error": "Invalid plan"}

        # 计算到期时间
        now = datetime.now()
        interval = plan.get("interval", "month")
        if interval == "month":
            from datetime import timedelta
            expires = now + timedelta(days=30)
        elif interval == "quarter":
            from datetime import timedelta
            expires = now + timedelta(days=90)
        else:
            from datetime import timedelta
            expires = now + timedelta(days=365)

        user_subscriptions[user_id] = {
            "plan_id": plan_id,
            "plan_name": plan["name"],
            "activated_at": now.isoformat(),
            "expires_at": expires.isoformat(),
            "status": "active"
        }

        return {
            "success": True,
            "plan_id": plan_id,
            "plan_name": plan["name"],
            "activated_at": now.isoformat(),
            "expires_at": expires.isoformat()
        }

    def cancel_subscription(self, user_id: str) -> Dict:
        """取消订阅"""
        if user_id in user_subscriptions:
            del user_subscriptions[user_id]
        return {"success": True, "message": "订阅已取消"}

    def get_billing_portal_session(self, user_id: str, return_url: str = "/") -> Dict:
        """获取 Stripe Billing Portal Session"""
        if not stripe_client:
            return {
                "success": False,
                "error": "Stripe not configured",
                "demo_mode": True
            }

        try:
            session = stripe_client.billing_portal.Session.create(
                customer="cus_demo_" + user_id,
                return_url=return_url
            )
            return {
                "success": True,
                "portal_url": session.url
            }
        except Exception as e:
            return {"success": False, "error": str(e)}


# 初始化 Stripe 管理器
stripe_manager = StripeManager()


def get_client_ip(request: Request) -> str:
    """获取客户端真实IP"""
    # 优先从 X-Forwarded-For 获取（如果通过代理）
    forwarded = request.headers.get("x-forwarded-for")
    if forwarded:
        return forwarded.split(",")[0].strip()
    # 然后尝试 X-Real-IP
    real_ip = request.headers.get("x-real-ip")
    if real_ip:
        return real_ip
    # 最后用客户端 IP
    if request.client:
        return request.client.host
    return ""


@app.get("/payment/plans")
def get_payment_plans(request: Request):
    """获取所有订阅套餐（根据IP显示不同货币）"""
    client_ip = get_client_ip(request)
    is_china = stripe_manager.is_china_ip(client_ip)
    return {
        "success": True,
        "plans": stripe_manager.get_plans(client_ip),
        "currency": "cny" if is_china else "dkk",
        "region": "china" if is_china else "international",
        "demo_note": "配置 STRIPE_SECRET_KEY 环境变量启用真实支付" if not STRIPE_SECRET_KEY else ""
    }


@app.get("/payment/subscription/{user_id}")
def get_subscription(user_id: str):
    """获取用户订阅状态"""
    return {
        "success": True,
        "subscription": stripe_manager.verify_subscription(user_id)
    }


@app.post("/payment/create-checkout")
async def create_payment_checkout(request: Request,
    plan_id: str = Form(...),
    user_id: str = Form(...)
):
    """创建支付会话（根据IP选择货币）"""
    try:
        client_ip = get_client_ip(request)
        result = stripe_manager.create_checkout_session(plan_id, user_id, client_ip=client_ip)
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/payment/activate-demo")
async def activate_demo_subscription(
    plan_id: str = Form(...),
    user_id: str = Form(...)
):
    """演示模式：直接激活订阅（无需真实支付）"""
    try:
        result = stripe_manager.activate_subscription(user_id, plan_id)
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/payment/cancel")
async def cancel_subscription(user_id: str = Form(...)):
    """取消订阅"""
    try:
        result = stripe_manager.cancel_subscription(user_id)
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/payment/webhook")
async def stripe_webhook(request: Request):
    """处理 Stripe Webhook"""
    if not stripe_client:
        return JSONResponse({"received": True}, status_code=200)

    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")

    try:
        event = stripe_client.webhook.construct_event(
            payload, sig_header, STRIPE_WEBHOOK_SECRET
        )

        if event["type"] == "checkout.session.completed":
            session = event["data"]["object"]
            user_id = session.get("metadata", {}).get("user_id")
            plan_id = session.get("metadata", {}).get("plan_id")
            if user_id and plan_id:
                stripe_manager.activate_subscription(user_id, plan_id)
                print(f"✅ 订阅激活: {user_id} -> {plan_id}")

        elif event["type"] == "customer.subscription.deleted":
            # 处理取消订阅
            pass

        return JSONResponse({"received": True}, status_code=200)
    except Exception as e:
        print(f"⚠️ Webhook 错误: {e}")
        return JSONResponse({"error": str(e)}, status_code=400)


# === 职位申请追踪 API ===

class ApplicationTracker:
    """职位申请追踪器"""

    def __init__(self):
        # 申请记录存储
        self._applications: Dict[str, List[Dict]] = {}

    def add_application(self, user_id: str, application: Dict) -> Dict:
        """添加一条申请记录"""
        if user_id not in self._applications:
            self._applications[user_id] = []
        
        app_entry = {
            "id": f"app_{user_id}_{len(self._applications[user_id]) + 1}",
            "user_id": user_id,
            "job_title": application.get("job_title", ""),
            "company": application.get("company", ""),
            "company_website": application.get("company_website", ""),
            "job_url": application.get("job_url", ""),
            "salary_range": application.get("salary_range", ""),
            "location": application.get("location", ""),
            "status": application.get("status", "applied"),
            "applied_date": application.get("applied_date", "2026-04-01"),
            "deadline": application.get("deadline", ""),
            "contact_name": application.get("contact_name", ""),
            "contact_email": application.get("contact_email", ""),
            "notes": application.get("notes", ""),
            "interview_date": "",
            "interview_notes": "",
            "offer_received": False,
            "rejected": False,
            "follow_up_date": "",
            "priority": application.get("priority", "normal"),  # high, normal, low
            "match_score": application.get("match_score", 0),
            "updated_at": "2026-04-01"
        }
        
        self._applications[user_id].append(app_entry)
        return app_entry

    def get_applications(self, user_id: str, status: str = "all") -> List[Dict]:
        """获取用户的申请列表"""
        apps = self._applications.get(user_id, [])
        if status != "all":
            apps = [a for a in apps if a["status"] == status]
        # 按更新时间和优先级排序
        priority_order = {"high": 0, "normal": 1, "low": 2}
        apps.sort(key=lambda x: (priority_order.get(x["priority"], 1), -x.get("match_score", 0)))
        return apps

    def update_status(self, user_id: str, application_id: str, 
                      new_status: str, notes: str = "") -> Dict:
        """更新申请状态"""
        for app in self._applications.get(user_id, []):
            if app["id"] == application_id:
                old_status = app["status"]
                app["status"] = new_status
                app["updated_at"] = "2026-04-01"
                if notes:
                    app["notes"] = notes
                if new_status == "rejected":
                    app["rejected"] = True
                if new_status == "offer":
                    app["offer_received"] = True
                return app
        return {}

    def get_statistics(self, user_id: str) -> Dict:
        """获取申请统计"""
        apps = self._applications.get(user_id, [])
        total = len(apps)
        
        status_counts = {}
        for app in apps:
            s = app.get("status", "unknown")
            status_counts[s] = status_counts.get(s, 0) + 1
        
        response_rate = 0
        if total > 0:
            responded = sum(1 for a in apps if a["status"] not in ["applied", "new"])
            response_rate = round(responded / total * 100, 1)
        
        # 本周申请数
        import datetime
        today = datetime.date.today()
        week_apps = [a for a in apps if a.get("applied_date", "") 
                     >= str(today - datetime.timedelta(days=7))]
        
        return {
            "total_applications": total,
            "status_breakdown": status_counts,
            "response_rate": response_rate,
            "this_week": len(week_apps),
            "pending_response": status_counts.get("applied", 0) + status_counts.get("new", 0),
            "interviews": status_counts.get("interview", 0),
            "offers": status_counts.get("offer", 0),
            "rejections": status_counts.get("rejected", 0)
        }


# 初始化申请追踪器
app_tracker = ApplicationTracker()


@app.post("/tracker/add")
async def add_application(
    user_id: str = Form(...),
    application: str = Form(...)  # JSON字符串
):
    """添加职位申请记录"""
    try:
        app_dict = json.loads(application)
        entry = app_tracker.add_application(user_id, app_dict)
        return {"success": True, "application": entry}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/tracker/list/{user_id}")
async def list_applications(
    user_id: str,
    status: str = "all"
):
    """获取申请列表"""
    apps = app_tracker.get_applications(user_id, status)
    return {
        "success": True,
        "count": len(apps),
        "applications": apps
    }


@app.post("/tracker/update-status")
async def update_application_status(
    user_id: str = Form(...),
    application_id: str = Form(...),
    new_status: str = Form(...),
    notes: str = Form("")
):
    """更新申请状态
    
    状态选项：
    - new: 新申请
    - applied: 已投递
    - screening: 筛选中
    - interview: 面试中
    - offer: 已拿到offer
    - rejected: 被拒绝
    - withdrawn: 自己撤回
    - accepted: 已接受
    """
    try:
        app = app_tracker.update_status(user_id, application_id, new_status, notes)
        if app:
            return {"success": True, "application": app}
        return {"success": False, "error": "Application not found"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/tracker/add-interview")
async def add_interview(
    user_id: str = Form(...),
    application_id: str = Form(...),
    interview_date: str = Form(...),
    interview_notes: str = Form("")
):
    """添加面试信息"""
    for app in app_tracker._applications.get(user_id, []):
        if app["id"] == application_id:
            app["interview_date"] = interview_date
            app["interview_notes"] = interview_notes
            app["status"] = "interview"
            app["updated_at"] = "2026-04-01"
            return {"success": True, "application": app}
    return {"success": False, "error": "Application not found"}


@app.get("/tracker/stats/{user_id}")
async def get_tracker_stats(user_id: str):
    """获取申请统计"""
    stats = app_tracker.get_statistics(user_id)
    return {"success": True, "stats": stats}


@app.get("/tracker/upcoming/{user_id}")
async def get_upcoming_activities(user_id: str):
    """获取即将到来的活动（面试/待跟进）"""
    apps = app_tracker._applications.get(user_id, [])
    upcoming = []
    
    for app in apps:
        if app.get("interview_date"):
            upcoming.append({
                "type": "interview",
                "application_id": app["id"],
                "job_title": app["job_title"],
                "company": app["company"],
                "date": app["interview_date"],
                "notes": app.get("interview_notes", "")
            })
        if app.get("follow_up_date"):
            upcoming.append({
                "type": "follow_up",
                "application_id": app["id"],
                "job_title": app["job_title"],
                "company": app["company"],
                "date": app["follow_up_date"],
                "notes": ""
            })
    
    upcoming.sort(key=lambda x: x.get("date", ""))
    return {"success": True, "upcoming": upcoming[:10]}


from early_bird import EarlyBirdCollector

# 初始化早期测试收集器
early_bird = EarlyBirdCollector()


# === 早期测试数据收集 API（Beta测试用） ===

# === LinkedIn 职位导入 ===

@app.post("/parse-job-url")
async def parse_job_url_endpoint(url: str = Form(...)):
    """解析 LinkedIn 职位 URL，提取职位信息
    
    用法：前端用户粘贴 LinkedIn 职位链接 → 后端抓取解析 → 返回结构化数据
    支持多种 LinkedIn 职位 URL 格式
    """
    try:
        result = linkedin_importer.parse_linkedin_url(url)
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/fetch-job-from-url")
async def fetch_job_from_url_endpoint(url: str = Form(...)):
    """从URL抓取职位信息（前端调用的端点）
    
    支持 LinkedIn、Jobindex、Jobnet 等职位网站
    返回格式与前端期望的一致
    
    策略：先用 LinkedIn API 抓取 -> 不完整则用 AI 提取 -> 再不完整则提示手动输入
    """
    try:
        result = linkedin_importer.parse_linkedin_url(url)
        
        if result.get("success") and result.get("job"):
            job = result["job"]
            title = job.get("title", "")
            company = job.get("company", "")
            description = job.get("description", "")
            raw_text = job.get("raw_text", "") or description
            
            # 如果 LinkedIn 抓取的数据不完整（缺少职位名或公司），用 AI 补充
            if (not title or not company) and raw_text and AI_AVAILABLE:
                print(f"LinkedIn 抓取数据不完整，尝试 AI 补充提取 (title='{title}', company='{company}')")
                try:
                    # 使用 AI 从 raw_text 中提取完整职位信息
                    extracted = extract_job_from_text(raw_text[:5000])
                    if extracted.get("success"):
                        title = title or extracted.get("title", "")
                        company = company or extracted.get("company", "")
                        if not description and extracted.get("description"):
                            description = extracted.get("description", "")
                        print(f"AI 补充提取成功: title='{title}', company='{company}'")
                except Exception as ai_err:
                    print(f"AI 补充提取失败: {ai_err}")
                    # 继续使用 LinkedIn 返回的部分数据
            
            return {
                "success": True,
                "job": {
                    "success": True,
                    "title": title or "职位名称（请手动填写）",
                    "company": company or "公司名称（请手动填写）",
                    "location": job.get("location", ""),
                    "description": description,
                    "url": url,
                    "source": job.get("source", "LinkedIn"),
                    "raw_text": raw_text
                }
            }
        else:
            error_msg = result.get("error", "无法从此链接抓取职位信息")
            return {
                "success": False,
                "job": {
                    "success": False,
                    "error": error_msg
                },
                "message": f"⚠️ {error_msg}。请尝试在职位页面复制全部内容，粘贴到下方。"
            }
    except Exception as e:
        return {
            "success": False,
            "job": {
                "success": False,
                "error": str(e)
            },
            "message": "⚠️ 抓取出错，请稍后重试或手动粘贴职位信息"
        }


@app.post("/import-linkedin")
async def import_linkedin_endpoint(
    url: str = Form(""),
    job_text: str = Form(""),
):
    """导入 LinkedIn 职位 — 支持两种方式：
    
    1. URL 方式：提供 LinkedIn 职位链接，自动抓取解析
    2. 粘贴方式：用户直接粘贴职位描述文本
    
    至少需要提供 url 或 job_text 之一
    """
    try:
        # 方式1: URL 自动解析
        if url:
            result = linkedin_importer.parse_linkedin_url(url)
            if result.get("success") and result.get("job"):
                job = result["job"]
                return {
                    "success": True,
                    "import_method": "url",
                    "job": job,
                    "message": "✅ 职位信息已成功解析"
                }
            # URL 解析失败，提示使用粘贴方式
            return {
                "success": False,
                "import_method": "url_failed",
                "error": result.get("error", "无法解析此链接"),
                "message": "⚠️ 无法自动解析此 LinkedIn 职位。请在 LinkedIn 职位页面复制全部内容，粘贴到下方。",
                "fallback": True
            }
        
        # 方式2: 手动粘贴解析
        elif job_text:
            result = linkedin_importer.parse_manual_input(job_text)
            if result.get("success") and result.get("job"):
                job = result["job"]
                return {
                    "success": True,
                    "import_method": "paste",
                    "job": job,
                    "message": "✅ 职位信息已从粘贴内容中提取"
                }
            return {
                "success": False,
                "error": result.get("error", "无法解析粘贴内容")
            }
        
        else:
            return {"success": False, "error": "请提供 LinkedIn 职位链接或职位描述文本"}
            
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/linkedin/generate-cover-letter")
async def linkedin_cover_letter_endpoint(
    resume_text: str = Form(...),
    job_title: str = Form(...),
    company: str = Form(""),
    job_description: str = Form(""),
    job_location: str = Form(""),
    language: str = Form("auto"),
    resume_highlights: str = Form(""),  # JSON string of highlights from polish
):
    """从 LinkedIn 导入的职位生成三语求职信（升级版：支持简历亮点）
    
    language 参数：
    - "auto": 自动检测职位描述语言
    - "zh": 中文
    - "en": 英文
    - "da": 丹麦文
    """
    try:
        # 自动检测语言
        if language == "auto":
            lang = linkedin_detect_language(job_description) if job_description else detect_language(job_title)
        else:
            lang = language

        job = {
            "title": job_title,
            "company": company,
            "description": job_description,
            "location": job_location,
            "source": "LinkedIn",
            "url": ""
        }

        # 解析简历亮点
        highlights = []
        if resume_highlights:
            try:
                highlights = json.loads(resume_highlights)
            except:
                pass
        
        # 如果没有提供亮点，先调用精修获取
        if not highlights and AI_AVAILABLE:
            job_context = {
                "job_title": job_title,
                "job_description": job_description,
                "company": company
            }
            polish_result = generate_polish_suggestions(resume_text, lang, job_context)
            if isinstance(polish_result, dict):
                highlights = polish_result.get('resume_highlights', [])

        cover_letter = generate_cover_letter_with_ai(resume_text, job, lang, highlights, None)  # 结构化数据从上传时获取

        return {
            "success": True,
            "cover_letter": cover_letter,
            "language": lang,
            "language_name": {"zh": "中文", "en": "English", "da": "Dansk"}.get(lang, "English"),
            "highlights_used": len(highlights) > 0,
            "job": {
                "title": job_title,
                "company": company,
                "source": "LinkedIn"
            }
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/linkedin/validate-url")
def validate_linkedin_url(url: str):
    """快速验证 LinkedIn URL 是否有效，无需完整解析"""
    is_valid, result = linkedin_importer.validate_url(url)
    return {
        "is_valid": is_valid,
        "job_id": result if is_valid else None,
        "error": None if is_valid else result
    }


# === Beta 测试系统 ===

@app.post("/beta/submit")
async def submit_beta_form(
    email: str = Form(...),
    name: str = Form(""),
    phone: str = Form(""),
    country: str = Form(""),
    city: str = Form(""),
    resume_text: str = Form(""),
    resume_file_name: str = Form(""),
    job_links: str = Form("[]"),  # JSON字符串
    social_accounts: str = Form("[]"),  # JSON字符串
    notes: str = Form(""),
    source: str = Form("website")
):
    """早期测试表单提交

    用于收集用户数据：
    - 简历文本或文件
    - 已申请/关注的职位链接
    - 社交媒体账号
    - 联系方式

    所有数据用于本地AI处理，帮助优化算法。
    """
    try:
        job_links_list = json.loads(job_links) if job_links else []
        social_list = json.loads(social_accounts) if social_accounts else []

        result = early_bird.submit_form(
            email=email,
            name=name,
            phone=phone,
            country=country,
            city=city,
            resume_text=resume_text,
            resume_file_name=resume_file_name,
            job_links=job_links_list,
            social_accounts=social_list,
            notes=notes,
            source=source
        )

        return result
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/beta/submission/{submission_id}")
async def get_beta_submission(submission_id: str):
    """获取提交详情"""
    submission = early_bird.get_submission(submission_id)
    if submission:
        return {"success": True, "submission": submission}
    return {"success": False, "error": "Submission not found"}


@app.get("/beta/submissions")
async def list_beta_submissions(
    status: str = "all",
    source: str = "all",
    limit: int = 50
):
    """获取提交列表"""
    submissions = early_bird.list_submissions(status, source, limit)
    return {"success": True, "submissions": submissions}


@app.post("/beta/update-status")
async def update_beta_status(
    submission_id: str = Form(...),
    status: str = Form(...),
    notes: str = Form("")
):
    """更新处理状态"""
    success = early_bird.update_status(submission_id, status, notes)
    return {"success": success}


@app.get("/beta/statistics")
async def get_beta_statistics():
    """获取统计数据"""
    stats = early_bird.get_statistics()
    return {"success": True, "statistics": stats}


@app.post("/beta/process")
async def process_beta_submission(
    submission_id: str = Form(...),
    action: str = Form(...),
    result: str = Form("")
):
    """记录处理操作"""
    early_bird.log_processing(submission_id, action, result)
    early_bird.update_status(submission_id, "processed", result)
    return {"success": True}


# ============================================================
# 新增：用户认证、简历管理、申请追踪 API
# ============================================================

from fastapi import Header, HTTPException
from typing import Optional

# 导入认证模块
from auth import (
    register_user, login_user, get_current_user, 
    update_user_profile, change_password, extract_token_from_header
)

# 导入简历导出模块
from resume_exporter import export_user_resume, get_resume_for_export, DOCX_AVAILABLE

# 导入申请追踪模块
from application_tracker import application_service

# 导入数据库和简历模块
from database import (
    create_resume, get_user_resumes, get_primary_resume,
    update_resume, delete_resume, get_resume_by_id,
    create_cover_letter, get_user_cover_letters, update_cover_letter, delete_cover_letter
)


def get_current_user_id(authorization: Optional[str] = Header(None)) -> Optional[str]:
    """从请求头获取当前用户ID"""
    if not authorization:
        return None
    token = extract_token_from_header(authorization)
    if not token:
        return None
    from auth import get_user_from_token
    user = get_user_from_token(token)
    return user["user_id"] if user else None


# === 认证 API ===

@app.post("/auth/register")
async def register_endpoint(
    email: str = Form(...),
    password: str = Form(...),
    name: str = Form(""),
    phone: str = Form(""),
    country: str = Form(""),
    city: str = Form("")
):
    """用户注册"""
    result = register_user(email, password, name, phone, country, city)
    return result


@app.post("/auth/login")
async def login_endpoint(
    email: str = Form(...),
    password: str = Form(...)
):
    """用户登录"""
    result = login_user(email, password)
    return result


@app.get("/auth/me")
async def get_me_endpoint(authorization: Optional[str] = Header(None)):
    """获取当前用户信息"""
    if not authorization:
        return {"success": False, "error": "Authorization header required"}
    
    result = get_current_user(authorization)
    return result


@app.post("/auth/update-profile")
async def update_profile_endpoint(
    authorization: Optional[str] = Header(None),
    name: str = Form(""),
    phone: str = Form(""),
    country: str = Form(""),
    city: str = Form(""),
    preferred_language: str = Form("")
):
    """更新用户资料"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    updates = {}
    if name: updates["name"] = name
    if phone: updates["phone"] = phone
    if country: updates["country"] = country
    if city: updates["city"] = city
    if preferred_language: updates["preferred_language"] = preferred_language
    
    return update_user_profile(user_id, **updates)


@app.post("/auth/change-password")
async def change_password_endpoint(
    authorization: Optional[str] = Header(None),
    old_password: str = Form(...),
    new_password: str = Form(...)
):
    """修改密码"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    return change_password(user_id, old_password, new_password)


# === 快速注册（保存进度）API ===

@app.post("/session/register")
async def session_register_endpoint(
    session_id: str = Form(...),
    email: str = Form(...),
    password: str = Form(""),
    interested_position: str = Form("")
):
    """
    快速注册：将匿名会话绑定到用户账户
    
    1. 如果邮箱已存在，直接绑定会话并登录
    2. 如果邮箱不存在，创建新用户并绑定
    3. 转移会话中的简历、职位数据到用户账户
    """
    from database import (
        save_session_data, get_all_session_data, bind_session_to_user,
        get_user_by_email, create_user
    )
    from auth import hash_password
    import json
    
    try:
        # 检查邮箱是否已注册
        existing_user = get_user_by_email(email)
        
        if existing_user:
            # 已注册用户：直接绑定会话
            user_id = existing_user["user_id"]
            bind_result = bind_session_to_user(session_id, user_id)
            
            return {
                "success": True,
                "is_new_user": False,
                "user_id": user_id,
                "message": "已绑定到您的账户，之前的进度已保存"
            }
        else:
            # 新用户：创建账户并绑定
            user_id = f"user_{int(datetime.now().timestamp())}"
            password_hash = hash_password(password) if password else hash_password(f"temp_{session_id}")
            
            # 创建用户
            create_user(
                user_id=user_id,
                email=email,
                password_hash=password_hash,
                name="",
                phone="",
                country="",
                city=""
            )
            
            # 绑定会话并转移数据
            bind_session_to_user(session_id, user_id)
            
            return {
                "success": True,
                "is_new_user": True,
                "user_id": user_id,
                "message": "注册成功！您的简历和进度已保存"
            }
            
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/session/save-progress")
async def session_save_progress_endpoint(
    session_id: str = Form(...),
    data_type: str = Form(...),
    data_content: str = Form(...)
):
    """
    保存会话进度（临时存储）
    
    data_type: resume | job_description | cover_letter | analysis_result
    """
    from database import save_session_data
    
    try:
        save_session_data(session_id, data_type, data_content)
        return {"success": True, "message": "进度已保存"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/session/{session_id}/data")
async def session_get_data_endpoint(session_id: str):
    """获取会话保存的数据"""
    from database import get_all_session_data
    
    try:
        data = get_all_session_data(session_id)
        return {"success": True, "data": data}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/session/check-email")
async def check_email_exists_endpoint(email: str = Form(...)):
    """检查邮箱是否已注册"""
    from database import get_user_by_email
    
    user = get_user_by_email(email)
    return {
        "exists": user is not None,
        "user_id": user["user_id"] if user else None
    }


# === 简历管理 API ===

@app.post("/resumes/upload")
async def upload_resume_endpoint(
    authorization: Optional[str] = Header(None),
    resume_text: str = Form(...),
    title: str = Form(""),
    language: str = Form("en"),
    file_name: str = Form(""),
    file_type: str = Form(""),
    is_primary: bool = Form(False)
):
    """上传简历"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    try:
        resume = create_resume(
            user_id=user_id,
            content=resume_text,
            title=title or "My Resume",
            language=language,
            file_name=file_name,
            file_type=file_type,
            is_primary=is_primary
        )
        return {
            "success": True,
            "resume": {
                "resume_id": resume["resume_id"],
                "title": resume["title"],
                "language": resume["language"],
                "created_at": resume["created_at"]
            }
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/resumes")
async def list_resumes_endpoint(
    authorization: Optional[str] = Header(None)
):
    """获取用户简历列表"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    resumes = get_user_resumes(user_id)
    return {
        "success": True,
        "resumes": [
            {
                "resume_id": r["resume_id"],
                "title": r["title"],
                "language": r["language"],
                "is_primary": bool(r.get("is_primary", 0)),
                "ats_score": r.get("ats_score"),
                "created_at": r["created_at"]
            }
            for r in resumes
        ]
    }


@app.get("/resumes/primary")
async def get_primary_resume_endpoint(
    authorization: Optional[str] = Header(None)
):
    """获取主要简历"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    resume = get_primary_resume(user_id)
    if resume:
        return {
            "success": True,
            "resume": {
                "resume_id": resume["resume_id"],
                "title": resume["title"],
                "content": resume["content"],
                "language": resume["language"],
                "is_primary": True,
                "ats_score": resume.get("ats_score"),
                "created_at": resume["created_at"]
            }
        }
    return {"success": False, "error": "No resume found"}


@app.get("/resumes/{resume_id}")
async def get_resume_detail_endpoint(
    resume_id: str,
    authorization: Optional[str] = Header(None)
):
    """获取简历详情"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    resume = get_resume_by_id(resume_id)
    if resume and resume["user_id"] == user_id:
        return {
            "success": True,
            "resume": resume
        }
    return {"success": False, "error": "Resume not found"}


@app.post("/resumes/{resume_id}/set-primary")
async def set_primary_resume_endpoint(
    resume_id: str,
    authorization: Optional[str] = Header(None)
):
    """设为主要简历"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    resume = get_resume_by_id(resume_id)
    if resume and resume["user_id"] == user_id:
        update_resume(resume_id, is_primary=True)
        return {"success": True, "message": "Primary resume set"}
    return {"success": False, "error": "Resume not found"}


@app.post("/resumes/{resume_id}/delete")
async def delete_resume_endpoint(
    resume_id: str,
    authorization: Optional[str] = Header(None)
):
    """删除简历"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    resume = get_resume_by_id(resume_id)
    if resume and resume["user_id"] == user_id:
        delete_resume(resume_id)
        return {"success": True, "message": "Resume deleted"}
    return {"success": False, "error": "Resume not found"}


# === 简历导出 API ===

@app.get("/export-resume")
async def export_resume_endpoint(
    authorization: Optional[str] = Header(None),
    resume_id: str = "",
    language: str = "en"
):
    """导出简历为 Word 文档
    
    参数：
    - resume_id: 简历ID（为空则导出主要简历）
    - language: 导出语言 (en, zh, da, auto)
    """
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    if not DOCX_AVAILABLE:
        return {"success": False, "error": "Word export not available. Please install python-docx."}
    
    try:
        doc_bytes = export_user_resume(user_id, resume_id, language)
        if doc_bytes:
            return StreamingResponse(
                io.BytesIO(doc_bytes),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=resume_{language}.docx"
                }
            )
        return {"success": False, "error": "Resume not found"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/export-resume-pdf")
async def export_resume_pdf_endpoint(
    authorization: Optional[str] = Header(None),
    resume_id: str = Form(""),
    language: str = Form("en"),
    polished_content: str = Form(""),
    user_name: str = Form(""),
    user_email: str = Form(""),
    user_phone: str = Form(""),
    user_location: str = Form("")
):
    """导出简历为专业 PDF 文档（使用 reportlab，支持中/英/丹麦语）
    
    参数：
    - resume_id: 简历ID（为空则导出主要简历）
    - language: 导出语言 (en, zh, da)
    - polished_content: 精修后的简历内容（可选，会覆盖数据库内容）
    - user_name/user_email/user_phone/user_location: 用户信息（精修内容时需要）
    """
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    if not PDF_ENGINE_AVAILABLE:
        return {"success": False, "error": "PDF engine not available. Please check dependencies."}
    
    try:
        # 优先使用精修内容
        if polished_content:
            resume_data = {
                "name": user_name or "Resume",
                "content": polished_content,
                "email": user_email,
                "phone": user_phone,
                "location": user_location,
            }
        else:
            # 从数据库获取
            if resume_id:
                resume_data_db = get_resume_by_id(resume_id)
            else:
                resume_data_db = get_primary_resume(user_id)
            
            if not resume_data_db:
                return {"success": False, "error": "Resume not found"}
            
            resume_data = {
                "name": resume_data_db.get("title", "Resume"),
                "content": resume_data_db.get("content", ""),
                "email": resume_data_db.get("email", ""),
                "phone": resume_data_db.get("phone", ""),
                "location": resume_data_db.get("location", ""),
            }
        
        # 生成 PDF
        from pdf_engine import generate_resume_pdf
        pdf_bytes = generate_resume_pdf(resume_data, language)
        
        if not pdf_bytes:
            return {"success": False, "error": "PDF generation failed"}
        
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=resume_{language}.pdf"
            }
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}


# === 求职信管理 API ===

@app.post("/cover-letters/save")
async def save_cover_letter_endpoint(
    authorization: Optional[str] = Header(None),
    company: str = Form(...),
    content: str = Form(...),
    job_title: str = Form(""),
    application_id: str = Form(""),
    language: str = Form("en"),
    is_template: bool = Form(False),
    quality_score: float = Form(0)
):
    """保存求职信"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    try:
        cl = create_cover_letter(
            user_id=user_id,
            company=company,
            content=content,
            job_title=job_title,
            application_id=application_id,
            language=language,
            is_template=is_template,
            quality_score=quality_score
        )
        return {
            "success": True,
            "cover_letter": {
                "cover_letter_id": cl["cover_letter_id"],
                "company": cl["company"],
                "language": cl["language"],
                "created_at": cl["created_at"]
            }
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/cover-letters")
async def list_cover_letters_endpoint(
    authorization: Optional[str] = Header(None),
    application_id: str = ""
):
    """获取求职信列表"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    letters = get_user_cover_letters(user_id, application_id)
    return {
        "success": True,
        "cover_letters": [
            {
                "cover_letter_id": cl["cover_letter_id"],
                "company": cl["company"],
                "job_title": cl.get("job_title", ""),
                "language": cl["language"],
                "is_template": bool(cl.get("is_template", 0)),
                "quality_score": cl.get("quality_score", 0),
                "created_at": cl["created_at"]
            }
            for cl in letters
        ]
    }


@app.get("/cover-letters/templates")
async def list_templates_endpoint(
    authorization: Optional[str] = Header(None)
):
    """获取求职信模板列表"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    templates = get_templates(user_id)
    return {
        "success": True,
        "templates": [
            {
                "cover_letter_id": cl["cover_letter_id"],
                "company": cl["company"],
                "job_title": cl.get("job_title", ""),
                "created_at": cl["created_at"]
            }
            for cl in templates
        ]
    }


@app.get("/cover-letters/{cover_letter_id}")
async def get_cover_letter_detail_endpoint(
    cover_letter_id: str,
    authorization: Optional[str] = Header(None)
):
    """获取求职信详情"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    cl = get_cover_letter_by_id(cover_letter_id)
    if cl and cl["user_id"] == user_id:
        return {
            "success": True,
            "cover_letter": cl
        }
    return {"success": False, "error": "Cover letter not found"}


@app.post("/cover-letters/{cover_letter_id}/update")
async def update_cover_letter_endpoint(
    cover_letter_id: str,
    authorization: Optional[str] = Header(None),
    content: str = Form(""),
    language: str = Form(""),
    is_template: bool = Form(None)
):
    """更新求职信"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    cl = get_cover_letter_by_id(cover_letter_id)
    if not cl or cl["user_id"] != user_id:
        return {"success": False, "error": "Cover letter not found"}
    
    updates = {}
    if content: updates["content"] = content
    if language: updates["language"] = language
    if is_template is not None: updates["is_template"] = is_template
    
    updated = update_cover_letter(cover_letter_id, **updates)
    return {
        "success": True,
        "cover_letter": updated
    }


@app.post("/cover-letters/{cover_letter_id}/delete")
async def delete_cover_letter_endpoint(
    cover_letter_id: str,
    authorization: Optional[str] = Header(None)
):
    """删除求职信"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    success = delete_cover_letter(cover_letter_id, user_id)
    if success:
        return {"success": True, "message": "Cover letter deleted"}
    return {"success": False, "error": "Cover letter not found"}


# === 申请追踪 API ===

@app.post("/applications/add")
async def add_application_endpoint(
    authorization: Optional[str] = Header(None),
    job_title: str = Form(...),
    company: str = Form(...),
    job_id: str = Form(""),
    company_website: str = Form(""),
    job_url: str = Form(""),
    salary_range: str = Form(""),
    location: str = Form(""),
    status: str = Form("new"),
    priority: str = Form("normal"),
    applied_date: str = Form(""),
    deadline: str = Form(""),
    contact_name: str = Form(""),
    contact_email: str = Form(""),
    contact_phone: str = Form(""),
    notes: str = Form(""),
    match_score: float = Form(0),
    resume_id: str = Form(""),
    cover_letter_id: str = Form("")
):
    """添加申请记录"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    return application_service.add_application(
        user_id=user_id,
        job_title=job_title,
        company=company,
        job_id=job_id,
        company_website=company_website,
        job_url=job_url,
        salary_range=salary_range,
        location=location,
        status=status,
        priority=priority,
        applied_date=applied_date,
        deadline=deadline,
        contact_name=contact_name,
        contact_email=contact_email,
        contact_phone=contact_phone,
        notes=notes,
        match_score=match_score,
        resume_id=resume_id,
        cover_letter_id=cover_letter_id
    )


@app.get("/applications")
async def list_applications_endpoint(
    authorization: Optional[str] = Header(None),
    status: str = "all",
    page: int = 1,
    page_size: int = 20
):
    """获取申请列表"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    return application_service.list_applications(user_id, status, page, page_size)


@app.get("/applications/{application_id}")
async def get_application_endpoint(
    application_id: str,
    authorization: Optional[str] = Header(None)
):
    """获取单个申请详情"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    return application_service.get_application(application_id, user_id)


@app.post("/applications/{application_id}/update")
async def update_application_endpoint(
    application_id: str,
    authorization: Optional[str] = Header(None),
    job_title: str = Form(""),
    company: str = Form(""),
    company_website: str = Form(""),
    job_url: str = Form(""),
    salary_range: str = Form(""),
    location: str = Form(""),
    status: str = Form(""),
    priority: str = Form(""),
    applied_date: str = Form(""),
    deadline: str = Form(""),
    contact_name: str = Form(""),
    contact_email: str = Form(""),
    contact_phone: str = Form(""),
    notes: str = Form(""),
    follow_up_date: str = Form("")
):
    """更新申请记录"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    updates = {}
    if job_title: updates["job_title"] = job_title
    if company: updates["company"] = company
    if company_website: updates["company_website"] = company_website
    if job_url: updates["job_url"] = job_url
    if salary_range: updates["salary_range"] = salary_range
    if location: updates["location"] = location
    if status: updates["status"] = status
    if priority: updates["priority"] = priority
    if applied_date: updates["applied_date"] = applied_date
    if deadline: updates["deadline"] = deadline
    if contact_name: updates["contact_name"] = contact_name
    if contact_email: updates["contact_email"] = contact_email
    if contact_phone: updates["contact_phone"] = contact_phone
    if notes: updates["notes"] = notes
    if follow_up_date: updates["follow_up_date"] = follow_up_date
    
    return application_service.update_application(application_id, user_id, **updates)


@app.post("/applications/{application_id}/update-status")
async def update_status_endpoint(
    application_id: str,
    authorization: Optional[str] = Header(None),
    new_status: str = Form(...),
    notes: str = Form("")
):
    """更新申请状态"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    return application_service.update_status(application_id, user_id, new_status, notes)


@app.post("/applications/{application_id}/add-interview")
async def add_interview_endpoint(
    application_id: str,
    authorization: Optional[str] = Header(None),
    interview_date: str = Form(...),
    interview_notes: str = Form("")
):
    """添加面试信息"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    return application_service.add_interview(application_id, user_id, interview_date, interview_notes)


@app.post("/applications/{application_id}/delete")
async def delete_application_endpoint(
    application_id: str,
    authorization: Optional[str] = Header(None)
):
    """删除申请记录"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    return application_service.delete_application(application_id, user_id)


@app.get("/applications/stats")
async def get_application_stats_endpoint(
    authorization: Optional[str] = Header(None)
):
    """获取申请统计"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    return application_service.get_statistics(user_id)


@app.get("/applications/upcoming")
async def get_upcoming_activities_endpoint(
    authorization: Optional[str] = Header(None)
):
    """获取即将到来的活动"""
    user_id = get_current_user_id(authorization)
    if not user_id:
        return {"success": False, "error": "Unauthorized"}
    
    return application_service.get_upcoming_activities(user_id)


@app.get("/applications/status-options")
async def get_status_options_endpoint(
    language: str = "en"
):
    """获取状态选项列表"""
    return application_service.get_status_options(language)


# ========== 简历精修 API ==========

class PolishRequest(BaseModel):
    resume_text: str
    language: str = "auto"
    job_title: str = ""
    job_description: str = ""
    company: str = ""
    # 结构化简历数据（新功能，来自 parse_resume_structure）
    resume_structure: Optional[Dict] = None

class PolishApplyRequest(BaseModel):
    original_text: str
    suggestions: List[Dict]

class DiscoverSkillsRequest(BaseModel):
    document_text: str
    resume_text: str = ""
    document_type: str = "document"

def generate_polish_suggestions(text: str, lang: str = 'en', job_context: Dict = None, resume_structure: Dict = None) -> List[Dict]:
    """使用 AI 生成逐条精修建议（升级版：简历亮点 + 职位匹配能力分析 + 结构化数据）
    
    Args:
        text: 简历文本
        lang: 语言代码
        job_context: 职位上下文
        resume_structure: 结构化简历数据（来自 parse_resume_structure）
    """
    if not AI_AVAILABLE:
        return generate_polish_fallback(text, lang)

    # 构建简历结构化数据上下文
    structure_context = ""
    if resume_structure and resume_structure.get('ai_enhanced'):
        edu_list = resume_structure.get('education', [])
        exp_list = resume_structure.get('experience', [])
        skills_data = resume_structure.get('skills', {})
        
        edu_str = ", ".join([f"{e.get('degree', '未知')}{e.get('field', '')}({e.get('school', '')}, {e.get('year', 'N/A')})" for e in edu_list]) if edu_list else "未识别到教育经历"
        exp_str = "\n".join([f"- {e.get('company', '公司')}: {e.get('role', '职位')} ({e.get('start', '')} - {e.get('end', '')})" for e in exp_list[:5]]) if exp_list else "未识别到工作经历"
        tech_skills = ", ".join(skills_data.get('technical', [])[:15]) if skills_data.get('technical') else "未识别到技术技能"
        
        structure_context = f"""
【简历结构化数据（已识别）】
学历: {edu_str}
工作经历（最近5个）:
{exp_str}
技术技能: {tech_skills}
总工作经验: {resume_structure.get('total_experience_years', 'N/A')}年

请结合以上已识别的简历结构数据，给出更精准的精修建议。
"""
    
    # 升级版提示词模板（中文）
    prompt_zh = """你是一位拥有20年HR高管经验的顶级职业顾问。请深度分析这份简历，提供专业精修建议，并提炼简历亮点。
{structure_context}
【核心要求】
1. 检查简历中的**每一句话**，找出所有需要改进的地方
2. 分析简历与职位的匹配度，找出核心能力亮点
3. 提炼出3个"简历亮点"，供后续求职信使用

【必须改进的句子类型】
1. **弱动词句子**：包含"负责"、"参与"、"协助"、"做了"等被动词
2. **缺乏量化**：没有数字、数据、百分比的成就描述
3. **笼统描述**：可以更具体却没有具体化的内容
4. **重复内容**：与前后句子意思重复的表达

【北欧职场适配 - 重要！】
- 强调协作、平等、成果导向的表达方式
- 丹麦HR喜欢：具体数据、直接表达、团队贡献、持续学习
- 避免：过度自我推销、模糊描述、不具体的承诺

【输出格式 - 必须严格遵循】
返回JSON格式，包含两个部分：

```json
{{
  "suggestions": [
    {{
      "id": 1,
      "original": "简历原文（精确匹配）",
      "suggested": "精修版本（保持原意但更专业）",
      "reason": "为什么这样改（50-100字）",
      "type": "weak_verb/quantification/vague/redundant",
      "priority": "high/medium/low"
    }}
  ],
  "resume_highlights": [
    {{
      "title": "亮点标题（3-8个字）",
      "description": "亮点描述（供求职信用，50-100字）",
      "evidence": "简历中的证据（原文或精修后）"
    }}
  ],
  "job_match_score": "85%",
  "recommendation": "针对这个职位的整体建议（30-50字）"
}}
```

【示例 - Resume Highlights】
简历中有："主导ERP实施项目，3个月内完成200+用户上线，零事故"
亮点输出：
```json
{{
  "title": "ERP实施专家",
  "description": "3个月内完成大型ERP项目上线，200+用户无事故切换，展现快速交付和风险管理能力",
  "evidence": "主导ERP实施项目，3个月内完成200+用户上线"
}}
```

简历内容：
"""

    # 升级版提示词模板
    prompts_by_lang = {
        'zh': prompt_zh.format(structure_context=structure_context) + text[:4000],
        'en': """You are a top career consultant with 20 years of HR executive experience. Analyze this resume and provide improvement suggestions with resume highlights.

【Core Requirements】
1. Check EVERY sentence for improvement opportunities
2. Analyze job match and identify core competency highlights
3. Extract 3 "resume highlights" for use in cover letters

【Must Improve These】
1. **Weak Verbs**: "responsible for", "participated in", "assisted", "helped"
2. **Missing Quantification**: Achievements without numbers/percentages
3. **Vague Descriptions**: Content that could be more specific
4. **Redundant Content**: Repetitive expressions

【Nordic Workplace Fit - IMPORTANT!】
- Emphasize: collaboration, equality, result-oriented expression
- Danish HR likes: concrete data, direct communication, team contribution, continuous learning
- Avoid: over-selling, vague descriptions, empty promises

【Output Format - STRICT】
```json
{
  "suggestions": [
    {
      "id": 1,
      "original": "exact text from resume",
      "suggested": "polished version",
      "reason": "why change (50-100 chars)",
      "type": "weak_verb/quantification/vague/redundant",
      "priority": "high/medium/low"
    }
  ],
  "resume_highlights": [
    {
      "title": "Highlight Title (3-8 words)",
      "description": "Highlight description for cover letter (50-100 chars)",
      "evidence": "Evidence from resume"
    }
  ],
  "job_match_score": "85%",
  "recommendation": "Overall recommendation for this job (30-50 chars)"
}
```

Resume:
{text[:4000]}""",
        'da': """Du er en top karrierekonsulent med 20 års erfaring som HR-direktør. Analyser dette CV og giv forbedringsforslag med CV-højdepunkter.

【Kernekrav】
1. Gennemgå HVER sætning for forbedringer
2. Analyser jobmatch og identificer kernekompetence-højdepunkter
3. Udtræk 3 "CV-højdepunkter" til brug i ansøgninger

【Skal Forbedres】
1. **Svage Verbum**: "ansvarlig for", "deltog i", "assisterede"
2. **Manglende Kvantificering**: Præstationer uden tal/procenter
3. **Vage Beskrivelser**: Indhold der kunne være mere specifikt
4. **Gentaget Indhold**: Repetitive udtryk

【Nordisk Arbejdspladstilpasning】
- Fremhæv: samarbejde, lighed, resultatorientering
- Danske HR kan lide: konkrete data, direkte kommunikation, teambidrag

【Output-format】
```json
{
  "suggestions": [...],
  "resume_highlights": [...],
  "job_match_score": "85%",
  "recommendation": "..."
}
```

CV:
{text[:4000]}""",
    }

    try:
        # 根据语言生成 prompt（支持结构化数据上下文）
        base_prompts = {
            'zh': """你是一位拥有20年HR高管经验的顶级职业顾问。请深度分析这份简历，提供专业精修建议，并提炼简历亮点。

【核心要求】
1. 检查简历中的**每一句话**，找出所有需要改进的地方
2. 分析简历与职位的匹配度，找出核心能力亮点
3. 提炼出3个"简历亮点"，供后续求职信使用

【必须改进的句子类型】
1. **弱动词句子**：包含"负责"、"参与"、"协助"、"做了"等被动词
2. **缺乏量化**：没有数字、数据、百分比的成就描述
3. **笼统描述**：可以更具体却没有具体化的内容
4. **重复内容**：与前后句子意思重复的表达

【北欧职场适配 - 重要！】
- 强调协作、平等、成果导向的表达方式
- 丹麦HR喜欢：具体数据、直接表达、团队贡献、持续学习
- 避免：过度自我推销、模糊描述、不具体的承诺

【输出格式 - 必须严格遵循】
返回JSON格式，包含两个部分：

```json
{
  "suggestions": [
    {
      "id": 1,
      "original": "简历原文（精确匹配）",
      "suggested": "精修版本（保持原意但更专业）",
      "reason": "为什么这样改（50-100字）",
      "type": "weak_verb/quantification/vague/redundant",
      "priority": "high/medium/low"
    }
  ],
  "resume_highlights": [
    {
      "title": "亮点标题（3-8个字）",
      "description": "亮点描述（供求职信用，50-100字）",
      "evidence": "简历中的证据（原文或精修后）"
    }
  ],
  "job_match_score": "85%",
  "recommendation": "针对这个职位的整体建议（30-50字）"
}
```

【示例 - Resume Highlights】
简历中有："主导ERP实施项目，3个月内完成200+用户上线，零事故"
亮点输出：
```json
{
  "title": "ERP实施专家",
  "description": "3个月内完成大型ERP项目上线，200+用户无事故切换，展现快速交付和风险管理能力",
  "evidence": "主导ERP实施项目，3个月内完成200+用户上线"
}
```

简历内容：
""",
            'en': """You are a top career consultant with 20 years of HR executive experience. Analyze this resume and provide improvement suggestions with resume highlights.

【Core Requirements】
1. Check EVERY sentence for improvement opportunities
2. Analyze job match and identify core competency highlights
3. Extract 3 "resume highlights" for use in cover letters

【Must Improve These】
1. **Weak Verbs**: "responsible for", "participated in", "assisted", "helped"
2. **Missing Quantification**: Achievements without numbers/percentages
3. **Vague Descriptions**: Content that could be more specific
4. **Redundant Content**: Repetitive expressions

【Nordic Workplace Fit - IMPORTANT!】
- Emphasize: collaboration, equality, result-oriented expression
- Danish HR likes: concrete data, direct communication, team contribution, continuous learning
- Avoid: over-selling, vague descriptions, empty promises

【Output Format - STRICT】
```json
{
  "suggestions": [
    {
      "id": 1,
      "original": "exact text from resume",
      "suggested": "polished version",
      "reason": "why change (50-100 chars)",
      "type": "weak_verb/quantification/vague/redundant",
      "priority": "high/medium/low"
    }
  ],
  "resume_highlights": [
    {
      "title": "Highlight Title (3-8 words)",
      "description": "Highlight description for cover letter (50-100 chars)",
      "evidence": "Evidence from resume"
    }
  ],
  "job_match_score": "85%",
  "recommendation": "Overall recommendation for this job (30-50 chars)"
}
```

Resume:
""",
            'da': """Du er en top karrierekonsulent med 20 års erfaring som HR-direktør. Analyser dette CV og giv forbedringsforslag med CV-højdepunkter.

【Kernekrav】
1. Gennemgå HVER sætning for forbedringer
2. Analyser jobmatch og identificer kernekompetence-højdepunkter
3. Udtræk 3 "CV-højdepunkter" til brug i ansøgninger

【Skal Forbedres】
1. **Svage Verbum**: "ansvarlig for", "deltog i", "assisterede"
2. **Manglende Kvantificering**: Præstationer uden tal/procenter
3. **Vage Beskrivelser**: Indhold der kunne være mere specifikt
4. **Gentaget Indhold**: Repetitive udtryk

【Nordisk Arbejdspladstilpasning】
- Fremhæv: samarbejde, lighed, resultatorientering
- Danske HR kan lide: konkrete data, direkte kommunikation, teambidrag

【Output-format】
```json
{
  "suggestions": [...],
  "resume_highlights": [...],
  "job_match_score": "85%",
  "recommendation": "..."
}
```

CV:
"""
        }
        
        # 拼接：结构化上下文 + 基础prompt + 简历文本
        base = base_prompts.get(lang, base_prompts['en'])
        prompt = structure_context + base + text[:4000]
        
        # 如果有职位上下文，追加职位信息
        if job_context and job_context.get('job_title'):
            job_context_section = {
                'zh': f"""
【目标职位上下文 - 请重点关注】
职位名称: {job_context.get('job_title', '')}
公司: {job_context.get('company', '未指定')}
职位描述: {job_context.get('job_description', '')[:800] or '无详细描述'}

【职位匹配要求】
1. 优先突出与职位最相关的经验、技能和成就
2. 分析简历与职位的匹配度，给出百分比评分
3. 如果简历中缺少某些技能，强调已有的可迁移技能
4. 调整表达方式，使用职位描述中的关键词（ATS友好）
""",
                'en': f"""
【Target Job Context - Focus on these】
Job Title: {job_context.get('job_title', '')}
Company: {job_context.get('company', 'Not specified')}
Job Description: {(job_context.get('job_description') or '')[:800] or 'No description'}

【Job Matching Requirements】
1. Prioritize experiences most relevant to the job
2. Analyze job match score (give percentage)
3. If skills are missing, emphasize transferable skills
4. Use job description keywords (ATS-friendly)
""",
                'da': f"""
【Målstilling Kontekst】
Jobtitel: {job_context.get('job_title', '')}
Virksomhed: {job_context.get('company', 'Ikke angivet')}
Jobbeskrivelse: {(job_context.get('job_description') or '')[:800] or 'Ingen beskrivelse'}

【Job Match Krav】
1. Prioriter mest relevante erfaringer
2. Analyser job match score (procent)
3. Fremhæv overførbare færdigheder
""",
            }
            lang_key = 'zh' if lang.startswith('zh') else ('da' if lang.startswith('da') else 'en')
            context_text = job_context_section.get(lang_key, job_context_section['en'])
            prompt = prompt.replace('{text[:4000]}', '{text[:3500]}' + context_text + '\n\nResume:\n{text[:3500]}')

        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a professional resume writing expert. Return ONLY valid JSON with suggestions array and resume_highlights array."},
                {"role": "user", "content": prompt}
            ],
            preferred_provider="groq",
            temperature=0.2,
            max_tokens=8000
        )

        content = response.choices[0].message.content.strip()
        # 清理可能的markdown格式
        if content.startswith('```'):
            content = content.split('```')[1]
            if content.startswith('json'):
                content = content[4:]
        
        result = json.loads(content)
        suggestions = result.get('suggestions', [])
        
        for i, s in enumerate(suggestions):
            if 'id' not in s:
                s['id'] = i + 1

        # 限制建议数量
        MAX_SUGGESTIONS = 30
        suggestions = suggestions[:MAX_SUGGESTIONS]
        
        # 返回完整结果（包括resume_highlights）
        return {
            'suggestions': suggestions,
            'resume_highlights': result.get('resume_highlights', []),
            'job_match_score': result.get('job_match_score', 'N/A'),
            'recommendation': result.get('recommendation', '')
        }

    except Exception as e:
        print(f"Polish suggestions failed: {e}, using fallback")
        return {
            'suggestions': generate_polish_fallback(text, lang, job_context),
            'resume_highlights': [],
            'job_match_score': 'N/A',
            'recommendation': ''
        }

def generate_polish_fallback(text: str, lang: str = 'en', job_context: Dict = None) -> List[Dict]:
    """本地规则生成精修建议（无需AI）- 增强版，基于职位上下文"""
    suggestions = []
    job_title = job_context.get('job_title', '') if job_context else ''
    job_desc = job_context.get('job_description', '') if job_context else ''
    
    # 弱动词映射及详细解释
    weak_verb_details = {
        'zh': [
            {'weak': '负责', 'strong': '主导', 'reason': '"负责"只是说明你参与了这项工作，但没有体现你的主动性和成果。"主导"则强调你是项目的推动者和决策者，展现了领导力和责任感。'},
            {'weak': '参与', 'strong': '推动', 'reason': '"参与"暗示你只是团队中的一员，贡献度不明确。"推动"表明你是项目的核心驱动力，能够主动解决问题。'},
            {'weak': '协助', 'strong': '协同', 'reason': '"协助"给人一种配角的感觉。"协同"强调你与团队平等合作，体现跨部门沟通协作能力。'},
            {'weak': '做了', 'strong': '实现', 'reason': '"做了"过于口语化和模糊。"实现"带有目标达成的意味，更符合职场专业表达。'},
            {'weak': '完成了', 'strong': '交付了', 'reason': '"完成了"只说明任务结束。"交付了"暗示你不仅完成，还提供了可用的成果，体现成果导向。'},
            {'weak': '支持', 'strong': '赋能', 'reason': '"支持"显得被动。"赋能"表明你帮助团队或业务实现更大价值，更有影响力。'},
            {'weak': '处理', 'strong': '优化', 'reason': '"处理"只是描述动作。"优化"暗示你不仅处理了问题，还改进了流程或效率。'},
            {'weak': '使用', 'strong': '精通', 'reason': '"使用"说明会用工具。"精通"表明你对该工具有深入掌握，更有竞争力。'}
        ],
        'en': [
            {'weak': 'responsible for', 'strong': 'led', 'reason': '"Responsible for" merely states involvement. "Led" demonstrates you were the driver and decision-maker, showcasing leadership.'},
            {'weak': 'worked on', 'strong': 'delivered', 'reason': '"Worked on" is vague. "Delivered" implies successful completion with tangible results.'},
            {'weak': 'helped with', 'strong': 'collaborated on', 'reason': '"Helped with" suggests a minor role. "Collaborated on" positions you as an equal partner.'},
            {'weak': 'participated in', 'strong': 'spearheaded', 'reason': '"Participated in" shows passive involvement. "Spearheaded" shows you took initiative and led.'},
            {'weak': 'assisted with', 'strong': 'enabled', 'reason': '"Assisted with" implies a supporting role. "Enabled" shows you empowered the team to achieve more.'},
            {'weak': 'handled', 'strong': 'streamlined', 'reason': '"Handled" describes the action. "Streamlined" implies you improved the process or efficiency.'},
            {'weak': 'used', 'strong': 'mastered', 'reason': '"Used" shows basic skill. "Mastered" demonstrates deep expertise.'}
        ],
        'da': [
            {'weak': 'ansvarlig for', 'strong': 'ledede', 'reason': '"Ansvarlig for" viser deltagelse. "Ledede" demonstrerer, at du var drivkraften.'},
            {'weak': 'arbejdede på', 'strong': 'leverede', 'reason': '"Arbejdede på" er vag. "Leverede" indebærer succesfuld gennemførelse.'},
            {'weak': 'deltog i', 'strong': 'ledede', 'reason': '"Deltog i" viser passiv deltagelse. "Ledede" viser initiativ og lederskab.'}
        ]
    }

    text_lower = text.lower()
    verb_details = weak_verb_details.get(lang, weak_verb_details['en'])
    found_verbs = set()

    for detail in verb_details:
        weak = detail['weak']
        strong = detail['strong']
        reason = detail['reason']

        if weak.lower() in text_lower:
            for line in text.split('\n'):
                if weak.lower() in line.lower() and weak.lower() not in found_verbs:
                    improved = line.lower().replace(weak.lower(), strong)
                    if improved and len(improved) > 2:
                        improved = improved[0].upper() + improved[1:]
                    suggestions.append({
                        'id': len(suggestions) + 1,
                        'original': line.strip()[:100],
                        'suggested': improved[:100] if improved else line.strip()[:100],
                        'reason': reason,
                        'type': 'weak_verb',
                        'priority': 'high'
                    })
                    found_verbs.add(weak.lower())
                    break

    # 量化建议
    quantify_reasons = {
        'zh': {'no_num': '你的简历缺少具体的数字和成果量化。建议添加：完成的项目数量（X个）、节省的成本金额（X元）、提升的效率百分比（X%）等具体指标，让招聘方直观看到你的价值。', 'with_job': '结合"{}"职位，建议量化：管理的团队规模（X人）、项目预算（X万）、业绩提升（X%）等与职位要求相关的指标。'.format(job_title)},
        'en': {'no_num': 'Your resume lacks quantified achievements. Add: number of projects, cost savings, efficiency improvement to show measurable value.', 'with_job': 'For "{}" role, quantify: team size (X people), project budget, performance improvement (X%) - metrics relevant to job requirements.'.format(job_title)},
        'da': {'no_num': 'Dit CV mangler kvantificerbare resultater. Tilføj: antal projekter, omkostningsbesparelser, effektivitetsforbedring.', 'with_job': 'For "{}" stilling, kvantificer: teamstørrelse, projektbudget, præstationsforbedring.'.format(job_title)}
    }
    q_texts = quantify_reasons.get(lang, quantify_reasons['en'])
    suggestions.append({
        'id': len(suggestions) + 1,
        'original': '量化成就描述' if lang == 'zh' else 'Quantify achievements',
        'suggested': '领导X人团队，X个月内完成项目，效率提升X%' if lang == 'zh' else 'Led team of X, delivered in X months, improved efficiency by X%',
        'reason': q_texts['with_job'] if job_title else q_texts['no_num'],
        'type': 'quantification',
        'priority': 'high'
    })

    # 职位匹配建议
    if job_title:
        job_reasons = {
            'zh': '您的简历与"{}"职位的匹配度可以通过以下方式提升：1) 使用职位描述中的关键词；2) 突出与职位要求最相关的工作经验；3) 强调可迁移技能。'.format(job_title),
            'en': 'Your match for "{}" can improve by: 1) Using keywords from job description; 2) Highlighting relevant experience; 3) Emphasizing transferable skills.'.format(job_title),
            'da': 'Din match for "{}" kan forbedres ved at: 1) Bruge nøgleord fra jobbeskrivelsen; 2) Fremhæve relevant erfaring.'.format(job_title)
        }
        suggestions.append({
            'id': len(suggestions) + 1,
            'original': '职位关键词匹配' if lang == 'zh' else 'Job keyword matching',
            'suggested': '在简历中融入职位描述关键词' if lang == 'zh' else 'Integrate job description keywords',
            'reason': job_reasons.get(lang, job_reasons['en']),
            'type': 'job_match',
            'priority': 'medium'
        })

    # ATS友好建议
    ats_reasons = {
        'zh': 'ATS(简历筛选系统)通常会扫描关键词。请确保简历包含：1) 职位名称中的关键词；2) 技术技能名称；3) 行业术语。这些能大幅提升简历通过初筛的概率。',
        'en': 'ATS scan for keywords. Ensure your resume includes: 1) Job title keywords; 2) Technical skill names; 3) Industry terminology. This improves initial screening chances.',
        'da': 'ATS scanner efter nøgleord. Sørg for, at dit CV inkluderer jobtitelnøgleord og tekniske færdighedsnavne.'
    }
    suggestions.append({
        'id': len(suggestions) + 1,
        'original': 'ATS关键词优化' if lang == 'zh' else 'ATS keyword optimization',
        'suggested': '添加职位描述中的专业术语和技术关键词' if lang == 'zh' else 'Add professional and technical keywords',
        'reason': ats_reasons.get(lang, ats_reasons['en']),
        'type': 'ats_friendly',
        'priority': 'medium'
    })

    # 北欧职场文化建议
    nordic_reasons = {
        'zh': '在丹麦求职，简历应强调：1) 团队协作和扁平化管理经验；2) 工作与生活平衡的价值观；3) 具体可量化的成果（北欧人喜欢数据）；4) 英语/丹麦语双语能力。',
        'en': 'For Danish jobs, emphasize: 1) Team collaboration; 2) Work-life balance values; 3) Concrete quantifiable results; 4) English/Danish bilingual ability.',
        'da': 'For danske job, fremhæv: 1) Teamsamarbejde; 2) Arbejdslivsbalance; 3) Kvantificerbare resultater; 4) Engelsk/dansk tosprogethed.'
    }
    suggestions.append({
        'id': len(suggestions) + 1,
        'original': '北欧职场文化适配' if lang == 'zh' else 'Nordic workplace culture fit',
        'suggested': '体现团队协作、平等沟通、成果导向的北欧价值观' if lang == 'zh' else 'Demonstrate team collaboration, equality, results-orientation',
        'reason': nordic_reasons.get(lang, nordic_reasons['en']),
        'type': 'culture_fit',
        'priority': 'medium'
    })

    return suggestions

def clean_text_for_pdf(text: str) -> str:
    """清理文本中的问题字符，确保PDF渲染正常"""
    if not text:
        return text
    # Unicode NFC规范化
    text = unicodedata.normalize('NFC', text)
    # 替换常见的特殊Unicode字符为标准ASCII等价物
    replacements = {
        '\u2018': "'",  # 左单引号
        '\u2019': "'",  # 右单引号
        '\u201c': '"',  # 左双引号
        '\u201d': '"',  # 右双引号
        '\u2013': '-',  # en dash
        '\u2014': '-',  # em dash
        '\u00a0': ' ',  # 不间断空格
        '\u3000': ' ',  # 全角空格
        '\u200b': '',   # 零宽空格
        '\ufeff': '',   # BOM
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # 移除所有控制字符（除换行和Tab外）
    text = ''.join(c for c in text if unicodedata.category(c) not in ['Cc'] or c in ['\n', '\t'])
    return text

def apply_polish_suggestions(original_text: str, accepted_suggestions: List[Dict]) -> str:
    """应用用户接受的精修建议"""
    result = original_text
    for suggestion in sorted(accepted_suggestions, key=lambda x: result.rfind(x['original']), reverse=True):
        if suggestion.get('accepted', True):
            original = suggestion['original']
            suggested = suggestion['suggested']
            if original in result:
                result = result.replace(original, suggested, 1)
    
    # 清理文本确保PDF渲染正常
    return clean_text_for_pdf(result)

def discover_skills_from_document(doc_text: str, resume_text: str = "", doc_type: str = "document") -> List[Dict]:
    """从文档中发现隐藏的能力和项目经验"""
    if not AI_AVAILABLE:
        return []

    try:
        prompt = f"""You are a career coach expert. Analyze the provided document and compare with the resume to discover skills and experiences that are NOT mentioned in the resume but are evident in the document.

Document content:
{doc_text[:3000]}

Existing resume:
{resume_text[:2000] if resume_text else 'No resume provided'}

Return a JSON array of discovered skills/experiences:
[
  {{
    "skill": "Skill or experience name",
    "evidence": "Specific evidence from the document",
    "confidence": 85
  }}
]

Return ONLY JSON array with 2-5 items. If no new skills found, return empty array."""

        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "You are a professional career coach. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            preferred_provider="groq",
            temperature=0.3,
            max_tokens=1000
        )

        content = response.choices[0].message.content.strip()
        if content.startswith('```'):
            content = content.split('```')[1]
            if content.startswith('json'):
                content = content[4:]
        return json.loads(content)

    except Exception as e:
        print(f"Discover skills failed: {e}")
        return []

@app.post("/api/resume/polish")
async def polish_resume(request: PolishRequest):
    """生成逐条精修建议 + 简历亮点总结"""
    try:
        if request.language == "auto":
            lang = detect_language(request.resume_text)
        else:
            lang = request.language

        # 构建职位上下文
        job_context = None
        if request.job_title or request.job_description:
            job_context = {
                "job_title": request.job_title,
                "job_description": request.job_description,
                "company": request.company
            }

        result = generate_polish_suggestions(request.resume_text, lang, job_context, request.resume_structure)

        # 兼容新旧格式（旧版返回list，新版返回dict）
        if isinstance(result, dict):
            suggestions = result.get('suggestions', [])
            resume_highlights = result.get('resume_highlights', [])
            job_match_score = result.get('job_match_score', 'N/A')
            recommendation = result.get('recommendation', '')
        else:
            suggestions = result
            resume_highlights = []
            job_match_score = 'N/A'
            recommendation = ''

        return {
            "success": True,
            "total": len(suggestions),
            "suggestions": suggestions,
            "resume_highlights": resume_highlights,
            "job_match_score": job_match_score,
            "recommendation": recommendation,
            "language": lang,
            "language_name": {"zh": "中文", "en": "English", "da": "Dansk"}.get(lang, "English"),
            "job_context": job_context
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/resume/apply-polish")
async def apply_polish(request: PolishApplyRequest):
    """应用用户选择的精修建议"""
    try:
        accepted = [s for s in request.suggestions if s.get('accepted', True)]
        polished_text = apply_polish_suggestions(request.original_text, accepted)

        return {
            "success": True,
            "accepted_count": len(accepted),
            "total_count": len(request.suggestions),
            "polished_resume": polished_text
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/resume/discover-skills")
async def discover_skills(request: DiscoverSkillsRequest):
    """从文档中发现隐藏能力"""
    try:
        discovered = discover_skills_from_document(
            request.document_text,
            request.resume_text,
            request.document_type
        )

        return {
            "success": True,
            "discovered_skills": discovered,
            "count": len(discovered)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ========== 简历语种转换 API ==========

class TranslateResumeRequest(BaseModel):
    resume_text: str
    source_lang: str = "auto"  # 源语言: zh, en, da, auto
    target_lang: str  # 目标语言: zh, en, da

def translate_resume_with_ai(resume_text: str, source_lang: str, target_lang: str) -> str:
    """使用 AI 翻译简历到目标语言（AI不可用时使用后备翻译）"""
    
    # 常用简历术语对照表
    TERM_TRANSLATIONS = {
        ('zh', 'en'): {
            '个人信息': 'Personal Information',
            '姓名': 'Name',
            '邮箱': 'Email',
            '电话': 'Phone',
            '地址': 'Address',
            '教育背景': 'Education',
            '工作经验': 'Work Experience',
            '工作经历': 'Work Experience',
            '项目经验': 'Project Experience',
            '项目经历': 'Project Experience',
            '技能特长': 'Skills',
            '专业技能': 'Professional Skills',
            '自我评价': 'Self Summary',
            '求职意向': 'Career Objective',
            '证书': 'Certificates',
            '语言能力': 'Language Skills',
            '培训经历': 'Training',
            '公司简介': 'Company Description',
        },
        ('en', 'zh'): {
            'personal information': '个人信息',
            'education': '教育背景',
            'work experience': '工作经验',
            'skills': '技能特长',
            'summary': '自我评价',
            'objective': '求职意向',
            'experience': '经历',
            'project': '项目',
        },
        ('zh', 'da'): {
            '个人信息': 'Personlige Oplysninger',
            '姓名': 'Navn',
            '邮箱': 'E-mail',
            '电话': 'Telefon',
            '教育背景': 'Uddannelse',
            '工作经验': 'Arbejdserfaring',
            '工作经历': 'Arbejdserfaring',
            '项目经验': 'Projekterfaring',
            '技能特长': 'Kompetencer',
            '自我评价': 'Profil',
        },
        ('en', 'da'): {
            'personal information': 'Personlige Oplysninger',
            'name': 'Navn',
            'email': 'E-mail',
            'phone': 'Telefon',
            'education': 'Uddannelse',
            'work experience': 'Arbejdserfaring',
            'skills': 'Kompetencer',
            'summary': 'Profil',
            'experience': 'Erfaring',
            'project': 'Projekt',
        },
        ('da', 'en'): {
            'personlige oplysninger': 'Personal Information',
            'navn': 'Name',
            'e-mail': 'Email',
            'telefon': 'Phone',
            'uddannelse': 'Education',
            'arbejdserfaring': 'Work Experience',
            'kompetencer': 'Skills',
            'profil': 'Summary',
        },
        ('da', 'zh'): {
            'personlige oplysninger': '个人信息',
            'navn': '姓名',
            'uddannelse': '教育背景',
            'arbejdserfaring': '工作经验',
        }
    }
    
    # 如果AI可用，使用AI翻译
    if AI_AVAILABLE:
        lang_names = {
            'zh': '中文',
            'en': 'English',
            'da': 'Dansk (Danish)'
        }
        
        source_name = lang_names.get(source_lang, source_lang)
        target_name = lang_names.get(target_lang, target_lang)
        
        prompt = f"""You are a professional resume translator. Translate the following resume from {source_name} to {target_name}.

Requirements:
1. Maintain professional resume format and structure
2. Adapt cultural conventions for {target_name} resumes
3. Keep all dates, numbers, and factual information accurate
4. Translate company names only if they have common {target_name} equivalents
5. Maintain professional tone appropriate for job applications

Resume to translate:
{resume_text[:4000]}

Return ONLY the translated resume text, nothing else."""

        try:
            response = smart_ai_request(
                messages=[
                    {"role": "system", "content": f"You are a professional resume translator specializing in {source_name} to {target_name} translation."},
                    {"role": "user", "content": prompt}
                ],
                preferred_provider="groq",
                temperature=0.3,
                max_tokens=3000
            )
            
            content = response.choices[0].message.content.strip()
            # 清理文本确保PDF渲染正常
            return clean_text_for_pdf(content) if content else content
        except Exception as e:
            print(f"Resume AI translation failed: {e}")
            # AI失败时继续使用后备翻译
    
    # 后备翻译：使用术语对照表
    translation_map = TERM_TRANSLATIONS.get((source_lang, target_lang), {})
    if not translation_map:
        return None  # 没有可用的翻译对照
    
    result = resume_text
    for source_term, target_term in translation_map.items():
        # 不区分大小写替换
        import re
        pattern = re.compile(re.escape(source_term), re.IGNORECASE)
        result = pattern.sub(target_term, result)
    
    return result

@app.post("/api/resume/translate")
async def translate_resume_endpoint(request: TranslateResumeRequest):
    """一键转换简历语种"""
    try:
        # 自动检测源语言
        if request.source_lang == "auto":
            source_lang = detect_language(request.resume_text)
        else:
            source_lang = request.source_lang
        
        # 如果源语言和目标语言相同，直接返回
        if source_lang == request.target_lang:
            return {
                "success": True,
                "translated_resume": request.resume_text,
                "source_lang": source_lang,
                "target_lang": request.target_lang,
                "message": "源语言和目标语言相同，无需翻译"
            }
        
        # 使用 AI 翻译
        if AI_AVAILABLE:
            translated = translate_resume_with_ai(
                request.resume_text,
                source_lang,
                request.target_lang
            )

            if translated:
                return {
                    "success": True,
                    "translated_resume": translated,
                    "source_lang": source_lang,
                    "target_lang": request.target_lang,
                    "source_lang_name": {"zh": "中文", "en": "English", "da": "Dansk"}.get(source_lang, source_lang),
                    "target_lang_name": {"zh": "中文", "en": "English", "da": "Dansk"}.get(request.target_lang, request.target_lang)
                }

        # AI 不可用或失败时，使用后备翻译（术语对照表）
        translated = translate_resume_with_ai(
            request.resume_text,
            source_lang,
            request.target_lang
        )

        if translated:
            return {
                "success": True,
                "translated_resume": translated,
                "source_lang": source_lang,
                "target_lang": request.target_lang,
                "source_lang_name": {"zh": "中文", "en": "English", "da": "Dansk"}.get(source_lang, source_lang),
                "target_lang_name": {"zh": "中文", "en": "English", "da": "Dansk"}.get(request.target_lang, request.target_lang),
                "note": "使用基础术语翻译，建议人工校对"
            }

        # 如果所有翻译都失败，返回错误
        return {
            "success": False,
            "error": "翻译服务暂时不可用，请稍后重试",
            "source_lang": source_lang,
            "target_lang": request.target_lang
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/upload-document")
async def upload_document(file: UploadFile = File(...)):
    """上传文档并解析内容"""
    try:
        content = await file.read()
        filename_lower = file.filename.lower()

        text = ""
        if filename_lower.endswith('.txt'):
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'gb18030']:
                try:
                    text = content.decode(encoding)
                    break
                except:
                    continue
        elif filename_lower.endswith('.pdf'):
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            text += t + "\n"
            except:
                pass
        elif filename_lower.endswith('.docx'):
            try:
                import docx
                document = docx.Document(io.BytesIO(content))
                for para in document.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
            except:
                pass
        elif filename_lower.endswith('.html') or filename_lower.endswith('.htm'):
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(content.decode('utf-8', errors='ignore'), 'html.parser')
                # 移除脚本和样式
                for script in soup(["script", "style"]):
                    script.decompose()
                # 获取纯文本
                text = soup.get_text(separator=' ', strip=True)
            except Exception as e:
                text = content.decode('utf-8', errors='ignore')
        elif filename_lower.endswith('.eml'):
            try:
                from bs4 import BeautifulSoup
                msg_content = content.decode('utf-8', errors='ignore')
                soup = BeautifulSoup(msg_content, 'html.parser')
                # 获取正文
                text = soup.get_text(separator=' ', strip=True)
            except Exception as e:
                text = content.decode('utf-8', errors='ignore')

        return {
            "success": True,
            "filename": file.filename,
            "text": text[:5000],
            "detected_language": detect_language(text)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============ Beta 测试 API ============
from early_bird import collector

class BetaSubmissionRequest(BaseModel):
    email: str
    name: str = ""
    phone: str = ""
    country: str = ""
    city: str = ""
    resume_text: str = ""
    resume_file_name: str = ""
    job_links: List[Dict] = []
    social_accounts: List[Dict] = []
    notes: str = ""
    source: str = "website"

@app.post("/beta/submit")
async def beta_submit(request: BetaSubmissionRequest):
    """Beta 测试用户提交"""
    try:
        result = collector.submit_form(
            email=request.email,
            name=request.name,
            phone=request.phone,
            country=request.country,
            city=request.city,
            resume_text=request.resume_text,
            resume_file_name=request.resume_file_name,
            job_links=request.job_links,
            social_accounts=request.social_accounts,
            notes=request.notes,
            source=request.source
        )
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/beta/submissions")
async def beta_submissions(status: str = "all", source: str = "all", limit: int = 50):
    """获取 Beta 测试提交列表（管理用）"""
    try:
        submissions = collector.list_submissions(status=status, source=source, limit=limit)
        return {"success": True, "submissions": submissions, "count": len(submissions)}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/beta/statistics")
async def beta_statistics():
    """获取 Beta 测试统计数据"""
    try:
        stats = collector.get_statistics()
        return {"success": True, "statistics": stats}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/beta/submission/{submission_id}")
async def beta_submission_detail(submission_id: str):
    """获取单个提交详情"""
    try:
        submission = collector.get_submission(submission_id)
        if submission:
            return {"success": True, "submission": submission}
        else:
            return {"success": False, "error": "Submission not found"}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/beta/process/{submission_id}")
async def beta_process_submission(submission_id: str, status: str = "processed", notes: str = ""):
    """更新提交处理状态"""
    try:
        success = collector.update_status(submission_id, status, notes)
        if success:
            collector.log_processing(submission_id, f"status_changed_to_{status}", notes)
            return {"success": True, "message": f"Submission {submission_id} updated to {status}"}
        else:
            return {"success": False, "error": "Submission not found"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============ 静态文件服务 ============
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
import os

# 获取当前文件所在目录
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FRONTEND_DIR = os.path.join(BASE_DIR, "frontend")

@app.get("/")
async def root():
    """API 根路径 - 返回 API 信息"""
    return {
        "service": "JobMatchAI API",
        "version": "2.0.0",
        "endpoints": {
            "health": "/health",
            "frontend": "/app",
            "analyze": "/analyze",
            "jobs_search": "/jobs/search",
            "cover_letter": "/cover-letter"
        }
    }

@app.get("/app")
async def app_page():
    """前端页面 - 语言选择"""
    return FileResponse(os.path.join(FRONTEND_DIR, "index.html"))

@app.get("/app-en")
async def app_en_page():
    """Frontend page - English version"""
    return FileResponse(os.path.join(FRONTEND_DIR, "index-en.html"))

@app.get("/app-zh")
async def app_zh_page():
    """前端页面 - 中文版"""
    return FileResponse(os.path.join(FRONTEND_DIR, "index-zh.html"))

@app.get("/beta.html")
async def beta_page():
    """Beta 测试页面"""
    return FileResponse(os.path.join(FRONTEND_DIR, "beta.html"))

@app.get("/admin")
async def admin_page():
    """管理后台页面"""
    admin_path = os.path.join(BASE_DIR, "admin.html")
    if os.path.exists(admin_path):
        return FileResponse(admin_path)
    return {"error": "Admin page not found"}

@app.get("/health")
async def health_check():
    """健康检查"""
    return {"status": "healthy", "version": "2.0.0", "ai_available": AI_AVAILABLE}

@app.get("/ai/status")
async def ai_status():
    """AI 服务状态检查"""
    return get_ai_status()

@app.post("/ai/check-groq")
async def check_groq_recovery_endpoint():
    """手动触发 Groq 配额恢复检测"""
    recovered = check_groq_recovery()
    status = get_ai_status()
    return {
        "groq_recovered": recovered,
        **status
    }

# ===== 后台任务：定期检测 Groq 恢复 =====
def groq_recovery_checker():
    """后台线程：每5分钟检测 Groq 配额是否恢复"""
    import time
    while True:
        try:
            time.sleep(300)  # 5分钟检查一次
            if not GROQ_AVAILABLE:
                print("🔄 [Background] 检测 Groq 配额状态...")
                if check_groq_recovery():
                    print("✅ [Background] Groq 配额已恢复！")
        except Exception as e:
            print(f"⚠️ [Background] Groq 检测出错: {e}")

# 启动后台任务
if groq_client and openai_client:
    recovery_thread = threading.Thread(target=groq_recovery_checker, daemon=True)
    recovery_thread.start()
    print("✅ [Background] Groq 恢复检测任务已启动（每5分钟）")

# ===== 意见反馈 API =====
class FeedbackRequest(BaseModel):
    userId: Optional[str] = None
    email: Optional[str] = None
    type: str  # suggestion, bug, other
    content: str

@app.post("/api/feedback")
async def submit_feedback(feedback: FeedbackRequest, request: Request):
    """提交用户反馈"""
    try:
        feedback_data = {
            "user_id": feedback.userId or "anonymous",
            "email": feedback.email or "",
            "type": feedback.type,
            "content": feedback.content,
            "timestamp": datetime.now().isoformat(),
            "user_agent": request.headers.get("user-agent", ""),
        }
        
        # 发送到管理员邮箱（如果配置了）
        if os.getenv("NOTIFY_EMAIL"):
            try:
                # 简化版：直接打印到日志，正式环境可用邮件服务
                print(f"[FEEDBACK] From: {feedback.email or 'Anonymous'} | Type: {feedback.type}")
                print(f"[FEEDBACK] Content: {feedback.content}")
            except Exception as e:
                print(f"[FEEDBACK] Email notification failed: {e}")
        
        # 返回成功
        return {"status": "ok", "message": "反馈已收到，感谢您的建议！"}
    except Exception as e:
        print(f"[FEEDBACK ERROR] {e}")
        return {"status": "error", "message": str(e)}

# ===== 客户数据追踪 API（匿名，仅用于AI训练）=====
class AnonymousTrackingRequest(BaseModel):
    session_id: str
    industry: Optional[str] = None  # 行业分类（不是个人信息）
    job_title: Optional[str] = None  # 职位类型（不是具体公司）
    resume_skills: Optional[List[str]] = None  # 提取的技能标签
    action_type: str  # analyze_resume, generate_cover, etc.

@app.post("/api/track")
async def track_anonymous_data(data: AnonymousTrackingRequest):
    """匿名追踪客户行为数据，用于AI训练"""
    try:
        tracking_data = {
            "session_id": data.session_id[:16] + "***",  # 只保留session前16位
            "industry": data.industry or "unknown",
            "job_title": data.job_title or "unknown",
            "resume_skills": data.resume_skills or [],
            "action_type": data.action_type,
            "timestamp": datetime.now().isoformat(),
        }
        
        # 只记录行业和技能类型，不记录任何个人信息
        # 这些数据用于训练AI更好地理解不同行业求职者的需求
        print(f"[TRACKING] Industry: {data.industry} | Action: {data.action_type} | Skills: {len(data.resume_skills or [])}")
        
        return {"status": "ok", "message": "Tracking recorded (anonymous)"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.get("/api/track/stats")
async def get_tracking_stats():
    """获取追踪统计（仅显示聚合数据，不含个人信息）"""
    # 这里可以返回一些聚合统计，如各行业的使用量
    # 实际实现需要查询数据库
    return {
        "status": "ok",
        "message": "Anonymous statistics only",
        "note": "No personal data is stored"
    }

# 简历模板文件路由
@app.get("/templates/resume-{template_name}.html")
async def serve_resume_template(template_name: str):
    """服务简历模板HTML文件"""
    file_path = os.path.join(FRONTEND_DIR, "templates", f"resume-{template_name}.html")
    if os.path.exists(file_path):
        return FileResponse(file_path, media_type="text/html")
    return {"detail": "Template not found"}

# 模板文件路由
@app.get("/static/templates/{filename}")
async def serve_template(filename: str):
    """服务模板JS文件"""
    file_path = os.path.join(FRONTEND_DIR, "templates", filename)
    if os.path.exists(file_path):
        return FileResponse(file_path, media_type="application/javascript")
    return {"detail": "File not found"}

# 通用静态文件路由
@app.get("/static/{filepath:path}")
async def serve_static(filepath: str):
    """服务其他静态文件"""
    file_path = os.path.join(FRONTEND_DIR, filepath)
    if os.path.exists(file_path):
        return FileResponse(file_path)
    return {"detail": "File not found"}


# === 简历模版 API ===

from resume_templates import (
    get_template_list,
    get_template_by_id,
    get_template_html,
    get_external_links,
    suggest_ats_keywords
)

from industry_tracker import (
    detect_industry,
    get_industry_by_id,
    get_all_industries,
    CoverLetterEvolution,
    get_salary_estimate,
    format_salary_range,
    get_salary_tips,
    get_industry_strategy
)

# 初始化求职信进化追踪器（按用户ID存储）
letter_evolution_store: Dict[str, CoverLetterEvolution] = {}


@app.get("/templates/list")
def list_templates(language: str = "en"):
    """获取简历模版列表"""
    return {
        "success": True,
        "templates": get_template_list(language),
        "external_links": get_external_links(language)
    }


@app.get("/templates/{template_id}")
def get_template(template_id: str, language: str = "en"):
    """获取指定模版"""
    template = get_template_by_id(template_id, language)
    if not template:
        return {"success": False, "error": "Template not found"}
    return {"success": True, "template": template}


@app.post("/templates/render")
async def render_template(
    template_id: str = Form(...),
    resume_data: str = Form(...)  # JSON
):
    """根据简历数据渲染模版"""
    try:
        import json
        data = json.loads(resume_data)
        html = get_template_html(template_id, data)
        if not html:
            return {"success": False, "error": "Failed to render template"}
        return {"success": True, "html": html}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ===== 简历模板AI个性化推荐 =====
@app.post("/api/resume/template-suggestion")
async def get_template_suggestion(
    resume_structure: dict = Body(...),
    job_description: str = Body(""),
    job_title: str = Body(""),
    language: str = Body("en")
):
    """
    AI驱动的简历模板个性化推荐
    
    分析用户简历和目标职位，推荐最适合的简历模板
    """
    if not AI_AVAILABLE:
        return {
            "success": False,
            "error": "AI服务暂不可用",
            "suggestions": get_default_template_suggestions(language)
        }
    
    try:
        # 提取简历关键信息
        tech_skills = resume_structure.get('skills', {}).get('technical', [])
        soft_skills = resume_structure.get('skills', {}).get('soft', [])
        education = resume_structure.get('education', [])
        experience = resume_structure.get('experience', [])
        total_years = resume_structure.get('total_experience_years', 0)
        
        # 行业判断
        industry_keywords = {
            'finance': ['财务', '金融', '会计', '银行', '投资', '财务', 'budget', 'accounting', 'finance', 'banking'],
            'tech': ['Python', 'Java', '开发', '工程师', '技术', '软件', 'IT', 'developer', 'engineer', 'software', 'data', 'machine learning'],
            'business': ['管理', '咨询', '运营', '商业', '市场', '销售', 'management', 'consulting', 'operations', 'business', 'marketing'],
            'creative': ['设计', '创意', 'UX', 'UI', '产品', '品牌', 'design', 'creative', 'UX', 'UI', 'product', 'brand'],
            'healthcare': ['医疗', '健康', '医院', '生物', '制药', 'healthcare', 'medical', 'hospital', 'pharma'],
            'education': ['教育', '教师', '学术', '研究', 'education', 'teacher', 'academic', 'research']
        }
        
        detected_industry = 'business'
        all_text = ' '.join(tech_skills + soft_skills).lower()
        for industry, keywords in industry_keywords.items():
            if any(kw.lower() in all_text for kw in keywords):
                detected_industry = industry
                break
        
        # 构建AI prompt
        prompt = f"""你是一名专业的简历设计顾问。根据候选人的背景和目标职位，推荐最适合的简历模板风格。

【候选人背景】
- 行业：{detected_industry}
- 工作经验：{total_years}年
- 技术技能：{', '.join(tech_skills[:8]) if tech_skills else '未指定'}
- 软技能：{', '.join(soft_skills[:5]) if soft_skills else '未指定'}
- 学历：{len(education)}段

【目标职位】
{job_title or '未指定'}

【职位描述摘要】
{job_description[:500] if job_description else '未提供'}

请返回JSON格式的模板推荐（只返回JSON，不要其他内容）：
{{
    "recommended_template": "modern/classic/creative 三选一",
    "confidence": 0.85,
    "reasoning": "推荐理由，用2-3句话说明为什么这个模板最适合",
    "alternative_templates": [
        {{
            "template": "modern/classic/creative",
            "reason": "备选理由"
        }}
    ],
    "customization_tips": [
        "个性化建议1",
        "个性化建议2"
    ],
    "color_suggestion": "深蓝色/绿色/紫色等",
    "layout_preference": "单栏/双栏"
}}

【模板风格说明】
- classic（经典型）：适合金融、咨询、法律、政府等保守行业，深蓝色系，结构清晰
- modern（现代型）：适合科技、互联网、创业公司，简洁专业，视觉层次分明
- creative（创意型）：适合设计、媒体、营销、创意行业，色彩丰富，视觉冲击力强</parameter>"""
        
        response = smart_ai_request(
            messages=[
                {"role": "system", "content": "你是一名专业的简历设计顾问。返回ONLY valid JSON, no markdown, no explanation."},
                {"role": "user", "content": prompt}
            ],
            preferred_provider="groq",
            temperature=0.3,
            max_tokens=800
        )
        
        content = response.choices[0].message.content.strip()
        if content.startswith('```'):
            content = content.split('```')[1]
            if content.startswith('json'):
                content = content[4:]
        
        result = json.loads(content)
        
        return {
            "success": True,
            "suggestions": result,
            "detected_industry": detected_industry,
            "total_years": total_years,
            "skills_count": len(tech_skills) + len(soft_skills)
        }
        
    except Exception as e:
        print(f"Template suggestion error: {e}")
        return {
            "success": False,
            "error": str(e),
            "suggestions": get_default_template_suggestions(language)
        }


def get_default_template_suggestions(language: str = "en") -> dict:
    # Default suggestions when AI is unavailable
    if language == "zh":
        return {
            "recommended_template": "modern",
            "confidence": 0.5,
            "reasoning": "推荐使用现代简约型模板，适合大多数职位申请",
            "alternative_templates": [
                {"template": "classic", "reason": "传统行业如金融、法律可考虑经典型"},
                {"template": "creative", "reason": "创意行业如设计、营销可考虑创意型"}
            ],
            "customization_tips": [
                "根据目标公司的行业特点选择模板风格",
                "保持简历内容简洁有力",
                "突出与职位相关的技能和经验"
            ],
            "color_suggestion": "蓝色",
            "layout_preference": "单栏"
        }
    else:
        return {
            "recommended_template": "modern",
            "confidence": 0.5,
            "reasoning": "Modern template recommended for most job applications",
            "alternative_templates": [
                {"template": "classic", "reason": "Consider classic for finance, law, government"},
                {"template": "creative", "reason": "Consider creative for design, marketing, media"}
            ],
            "customization_tips": [
                "Choose template style based on target company's industry",
                "Keep resume content concise and impactful",
                "Highlight skills and experience relevant to the position"
            ],
            "color_suggestion": "Blue",
            "layout_preference": "Single column"
        }


@app.get("/ats/keywords")
def get_ats_keywords_suggestions(
    job_title: str = "",
    industry: str = ""
):
    """获取ATS关键词建议"""
    keywords = suggest_ats_keywords(job_title, industry)
    return {
        "success": True,
        "keywords": keywords,
        "note": "在简历中使用这些关键词可以提高ATS通过率"
    }


# === 行业追踪 API ===

@app.get("/industries/list")
def list_industries(language: str = "en"):
    """获取行业列表"""
    return {
        "success": True,
        "industries": get_all_industries(language)
    }


@app.post("/industries/detect")
async def detect_job_industry(
    job_title: str = Form(...),
    description: str = Form("")
):
    """根据职位信息识别行业"""
    result = detect_industry(job_title, description)
    return {
        "success": True,
        "detection": result
    }


@app.get("/industries/{industry_id}/strategy")
def get_industry_application_strategy(industry_id: str):
    """获取特定行业的申请策略"""
    strategy = get_industry_strategy(industry_id)
    return {
        "success": True,
        "strategy": strategy
    }


# === 求职信进化追踪 API ===

@app.get("/letter-evolution/{user_id}")
def get_letter_evolution(user_id: str):
    """获取用户的求职信进化历史"""
    if user_id not in letter_evolution_store:
        letter_evolution_store[user_id] = CoverLetterEvolution()
    
    tracker = letter_evolution_store[user_id]
    summary = tracker.get_evolution_summary()
    
    return {
        "success": True,
        "evolution": summary
    }


@app.post("/letter-evolution/add")
async def add_cover_letter_record(
    user_id: str = Form(...),
    company: str = Form(...),
    job_title: str = Form(...),
    content: str = Form(...),
    industry: str = Form("")
):
    """记录一封求职信"""
    if user_id not in letter_evolution_store:
        letter_evolution_store[user_id] = CoverLetterEvolution()
    
    tracker = letter_evolution_store[user_id]
    letter_info = tracker.add_letter(company, job_title, content, industry)
    
    return {
        "success": True,
        "letter": letter_info
    }


@app.post("/letter-evolution/suggestions")
async def get_letter_improvement_suggestions(
    user_id: str = Form(...),
    new_job_title: str = Form(...),
    new_industry: str = Form("")
):
    """获取改进建议"""
    if user_id not in letter_evolution_store:
        letter_evolution_store[user_id] = CoverLetterEvolution()
    
    tracker = letter_evolution_store[user_id]
    suggestions = tracker.suggest_improvements(new_job_title, new_industry)
    
    return {
        "success": True,
        "suggestions": suggestions
    }


# === 薪资范围查询 API ===

@app.get("/salary/estimate")
def get_salary(
    job_title: str = "",
    location: str = "",
    country: str = "DK"
):
    """获取薪资估算"""
    salary_data = get_salary_estimate(job_title, location, country)
    formatted = format_salary_range(salary_data)
    tips = get_salary_tips(country)
    
    return {
        "success": True,
        "salary": salary_data,
        "formatted": formatted,
        "tips": tips,
        "disclaimer": salary_data.get("disclaimer_en", "")
    }


# === 启动 ===
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
