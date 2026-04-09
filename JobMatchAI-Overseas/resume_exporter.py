"""
JobMatchAI - 简历导出模块
使用 python-docx 生成 Word 文档

支持语言：
- en: English
- zh: 中文
- da: 丹麦语

Copyright © 2026 JobMatchAI. All rights reserved.
"""

import os
import io
import sys
import re
import unicodedata
from typing import Dict, Any, Optional, List
from datetime import datetime


def strip_emoji(text: str) -> str:
    """
    移除或替换 emoji 字符，防止 Word/PDF 字体渲染问题。
    保留普通 Unicode 文字（包括中文、丹麦语特殊字母）。
    """
    if not text:
        return text
    
    # 常见 emoji 替换为纯文字
    EMOJI_REPLACEMENTS = {
        "📞": "Tel:",
        "📧": "Email:",
        "📍": "Location:",
        "🔗": "LinkedIn:",
        "⭐": "*",
        "✅": "[OK]",
        "❌": "[X]",
        "🎯": "->",
        "💼": "",
        "🏢": "",
        "📝": "",
        "🌍": "",
        "📊": "",
        "🔧": "",
        "💡": "",
        "🚀": "",
    }
    
    for emoji, replacement in EMOJI_REPLACEMENTS.items():
        text = text.replace(emoji, replacement)
    
    # 过滤掉剩余的 emoji（Unicode 范围）
    result = []
    for char in text:
        cp = ord(char)
        # 保留基本拉丁、扩展拉丁（含丹麦语）、CJK（中文）、常用标点
        if (cp < 0x2600 or  # 基础字符
            (0x4E00 <= cp <= 0x9FFF) or   # CJK 中文
            (0x00C0 <= cp <= 0x024F) or   # 扩展拉丁（丹麦语 Æ Ø Å 等）
            (0x2000 <= cp <= 0x206F)):    # 常用标点
            result.append(char)
        elif unicodedata.category(char).startswith('L'):
            # 保留所有字母类字符
            result.append(char)
        elif char in ' \t\n\r.,;:!?()[]{}/-_@#$%&+=<>|\\\'\"':
            result.append(char)
        # 其余（emoji 等）跳过
    
    return "".join(result)

# 尝试导入 docx
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. Resume export will not be available.")

# 导入数据库模块
sys.path.insert(0, os.path.dirname(__file__))
from database import get_primary_resume, get_resume_by_id, get_user_resumes


class ResumeExporter:
    """简历导出器"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for resume export")
    
    def create_document(self) -> Document:
        """创建新文档"""
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        return doc
    
    def add_heading(self, doc: Document, text: str, level: int = 1, color: RGBColor = None):
        """添加标题"""
        heading = doc.add_heading(text, level=level)
        if color:
            for run in heading.runs:
                run.font.color.rgb = color
        return heading
    
    def add_paragraph(self, doc: Document, text: str, bold: bool = False, 
                      italic: bool = False, size: int = 11, align: str = "left"):
        """添加段落（自动过滤 emoji 防止 docx 损坏）"""
        # 过滤 emoji，防止 Word 字体无法渲染导致文件损坏
        text = strip_emoji(text)
        para = doc.add_paragraph()
        
        # 设置对齐
        if align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "justify":
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 添加文本
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        
        return para
    
    def add_bullet_list(self, doc: Document, items: List[str], indent: int = 0):
        """添加项目符号列表"""
        for item in items:
            para = doc.add_paragraph(item, style='List Bullet')
            if indent > 0:
                para.paragraph_format.left_indent = Inches(indent * 0.25)
        return para
    
    def add_table(self, doc: Document, headers: List[str], rows: List[List[str]], 
                  col_widths: List[float] = None):
        """添加表格"""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        
        # 添加表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # 添加数据行
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_text in enumerate(row_data):
                row_cells[col_idx].text = str(cell_text)
        
        # 设置列宽
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(width)
        
        return table
    
    def export_english_resume(self, resume_data: Dict) -> Document:
        """导出英文简历"""
        doc = self.create_document()
        
        # 姓名和联系方式
        name = resume_data.get('name', resume_data.get('content', '').split('\n')[0] if resume_data.get('content') else 'Your Name')
        if '|' in name or '#' in name:
            name = name.split('|')[0].split('#')[0].strip()
        
        self.add_paragraph(doc, name, bold=True, size=20, align="center")
        
        # 联系方式行（纯文字，不用 emoji 避免 docx 损坏）
        contact_parts = []
        if resume_data.get('phone'): contact_parts.append(f"Tel: {resume_data['phone']}")
        if resume_data.get('email'): contact_parts.append(f"Email: {resume_data['email']}")
        if resume_data.get('location'): contact_parts.append(f"Location: {resume_data['location']}")
        if resume_data.get('linkedin'): contact_parts.append(f"LinkedIn: {resume_data['linkedin']}")
        
        if contact_parts:
            self.add_paragraph(doc, "  |  ".join(contact_parts), size=10, align="center")
        
        doc.add_paragraph()
        
        # 解析内容
        content = resume_data.get('content', '')
        
        # 专业简介
        summary_match = re.search(r'##\s*(?:Professional\s*)?Summary[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if summary_match:
            self.add_heading(doc, "Professional Summary", level=2)
            summary_text = summary_match.group(1).strip()
            for para in summary_text.split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
            doc.add_paragraph()
        
        # 核心能力
        competencies_match = re.search(r'##\s*(?:Core\s*)?Competencies[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if competencies_match:
            self.add_heading(doc, "Core Competencies", level=2)
            for line in competencies_match.group(1).strip().split('\n'):
                if line.strip().startswith('-'):
                    self.add_paragraph(doc, line.strip().replace('- ', ''))
                elif line.strip():
                    self.add_paragraph(doc, line.strip())
            doc.add_paragraph()
        
        # 工作经历
        experience_match = re.search(r'##\s*(?:Professional\s*)?Experience[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if experience_match:
            self.add_heading(doc, "Professional Experience", level=2)
            exp_content = experience_match.group(1)
            
            # 分割各个职位
            job_pattern = r'(?:###\s*)?([^\n]+?)\s*\|\s*([^\n]+?)\s*\n\*\*(\d{4})(?:\s*[-–]\s*(\d{4}|Present))?\*\*\s*\n(.*?)(?=(?:###|\n##|$))'
            jobs = re.findall(job_pattern, exp_content, re.DOTALL)
            
            for job in jobs:
                company = job[0].strip()
                title = job[1].strip()
                start_year = job[2]
                end_year = job[3] if len(job) > 3 else ""
                description = job[4] if len(job) > 4 else ""
                
                # 职位标题和公司
                para = doc.add_paragraph()
                run = para.add_run(f"{title} | {company}")
                run.bold = True
                run.font.size = Pt(12)
                
                # 时间
                if start_year:
                    time_str = f"{start_year}"
                    if end_year:
                        time_str += f" - {end_year}"
                    time_para = doc.add_paragraph(time_str)
                    time_para.runs[0].italic = True
                    time_para.runs[0].font.size = Pt(10)
                
                # 职责描述
                for line in description.split('\n'):
                    if line.strip().startswith('-'):
                        self.add_paragraph(doc, line.strip().replace('- ', '• '))
                    elif line.strip().startswith('**') and 'highlight' in line.lower():
                        self.add_paragraph(doc, line.strip().replace('**', ''), italic=True)
                    elif line.strip():
                        self.add_paragraph(doc, line.strip())
                doc.add_paragraph()
        
        # 教育背景
        education_match = re.search(r'##\s*Education[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if education_match:
            self.add_heading(doc, "Education", level=2)
            for line in education_match.group(1).strip().split('\n'):
                if line.strip().startswith('-'):
                    self.add_paragraph(doc, line.strip().replace('- ', ''))
                elif line.strip():
                    self.add_paragraph(doc, line.strip())
            doc.add_paragraph()
        
        # 证书
        cert_match = re.search(r'##\s*Certifications?[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if cert_match:
            self.add_heading(doc, "Certifications", level=2)
            for line in cert_match.group(1).strip().split('\n'):
                if line.strip().startswith('-'):
                    self.add_paragraph(doc, line.strip().replace('- ', '• '))
                elif line.strip():
                    self.add_paragraph(doc, line.strip())
            doc.add_paragraph()
        
        # 技能
        skills_match = re.search(r'##\s*(?:Technical\s*)?Skills[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if skills_match:
            self.add_heading(doc, "Technical Skills", level=2)
            for line in skills_match.group(1).strip().split('\n'):
                if line.strip():
                    self.add_paragraph(doc, line.strip())
            doc.add_paragraph()
        
        # 语言能力
        lang_match = re.search(r'##\s*Languages?[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if lang_match:
            self.add_heading(doc, "Languages", level=2)
            for line in lang_match.group(1).strip().split('\n'):
                if line.strip():
                    self.add_paragraph(doc, line.strip())
        
        return doc
    
    def export_chinese_resume(self, resume_data: Dict) -> Document:
        """导出中文简历"""
        doc = self.create_document()
        
        # 设置中文字体支持
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        name = resume_data.get('name', resume_data.get('content', '').split('\n')[0] if resume_data.get('content') else '您的姓名')
        
        # 姓名
        self.add_paragraph(doc, name, bold=True, size=18, align="center")
        
        # 联系方式（纯文字，不用 emoji 防止 docx 损坏）
        contact_parts = []
        if resume_data.get('phone'): contact_parts.append(f"电话: {resume_data['phone']}")
        if resume_data.get('email'): contact_parts.append(f"邮箱: {resume_data['email']}")
        if resume_data.get('location'): contact_parts.append(f"地址: {resume_data['location']}")
        
        if contact_parts:
            self.add_paragraph(doc, " | ".join(contact_parts), size=10, align="center")
        
        doc.add_paragraph()
        
        content = resume_data.get('content', '')
        
        # 个人简介
        summary_match = re.search(r'##\s*(?:个人简介|简介|Summary)[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if summary_match:
            self.add_heading(doc, "个人简介", level=2)
            for para in summary_match.group(1).strip().split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
            doc.add_paragraph()
        
        # 工作经历
        exp_match = re.search(r'##\s*(?:工作经历|工作经验|Experience)[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if exp_match:
            self.add_heading(doc, "工作经历", level=2)
            for para in exp_match.group(1).strip().split('\n'):
                if para.strip().startswith('-') or para.strip().startswith('•'):
                    self.add_paragraph(doc, para.strip())
                elif para.strip():
                    self.add_paragraph(doc, para.strip())
            doc.add_paragraph()
        
        # 教育背景
        edu_match = re.search(r'##\s*(?:教育背景|教育|Education)[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if edu_match:
            self.add_heading(doc, "教育背景", level=2)
            for para in edu_match.group(1).strip().split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
            doc.add_paragraph()
        
        # 技能特长
        skill_match = re.search(r'##\s*(?:技能特长|专业技能|Skills)[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if skill_match:
            self.add_heading(doc, "技能特长", level=2)
            for para in skill_match.group(1).strip().split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
            doc.add_paragraph()
        
        # 语言能力
        lang_match = re.search(r'##\s*(?:语言能力|语言|Languages)[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if lang_match:
            self.add_heading(doc, "语言能力", level=2)
            for para in lang_match.group(1).strip().split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
        
        return doc
    
    def export_danish_resume(self, resume_data: Dict) -> Document:
        """导出丹麦语简历 (CV)"""
        doc = self.create_document()
        
        name = resume_data.get('name', resume_data.get('content', '').split('\n')[0] if resume_data.get('content') else 'Dit Navn')
        
        # 姓名
        self.add_paragraph(doc, name, bold=True, size=18, align="center")
        
        # 联系方式（纯文字，不用 emoji 防止 docx 损坏）
        contact_parts = []
        if resume_data.get('phone'): contact_parts.append(f"Tlf: {resume_data['phone']}")
        if resume_data.get('email'): contact_parts.append(f"Email: {resume_data['email']}")
        if resume_data.get('location'): contact_parts.append(f"Adresse: {resume_data['location']}")
        
        if contact_parts:
            self.add_paragraph(doc, " | ".join(contact_parts), size=10, align="center")
        
        doc.add_paragraph()
        
        content = resume_data.get('content', '')
        
        # 专业简介
        summary_match = re.search(r'##\s*(?:Professionel\s*)?Profil[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if not summary_match:
            summary_match = re.search(r'##\s*(?:Professional\s*)?Summary[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        
        if summary_match:
            self.add_heading(doc, "Profil", level=2)
            for para in summary_match.group(1).strip().split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
            doc.add_paragraph()
        
        # 工作经历
        exp_match = re.search(r'##\s*(?:Erhvervserfaring|Professionel\s*)?Erfaring[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if not exp_match:
            exp_match = re.search(r'##\s*(?:Professional\s*)?Experience[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        
        if exp_match:
            self.add_heading(doc, "Erhvervserfaring", level=2)
            for para in exp_match.group(1).strip().split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
            doc.add_paragraph()
        
        # 教育背景
        edu_match = re.search(r'##\s*(?:Uddannelse|Education)[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if edu_match:
            self.add_heading(doc, "Uddannelse", level=2)
            for para in edu_match.group(1).strip().split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
            doc.add_paragraph()
        
        # 技能
        skill_match = re.search(r'##\s*(?:Kompetencer|Skills)[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if skill_match:
            self.add_heading(doc, "Kompetencer", level=2)
            for para in skill_match.group(1).strip().split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
            doc.add_paragraph()
        
        # 证书
        cert_match = re.search(r'##\s*(?:Certificeringer|Certifications)[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if cert_match:
            self.add_heading(doc, "Certificeringer", level=2)
            for para in cert_match.group(1).strip().split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
            doc.add_paragraph()
        
        # 语言
        lang_match = re.search(r'##\s*(?:Sprog|Languages)[:]?\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if lang_match:
            self.add_heading(doc, "Sprog", level=2)
            for para in lang_match.group(1).strip().split('\n'):
                if para.strip():
                    self.add_paragraph(doc, para.strip())
        
        return doc
    
    def export_resume(self, resume_data: Dict, language: str = "en") -> Document:
        """导出简历（根据语言选择模板）"""
        if language == "zh":
            return self.export_chinese_resume(resume_data)
        elif language == "da":
            return self.export_danish_resume(resume_data)
        else:
            return self.export_english_resume(resume_data)
    
    def export_to_bytes(self, resume_data: Dict, language: str = "en") -> bytes:
        """导出简历到字节流"""
        doc = self.export_resume(resume_data, language)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


# 全局导出器实例
resume_exporter = ResumeExporter() if DOCX_AVAILABLE else None


def export_user_resume(user_id: str, resume_id: str = "", language: str = "en") -> Optional[bytes]:
    """导出用户简历
    
    Args:
        user_id: 用户ID
        resume_id: 简历ID（为空则使用主要简历）
        language: 语言 (en, zh, da)
    
    Returns:
        Word文档字节流
    """
    if not DOCX_AVAILABLE:
        return None
    
    # 获取简历数据
    if resume_id:
        resume_data = get_resume_by_id(resume_id)
    else:
        resume_data = get_primary_resume(user_id)
    
    if not resume_data:
        return None
    
    # 自动检测语言
    if language == "auto":
        content = resume_data.get('content', '')
        # 简单检测
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', content))
        danish_chars = len(re.findall(r'[æøåÆØÅ]', content))
        
        if chinese_chars > 10:
            language = "zh"
        elif danish_chars > 5:
            language = "da"
        else:
            language = "en"
    
    # 导出
    return resume_exporter.export_to_bytes(resume_data, language)


def get_resume_for_export(user_id: str, resume_id: str = "") -> Optional[Dict]:
    """获取用于导出的简历数据"""
    if resume_id:
        return get_resume_by_id(resume_id)
    return get_primary_resume(user_id)
